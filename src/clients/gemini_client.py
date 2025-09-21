import asyncio
import json
import os
from typing import Dict, List, AsyncGenerator

# Use google-generativeai package
import google.generativeai as genai
from src.utils.logger import setup_logger
import requests

logger = setup_logger()


class GeminiClient:
    def __init__(self, api_key: str):
        """
        Initialize Gemini client with google-generativeai package.
        """
        self.api_key = api_key

        # Configure API key
        genai.configure(api_key=api_key)
        self.client = genai
        self.logger = logger

        # Default models
        self.text_model = "gemini-1.5-flash"
        self.vision_model = "gemini-1.5-flash"

        logger.info("✅ Gemini client initialized with google-generativeai package")

    async def chat_completion(
        self, messages: List[Dict], temperature: float = 0.7
    ) -> str:
        """
        Basic chat completion without streaming using google-generativeai.
        """
        try:
            # Convert messages to Gemini format
            contents = []
            for msg in messages:
                role = "user" if msg["role"] in ["user", "system"] else "model"
                contents.append({
                    "role": role,
                    "parts": [{"text": msg["content"]}]
                })

            # Generate content using correct format
            model = genai.GenerativeModel(self.text_model)
            response = model.generate_content(contents)

            return response.text

        except Exception as e:
            self.logger.error(f"❌ Gemini chat completion error: {e}")
            raise

    async def chat_completion_stream(
        self, messages: List[Dict], temperature: float = 0.7
    ) -> AsyncGenerator[str, None]:
        """
        Streaming chat completion using new API.
        """
        try:
            # Convert messages to Gemini format
            contents = []
            for msg in messages:
                role = "user" if msg["role"] in ["user", "system"] else "model"
                contents.append({
                    "role": role,
                    "parts": [{"text": msg["content"]}]
                })

            # Generate streaming content
            model = genai.GenerativeModel(self.text_model)
            response_stream = model.generate_content(
                contents,
                stream=True
            )

            for chunk in response_stream:
                if hasattr(chunk, 'text') and chunk.text:
                    yield chunk.text
                elif hasattr(chunk, 'candidates') and chunk.candidates:
                    for candidate in chunk.candidates:
                        if hasattr(candidate, 'content') and candidate.content:
                            for part in candidate.content.parts:
                                if hasattr(part, 'text') and part.text:
                                    yield part.text

        except Exception as e:
            self.logger.error(f"❌ Gemini chat completion stream error: {e}")
            yield f"Lỗi Gemini: {str(e)}"

    async def upload_file_and_analyze(
        self, file_content: bytes, file_name: str, prompt: str = ""
    ) -> str:
        """
        Upload file to Gemini and analyze it.
        Supports: PDF (direct upload) and DOCX (text extraction).
        Frontend should limit file uploads to these formats only.
        """
        try:
            original_file_name = file_name
            processed_file_content = file_content

            # Handle different file types - only PDF and DOCX supported
            file_extension = file_name.lower().split('.')[-1] if '.' in file_name else ''

            # 1. Handle DOCX files - convert to PDF
            if file_extension == 'docx':
                logger.info("📄 DOCX file detected, converting to PDF...")
                processed_file_content, file_name, mime_type = await self._convert_docx_to_pdf(file_content, file_name)

            # 2. Handle PDF files directly
            elif file_extension == 'pdf':
                logger.info("� PDF file detected, uploading directly...")
                mime_type = 'application/pdf'

            else:
                # Unsupported format - frontend should prevent this
                logger.error(f"❌ Unsupported file format: {file_extension}")
                raise Exception(f"File format '{file_extension}' not supported. Please use PDF or DOCX files only.")

            # Create temporary file for Gemini upload
            import tempfile
            import os

            with tempfile.NamedTemporaryFile(
                delete=False, suffix=f"_{file_name}", mode='wb'
            ) as tmp_file:
                tmp_file.write(processed_file_content)
                tmp_file_path = tmp_file.name

            try:
                # Upload file using google-generativeai
                uploaded_file = genai.upload_file(
                    path=tmp_file_path,
                    mime_type=mime_type,
                    display_name=file_name
                )
                logger.info(
                    f"✅ Gemini file upload successful: {file_name} ({mime_type})"
                )

                # Wait for file processing
                import time
                timeout_start = time.time()
                timeout = 90  # Increased timeout for larger files

                while uploaded_file.state.name == "PROCESSING":
                    if time.time() - timeout_start > timeout:
                        raise Exception("File processing timeout")
                    time.sleep(2)
                    # Refresh file status
                    uploaded_file = self.client.files.get(uploaded_file.name)

                if uploaded_file.state.name != "ACTIVE":
                    raise Exception(
                        f"File processing failed with state: {uploaded_file.state.name}"
                    )

                logger.info(f"✅ File processed successfully: {uploaded_file.name}")

                # Generate content with uploaded file using new API
                if not prompt:
                    prompt = f"Hãy phân tích và trích xuất thông tin từ file {original_file_name}."

                response = self.client.models.generate_content(
                    model=self.vision_model,
                    contents=[prompt, uploaded_file]  # Simple format that works
                )

                # Clean up uploaded file
                try:
                    self.client.files.delete(name=uploaded_file.name)
                    logger.info("✅ Uploaded file cleaned up successfully")
                except Exception as cleanup_error:
                    logger.warning(
                        f"⚠️ Could not cleanup uploaded file: {cleanup_error}"
                    )

                return response.text

            except Exception as upload_error:
                logger.error(f"❌ Gemini file upload failed: {upload_error}")

                # FALLBACK: Use ChatGPT instead of Gemini
                logger.info("🔄 Falling back to ChatGPT for file processing...")
                raise Exception(
                    f"Gemini upload failed, fallback to ChatGPT: {upload_error}"
                )

            finally:
                # Clean up temporary file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)

        except Exception as e:
            self.logger.error(f"❌ Gemini file upload error: {e}")
            raise

    async def _convert_docx_to_pdf(self, file_content: bytes, file_name: str):
        """Convert DOCX file to PDF using python-docx + reportlab"""
        try:
            from docx import Document
            import tempfile
            import os
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch

            # Create temporary DOCX file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                tmp_docx.write(file_content)
                tmp_docx_path = tmp_docx.name

            # Create temporary PDF path
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_pdf:
                tmp_pdf_path = tmp_pdf.name

            try:
                # Read DOCX document
                doc = Document(tmp_docx_path)

                # Create PDF document
                pdf_doc = SimpleDocTemplate(tmp_pdf_path, pagesize=A4)
                styles = getSampleStyleSheet()
                elements = []

                # Custom style for better formatting
                normal_style = ParagraphStyle(
                    'CustomNormal',
                    parent=styles['Normal'],
                    fontSize=10,
                    spaceAfter=6,
                    leftIndent=0.25*inch,
                    rightIndent=0.25*inch
                )

                # Add document title
                title = Paragraph(f"Document: {file_name}", styles['Title'])
                elements.append(title)
                elements.append(Spacer(1, 12))

                # Process paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        # Clean and format text
                        text = paragraph.text.strip()
                        # Escape special characters for reportlab
                        text = text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

                        para = Paragraph(text, normal_style)
                        elements.append(para)
                        elements.append(Spacer(1, 3))

                # Process tables if any
                for table in doc.tables:
                    # Add table title
                    table_title = Paragraph("Table:", styles['Heading2'])
                    elements.append(table_title)
                    elements.append(Spacer(1, 6))

                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                        if row_text:
                            row_text = row_text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                            para = Paragraph(row_text, normal_style)
                            elements.append(para)
                            elements.append(Spacer(1, 3))

                    elements.append(Spacer(1, 12))

                # Build PDF
                pdf_doc.build(elements)

                # Read the generated PDF
                with open(tmp_pdf_path, 'rb') as pdf_file:
                    pdf_content = pdf_file.read()

                processed_name = file_name.replace('.docx', '.pdf')

                logger.info(f"✅ DOCX to PDF conversion completed: {len(pdf_content)} bytes")

                return pdf_content, processed_name, 'application/pdf'

            finally:
                # Clean up temporary files
                if os.path.exists(tmp_docx_path):
                    os.unlink(tmp_docx_path)
                if os.path.exists(tmp_pdf_path):
                    os.unlink(tmp_pdf_path)

        except ImportError as e:
            logger.warning(f"⚠️ Required packages not available: {e}, falling back to text extraction")
            # Fallback to text extraction if packages are not available
            return await self._process_docx_file(file_content, file_name)
        except Exception as docx_error:
            logger.error(f"❌ DOCX to PDF conversion failed: {docx_error}")
            logger.info("🔄 Falling back to text extraction...")
            # Fallback to text extraction if conversion fails
            return await self._process_docx_file(file_content, file_name)

    async def _process_docx_file(self, file_content: bytes, file_name: str):
        """Extract text from DOCX file"""
        try:
            from docx import Document
            import tempfile
            import os

            # Create temporary DOCX file
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
                tmp_docx.write(file_content)
                tmp_docx_path = tmp_docx.name

            try:
                # Extract text from DOCX
                doc = Document(tmp_docx_path)
                extracted_text = ""

                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        extracted_text += paragraph.text.strip() + "\n"

                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                extracted_text += cell.text.strip() + "\t"
                        extracted_text += "\n"

                logger.info(f"✅ Extracted {len(extracted_text)} characters from DOCX")

                # Convert extracted text to bytes
                processed_content = extracted_text.encode('utf-8')
                processed_name = file_name.replace('.docx', '_extracted.txt')

                return processed_content, processed_name, 'text/plain'

            finally:
                if os.path.exists(tmp_docx_path):
                    os.unlink(tmp_docx_path)

        except ImportError:
            logger.warning("⚠️ python-docx not available")
            raise Exception("DOCX processing requires python-docx package")
        except Exception as docx_error:
            logger.error(f"❌ DOCX extraction failed: {docx_error}")
            raise Exception(f"Failed to extract text from DOCX: {docx_error}")

    async def chat_with_file_stream(
        self,
        messages: List[Dict],
        file_content: bytes,
        file_name: str,
        temperature: float = 0.7,
    ) -> AsyncGenerator[str, None]:
        """
        Stream chat with file upload.
        """
        try:
            # Upload and analyze file first
            file_analysis = await self.upload_file_and_analyze(
                file_content, file_name, "Hãy phân tích file này."
            )

            # Add file analysis to conversation
            enhanced_messages = messages.copy()
            enhanced_messages.append(
                {
                    "role": "system",
                    "content": f"Thông tin từ file {file_name}:\n{file_analysis}",
                }
            )

            # Stream response with enhanced context
            async for chunk in self.chat_completion_stream(
                enhanced_messages, temperature
            ):
                yield chunk

        except Exception as e:
            self.logger.error(f"❌ Gemini chat with file stream error: {e}")
            yield f"Lỗi xử lý file: {str(e)}"
