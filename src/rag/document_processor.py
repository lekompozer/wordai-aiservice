import json
import re
from pathlib import Path
from typing import List, Dict, Optional
from PyPDF2 import PdfReader
import docx
import os

from src.utils.logger import setup_logger

logger = setup_logger()

from typing import Any, Dict

class Document:
    """Lớp Document đại diện cho một đoạn văn bản đã xử lý với metadata"""
    def __init__(self, content: str, metadata: Dict[str, str]):
        self.content = content
        self.metadata = metadata
       
    def __getitem__(self, key):
        if isinstance(key, slice):
            # Handle slice operations for content
            return self.content[key]
        elif key == "content":
            return self.content
        elif key in self.metadata:
            return self.metadata[key]
        else:
            raise KeyError(f"Key '{key}' not found in Document")

    def __str__(self) -> str:
        return f"Document(content={self.content[:50]}..., metadata={self.metadata})"

class DocumentProcessor:
    def __init__(self):
        self.tone_config = self._load_tone_config()
        self.supported_extensions = {".pdf", ".docx", ".txt"}
        self.min_chunk_length = 100
        self.max_chunk_length = 1000  # Độ dài tối đa của một chunk

    def _load_tone_config(self):
        config_path = (
            Path(__file__).parent.parent.parent / "config" / "tone_config.json"
        )
        with open(config_path, "r", encoding="utf-8") as f:
            return json.load(f)

    def process_folder(self, folder_path: str) -> Dict[str, List[Document]]:
        """Xử lý toàn bộ thư mục và trả về dict {filename: document_objects}"""
        if not os.path.isdir(folder_path):
            raise NotADirectoryError(f"Folder not found: {folder_path}")

        results = {}
        for filepath in Path(folder_path).iterdir():
            if filepath.suffix.lower() in self.supported_extensions:
                try:
                    documents = self.process_file(str(filepath))
                    if documents:
                        results[filepath.name] = documents
                        logger.info(f"Processed {filepath.name}: {len(documents)} documents")
                    else:
                        logger.warning(f"Skipped {filepath.name} (no valid content)")
                except Exception as e:
                    logger.error(f"Error processing {filepath.name}: {str(e)}")
        return results

    def process_file(self, file_path: str) -> List[Document]:
        """Xử lý từng file và trả về các Document object đã chuẩn hóa"""
        if not Path(file_path).exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        content = self._read_file_content(file_path)
        filename = Path(file_path).name
        filetype = Path(file_path).suffix.lower()
        file_size = os.path.getsize(file_path)
       
        # Metadata cơ bản cho tất cả các documents
        base_metadata = {
            "source": filename,
            "filetype": filetype,
            "file_size": str(file_size),
            "processed_date": str(Path(file_path).stat().st_mtime)
        }
       
        return self._chunk_content(content, base_metadata)

    def _read_file_content(self, file_path: str) -> str:
        """Đọc nội dung file theo định dạng"""
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            with open(file_path, "rb") as f:
                return "\n".join(
                    page.extract_text()
                    for page in PdfReader(f).pages
                    if page.extract_text()
                )
        elif ext == ".docx":
            doc = docx.Document(file_path)
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        elif ext == ".txt":
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        else:
            raise ValueError(f"Unsupported file type: {ext}")

    def _chunk_content(self, content: str, base_metadata: Dict[str, str]) -> List[Document]:
        """Chia nội dung thành các Document objects có độ dài phù hợp"""
        content = re.sub(r"\s+", " ", content).strip()
        paragraphs = [p.strip() for p in re.split(r'\n\n+|\*{3,}', content) if p.strip()]

        documents = []
        current_chunk = ""
        chunk_index = 0

        for para in paragraphs:
            if len(current_chunk) + len(para) <= self.max_chunk_length:
                current_chunk += "\n\n" + para if current_chunk else para
            else:
                if len(current_chunk) >= self.min_chunk_length:
                    extra_meta = {"chunk_index": str(chunk_index)}
                    document = self._create_document(
                        current_chunk,
                        base_metadata,
                        extra_meta
                    )
                    # Log chỉ 50 ký tự đầu tiên để giảm bộ nhớ và tránh trùng lặp log
                    logger.debug(f"Chunk {chunk_index}: {document.content[:50]}...")
                    documents.append(document)
                    chunk_index += 1
                current_chunk = para

        if current_chunk and len(current_chunk) >= self.min_chunk_length:
            extra_meta = {"chunk_index": str(chunk_index)}
            document = self._create_document(
                current_chunk,
                base_metadata,
                extra_meta
            )
            # Log chỉ 50 ký tự đầu tiên
            logger.debug(f"Chunk {chunk_index}: {document.content[:50]}...")
            documents.append(document) 
        logger.info(f"Created {len(documents)} chunks from content")    
        return documents
   
    def _create_document(self, content: str, base_metadata: Dict[str, str],
                         extra_metadata: Dict[str, str] = None) -> Document:
        """Helper để tạo Document object với metadata"""
        metadata = base_metadata.copy()
        if extra_metadata:
            metadata.update(extra_metadata)
        # Thêm kích thước để theo dõi bộ nhớ được sử dụng
        metadata["content_length"] = str(len(content))
        return Document(content=content, metadata=metadata)
