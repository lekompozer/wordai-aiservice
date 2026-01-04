"""
Document Export Service
Convert HTML documents to various formats (PDF, DOCX, TXT, HTML)
"""

import os
import uuid
import logging
import tempfile
from typing import Optional, Tuple, List, Dict
from datetime import datetime, timedelta
from io import BytesIO

logger = logging.getLogger("chatbot")


class DocumentExportService:
    """
    Service để export HTML documents sang các format khác
    - PDF: weasyprint
    - DOCX: python-docx + htmldocx
    - TXT: BeautifulSoup
    - HTML: Raw HTML file
    """

    def __init__(self, r2_client, db):
        """
        Initialize DocumentExportService

        Args:
            r2_client: R2Client instance for file upload
            db: MongoDB database instance
        """
        self.r2_client = r2_client
        self.db = db
        self.exports = db["document_exports"]  # Track export history for rate limiting

        # Create indexes for rate limiting
        self._create_indexes()

    def _create_indexes(self):
        """Create indexes for export tracking collection"""
        try:
            # Index for user + document (per-document cooldown)
            self.exports.create_index(
                [("user_id", 1), ("document_id", 1), ("exported_at", -1)],
                name="user_doc_export_idx",
            )
            # Index for user exports (global rate limit)
            self.exports.create_index(
                [("user_id", 1), ("exported_at", -1)], name="user_export_history_idx"
            )
            logger.info("✅ Export indexes created")
        except Exception as e:
            logger.warning(f"⚠️ Export index creation warning: {e}")

    def check_rate_limits(self, user_id: str, document_id: str) -> Tuple[bool, str]:
        """
        Check if user can export document based on rate limits

        Rate Limits:
        1. Per-document cooldown: 15 seconds between exports of same document
        2. Global limit: Max 10 exports in 30 minutes

        Args:
            user_id: User's Firebase UID
            document_id: Document ID

        Returns:
            (can_export: bool, error_message: str)
        """
        now = datetime.utcnow()

        # Check 1: Per-document cooldown (15 seconds)
        last_export = self.exports.find_one(
            {
                "user_id": user_id,
                "document_id": document_id,
                "exported_at": {"$gte": now - timedelta(seconds=15)},
            },
            sort=[("exported_at", -1)],
        )

        if last_export:
            seconds_ago = (now - last_export["exported_at"]).total_seconds()
            wait_time = 15 - int(seconds_ago)
            return (
                False,
                f"Please wait {wait_time} seconds before exporting this document again",
            )

        # Check 2: Global rate limit (10 exports in 30 minutes)
        exports_count = self.exports.count_documents(
            {"user_id": user_id, "exported_at": {"$gte": now - timedelta(minutes=30)}}
        )

        if exports_count >= 10:
            return (
                False,
                "Export limit reached: Maximum 10 exports per 30 minutes. Please try again later.",
            )

        return True, ""

    def track_export(self, user_id: str, document_id: str, format: str, file_size: int):
        """
        Track export in database for rate limiting

        Args:
            user_id: User's Firebase UID
            document_id: Document ID
            format: Export format (pdf/docx/txt/html)
            file_size: Exported file size in bytes
        """
        try:
            self.exports.insert_one(
                {
                    "export_id": f"exp_{uuid.uuid4().hex[:12]}",
                    "user_id": user_id,
                    "document_id": document_id,
                    "format": format,
                    "file_size": file_size,
                    "exported_at": datetime.utcnow(),
                }
            )
            logger.info(f"📊 Tracked export: {document_id} → {format}")
        except Exception as e:
            # Don't fail export if tracking fails
            logger.error(f"❌ Failed to track export: {e}")

    def _sanitize_filename(self, title: str) -> str:
        """
        Sanitize document title for filename
        Remove special characters, limit length
        """
        import re

        # Remove special characters
        filename = re.sub(r'[<>:"/\\|?*]', "", title)
        # Replace spaces with underscores
        filename = filename.replace(" ", "_")
        # Limit length
        filename = filename[:50] if len(filename) > 50 else filename
        # Add timestamp
        timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
        return f"{filename}_{timestamp}"

    async def export_to_pdf_playwright(
        self, html_content: str, title: str = "document", document_type: str = "doc"
    ) -> Tuple[bytes, str]:
        """
        Convert HTML to PDF using Playwright (Chromium rendering engine)

        **Best for slide presentations** - provides browser-quality rendering

        Args:
            html_content: HTML content to convert
            title: Document title for filename
            document_type: Document type ("doc", "slide", "note") for page sizing

        Returns:
            (pdf_bytes, filename)
        """
        try:
            from playwright.async_api import async_playwright
            import tempfile
            import os

            # Determine page size and styling
            if document_type == "slide":
                width = "1920px"
                height = "1080px"
                landscape = True

                # Enhanced CSS for slide presentations
                enhanced_css = """
                <style>
                * {
                    box-sizing: border-box;
                    -webkit-print-color-adjust: exact;
                    print-color-adjust: exact;
                }

                body {
                    margin: 0;
                    padding: 0;
                    width: 1920px;
                    height: 1080px;
                    overflow: hidden;
                    background: white;
                }

                .slide {
                    width: 1920px !important;
                    height: 1080px !important;
                    max-width: none !important;
                    margin: 0 !important;
                    padding: 60px !important;
                    box-sizing: border-box;
                    page-break-after: always;
                    page-break-inside: avoid;
                    position: relative;
                    overflow: hidden;
                }

                /* Enhanced typography for slides */
                .slide h1 {
                    font-size: 64px !important;
                    font-weight: bold;
                    margin-bottom: 30px;
                    line-height: 1.2;
                }

                .slide h2 {
                    font-size: 48px !important;
                    margin-bottom: 25px;
                    line-height: 1.3;
                }

                .slide h3 {
                    font-size: 36px !important;
                    margin-bottom: 20px;
                    line-height: 1.4;
                }

                .slide p, .slide li {
                    font-size: 28px !important;
                    line-height: 1.6;
                    margin-bottom: 15px;
                }

                .slide ul, .slide ol {
                    font-size: 28px !important;
                    line-height: 1.6;
                }

                /* Table styling */
                .slide table {
                    font-size: 24px !important;
                    border-collapse: collapse;
                }

                .slide th, .slide td {
                    padding: 12px !important;
                    border: 1px solid #ddd;
                }

                /* Overlay elements - absolute positioning */
                .overlay-textbox {
                    position: absolute !important;
                    display: flex;
                    align-items: center;
                    word-wrap: break-word;
                    overflow-wrap: break-word;
                    white-space: pre-wrap;
                }

                .overlay-image {
                    position: absolute !important;
                    object-fit: cover;
                }

                .overlay-shape {
                    position: absolute !important;
                }

                .overlay-video-placeholder {
                    position: absolute !important;
                    display: flex;
                    align-items: center;
                    justify-content: center;
                    background: #1a1a1a;
                    border-radius: 8px;
                }

                /* YouTube embed iframe */
                .youtube-embed {
                    position: absolute !important;
                    border: none;
                }

                /* Print optimization */
                @page {
                    size: 1920px 1080px;
                    margin: 0;
                }

                @media print {
                    body {
                        width: 1920px;
                        height: 1080px;
                    }
                    .slide {
                        page-break-after: always;
                        page-break-inside: avoid;
                    }
                }
                </style>
                """
                logger.info(f"📄 Using Playwright with FullHD (1920x1080) for slides")

            else:
                # A4 for documents and notes
                width = "210mm"
                height = "297mm"
                landscape = False

                enhanced_css = """
                <style>
                @page {
                    size: A4;
                    margin: 20mm;
                }

                * {
                    -webkit-print-color-adjust: exact;
                    print-color-adjust: exact;
                }

                body {
                    font-family: Arial, sans-serif;
                    font-size: 12pt;
                    line-height: 1.6;
                    color: #333;
                    margin: 0;
                    padding: 0;
                }

                h1 { font-size: 24pt; margin-bottom: 12pt; }
                h2 { font-size: 20pt; margin-bottom: 10pt; }
                h3 { font-size: 16pt; margin-bottom: 8pt; }
                p { margin-bottom: 8pt; }
                img { max-width: 100%; height: auto; }
                table { border-collapse: collapse; width: 100%; }
                th, td { border: 1px solid #ddd; padding: 8px; }
                </style>
                """
                logger.info(f"📄 Using Playwright with A4 for {document_type}")

            # Wrap HTML with proper structure
            full_html = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    {enhanced_css}
</head>
<body>
{html_content}
</body>
</html>"""

            # Create temp file for PDF output
            temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            temp_pdf_path = temp_pdf.name
            temp_pdf.close()

            try:
                async with async_playwright() as p:
                    # Launch Chromium browser (headless)
                    browser = await p.chromium.launch(headless=True)
                    page = await browser.new_page()

                    # Set viewport for slides
                    if document_type == "slide":
                        await page.set_viewport_size({"width": 1920, "height": 1080})

                    # Load HTML content
                    await page.set_content(full_html, wait_until="networkidle")

                    # Wait for any images/resources to load
                    await page.wait_for_timeout(1000)  # 1 second delay for rendering

                    # Generate PDF
                    await page.pdf(
                        path=temp_pdf_path,
                        format=None if document_type == "slide" else "A4",
                        width=width if document_type == "slide" else None,
                        height=height if document_type == "slide" else None,
                        landscape=landscape,
                        margin=(
                            {"top": "0", "right": "0", "bottom": "0", "left": "0"}
                            if document_type == "slide"
                            else {
                                "top": "20mm",
                                "right": "20mm",
                                "bottom": "20mm",
                                "left": "20mm",
                            }
                        ),
                        print_background=True,  # Include background colors/images
                        prefer_css_page_size=True,
                    )

                    await browser.close()

                # Read PDF bytes
                with open(temp_pdf_path, "rb") as f:
                    pdf_bytes = f.read()

                filename = f"{self._sanitize_filename(title)}.pdf"

                logger.info(
                    f"✅ Generated PDF with Playwright: {filename} "
                    f"({len(pdf_bytes)} bytes, {document_type} format)"
                )

                return pdf_bytes, filename

            finally:
                # Clean up temp file
                if os.path.exists(temp_pdf_path):
                    os.unlink(temp_pdf_path)

        except Exception as e:
            logger.error(f"❌ Error generating PDF with Playwright: {e}")
            raise Exception(f"PDF generation failed: {str(e)}")

    def _sanitize_html_for_weasyprint(self, html_content: str) -> str:
        """
        Sanitize HTML for WeasyPrint compatibility
        Remove problematic CSS while preserving content structure

        Args:
            html_content: Original HTML content

        Returns:
            Sanitized HTML content
        """
        from bs4 import BeautifulSoup
        import re

        try:
            # Parse HTML
            soup = BeautifulSoup(html_content, "html.parser")

            # Remove all style attributes that might contain unsupported CSS
            for tag in soup.find_all(style=True):
                style_str = tag.get("style", "")

                # List of problematic CSS properties to remove
                problematic_patterns = [
                    r"lab\([^)]+\)",  # lab() colors
                    r"text-decoration-thickness:[^;]+;?",
                    r"margin-inline:[^;]+;?",
                    r"transform:[^;]+;?",
                    r"-webkit-[^:]+:[^;]+;?",  # webkit prefixes
                    r"-moz-[^:]+:[^;]+;?",  # mozilla prefixes
                ]

                # Clean the style string
                for pattern in problematic_patterns:
                    style_str = re.sub(pattern, "", style_str)

                # Replace lab() colors in remaining style
                style_str = re.sub(r"lab\([^)]+\)", "rgb(0,0,0)", style_str)

                # Clean up multiple semicolons and spaces
                style_str = re.sub(r";+", ";", style_str)
                style_str = re.sub(r";\s*;", ";", style_str)
                style_str = style_str.strip(";").strip()

                # Update or remove style attribute
                if style_str:
                    tag["style"] = style_str
                else:
                    del tag["style"]

            # Also clean <style> tags
            for style_tag in soup.find_all("style"):
                css_content = style_tag.string or ""
                for pattern in [
                    r"lab\([^)]+\)",
                    r"text-decoration-thickness:[^;]+;?",
                    r"margin-inline:[^;]+;?",
                    r"transform:[^;]+;?",
                ]:
                    css_content = re.sub(pattern, "", css_content)
                css_content = re.sub(r"lab\([^)]+\)", "rgb(0,0,0)", css_content)
                style_tag.string = css_content

            return str(soup)

        except Exception as e:
            logger.warning(
                f"⚠️ Error sanitizing HTML with BeautifulSoup: {e}, using basic regex fallback"
            )
            # Fallback to basic replacement
            html_content = re.sub(r"lab\([^)]+\)", "rgb(0,0,0)", html_content)
            return html_content

    def export_to_pdf(
        self, html_content: str, title: str = "document", document_type: str = "doc"
    ) -> Tuple[bytes, str]:
        """
        Convert HTML to PDF using weasyprint (legacy method)

        Args:
            html_content: HTML content to convert
            title: Document title for filename
            document_type: Document type ("doc", "slide", "note") for page sizing

        Returns:
            (pdf_bytes, filename)
        """
        try:
            from weasyprint import HTML, CSS

            # ✅ Sanitize HTML for WeasyPrint compatibility
            html_content = self._sanitize_html_for_weasyprint(html_content)
            logger.info("🧹 HTML sanitized for WeasyPrint compatibility")

            # Choose page size based on document type
            if document_type == "slide":
                # FullHD 1920x1080 (16:9 landscape) for slides
                page_size = "1920px 1080px"
                page_margin = "0"  # No margin for slides (full bleed)
                logger.info(f"📄 Using FullHD (1920x1080) page size for slide document")
            else:
                # A4 portrait for docs and notes
                page_size = "A4"
                page_margin = "20mm"
                logger.info(f"📄 Using A4 page size for {document_type} document")

            # Add CSS for better PDF rendering
            css = f"""
            @page {{
                size: {page_size};
                margin: {page_margin};
            }}
            body {{
                font-family: Arial, sans-serif;
                font-size: 12pt;
                line-height: 1.6;
                color: #333;
            }}
            h1 {{ font-size: 24pt; margin-bottom: 12pt; }}
            h2 {{ font-size: 20pt; margin-bottom: 10pt; }}
            h3 {{ font-size: 16pt; margin-bottom: 8pt; }}
            p {{ margin-bottom: 8pt; }}
            img {{ max-width: 100%; height: auto; }}
            table {{ border-collapse: collapse; width: 100%; }}
            th, td {{ border: 1px solid #ddd; padding: 8px; }}
            """

            html_with_style = f"<style>{css}</style>{html_content}"

            # Generate PDF in memory
            pdf_file = BytesIO()
            HTML(string=html_with_style).write_pdf(pdf_file)
            pdf_bytes = pdf_file.getvalue()

            filename = f"{self._sanitize_filename(title)}.pdf"

            logger.info(f"✅ Generated PDF: {filename} ({len(pdf_bytes)} bytes)")
            return pdf_bytes, filename

        except Exception as e:
            logger.error(f"❌ Error generating PDF: {e}")
            raise Exception(f"PDF generation failed: {str(e)}")

    def export_to_docx(
        self, html_content: str, title: str = "document"
    ) -> Tuple[bytes, str]:
        """
        Convert HTML to DOCX using Pandoc (via pypandoc)

        Args:
            html_content: HTML content to convert
            title: Document title for filename

        Returns:
            (docx_bytes, filename)
        """
        try:
            import pypandoc

            # Add proper HTML structure if missing
            if not html_content.strip().startswith(
                "<!DOCTYPE"
            ) and not html_content.strip().startswith("<html"):
                html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>{title}</title>
</head>
<body>
{html_content}
</body>
</html>"""

            # Convert HTML to DOCX using Pandoc
            # Pandoc is much more powerful than htmldocx and handles complex HTML better
            # For docx format, we need to write to a temporary file
            import tempfile
            import os

            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_file:
                tmp_path = tmp_file.name

            try:
                # Convert HTML to DOCX file
                pypandoc.convert_text(
                    html_content,
                    "docx",
                    format="html",
                    outputfile=tmp_path,
                    extra_args=[
                        f"--metadata=title:{title}",  # Set document title
                    ],
                )

                # Post-process DOCX to add table borders
                # Pandoc doesn't preserve HTML table borders, so we add them manually
                self._add_table_borders_to_docx(tmp_path)

                # Read the generated DOCX file
                with open(tmp_path, "rb") as f:
                    docx_bytes = f.read()

            finally:
                # Clean up temporary file
                if os.path.exists(tmp_path):
                    os.remove(tmp_path)

            filename = f"{self._sanitize_filename(title)}.docx"

            logger.info(
                f"✅ Generated DOCX with Pandoc: {filename} ({len(docx_bytes)} bytes)"
            )
            return docx_bytes, filename

        except Exception as e:
            logger.error(f"❌ Error generating DOCX with Pandoc: {e}")
            raise Exception(f"DOCX generation failed: {str(e)}")

    def _add_table_borders_to_docx(self, docx_path: str) -> None:
        """
        Add borders to all tables in a DOCX file.
        Pandoc doesn't preserve HTML table borders, so we add them manually.

        Args:
            docx_path: Path to the DOCX file to modify
        """
        try:
            from docx import Document
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            # Open the document
            doc = Document(docx_path)

            # Process each table
            for table in doc.tables:
                # Set table borders
                tbl = table._element
                tblPr = tbl.tblPr
                if tblPr is None:
                    tblPr = OxmlElement("w:tblPr")
                    tbl.insert(0, tblPr)

                # Create table borders element
                tblBorders = OxmlElement("w:tblBorders")

                # Define border style (single line, 1pt, black)
                border_attrs = {
                    "w:val": "single",
                    "w:sz": "4",  # 4/8 = 0.5pt
                    "w:space": "0",
                    "w:color": "000000",
                }

                # Add all borders
                for border_name in [
                    "top",
                    "left",
                    "bottom",
                    "right",
                    "insideH",
                    "insideV",
                ]:
                    border = OxmlElement(f"w:{border_name}")
                    for attr, value in border_attrs.items():
                        border.set(qn(attr), value)
                    tblBorders.append(border)

                # Remove existing borders if any
                existing_borders = tblPr.find(qn("w:tblBorders"))
                if existing_borders is not None:
                    tblPr.remove(existing_borders)

                # Add new borders
                tblPr.append(tblBorders)

            # Save the modified document
            doc.save(docx_path)
            logger.debug(f"✅ Added borders to {len(doc.tables)} tables in DOCX")

        except Exception as e:
            # Don't fail the entire export if border addition fails
            logger.warning(f"⚠️ Could not add table borders to DOCX: {e}")

    def export_to_txt(
        self, html_content: str, title: str = "document"
    ) -> Tuple[bytes, str]:
        """
        Convert HTML to plain text using BeautifulSoup

        Args:
            html_content: HTML content to convert
            title: Document title for filename

        Returns:
            (txt_bytes, filename)
        """
        try:
            from bs4 import BeautifulSoup
            import re

            # Parse HTML
            soup = BeautifulSoup(html_content, "html.parser")

            # Remove script and style tags
            for tag in soup(["script", "style"]):
                tag.decompose()

            # Handle line breaks and paragraphs
            for br in soup.find_all("br"):
                br.replace_with("\n")

            for p in soup.find_all("p"):
                p.append("\n\n")

            # Extract text with proper formatting
            text = soup.get_text()

            # Clean up whitespace
            # Replace multiple spaces with single space
            text = re.sub(r" +", " ", text)
            # Replace multiple newlines with double newline (paragraph break)
            text = re.sub(r"\n\s*\n+", "\n\n", text)
            # Remove leading/trailing whitespace on each line
            text = "\n".join(line.strip() for line in text.split("\n"))
            # Remove empty lines at start/end
            text = text.strip()

            # Add UTF-8 BOM for better compatibility with Windows/Notepad
            # BOM helps Windows recognize the file as UTF-8
            txt_bytes = "\ufeff".encode("utf-8") + text.encode("utf-8")

            filename = f"{self._sanitize_filename(title)}.txt"

            logger.info(f"✅ Generated TXT: {filename} ({len(txt_bytes)} bytes)")
            return txt_bytes, filename

        except Exception as e:
            logger.error(f"❌ Error generating TXT: {e}")
            raise Exception(f"TXT generation failed: {str(e)}")

    def export_to_html(
        self, html_content: str, title: str = "document"
    ) -> Tuple[bytes, str]:
        """
        Save HTML as standalone file with CSS

        Args:
            html_content: HTML content to save
            title: Document title for filename

        Returns:
            (html_bytes, filename)
        """
        try:
            # Wrap HTML in complete document structure
            html_template = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, 'Helvetica Neue', Arial, sans-serif;
            line-height: 1.6;
            max-width: 800px;
            margin: 40px auto;
            padding: 20px;
            color: #333;
        }}
        h1, h2, h3, h4, h5, h6 {{ margin-top: 24px; margin-bottom: 16px; }}
        p {{ margin-bottom: 16px; }}
        img {{ max-width: 100%; height: auto; }}
        table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
        th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
        th {{ background-color: #f2f2f2; }}
    </style>
</head>
<body>
{html_content}
</body>
</html>"""

            html_bytes = html_template.encode("utf-8")

            filename = f"{self._sanitize_filename(title)}.html"

            logger.info(f"✅ Generated HTML: {filename} ({len(html_bytes)} bytes)")
            return html_bytes, filename

        except Exception as e:
            logger.error(f"❌ Error generating HTML: {e}")
            raise Exception(f"HTML generation failed: {str(e)}")

    async def export_and_upload(
        self,
        user_id: str,
        document_id: str,
        html_content: str,
        title: str,
        format: str,
        document_type: str = "doc",  # Add document_type parameter
    ) -> dict:
        """
        Export document to specified format and upload to R2

        Args:
            user_id: User Firebase UID
            document_id: Document ID
            html_content: HTML content to export
            title: Document title
            format: Export format ("pdf", "docx", "txt", "html")
            document_type: Document type ("doc", "slide", "note") for proper page sizing

        Returns:
            {
                "download_url": "https://r2.wordai.pro/...",
                "filename": "document_20251009_153045.pdf",
                "file_size": 123456,
                "expires_in": 3600,
                "expires_at": "2025-10-09T16:30:45Z"
            }
        """
        try:
            # Generate file based on format
            if format == "pdf":
                # ✅ FIX: Use Playwright for ALL document types
                # WeasyPrint has compatibility issues with modern HTML/CSS
                logger.info(
                    f"🎬 Using Playwright for PDF export (document_type={document_type})"
                )
                file_bytes, filename = await self.export_to_pdf_playwright(
                    html_content, title, document_type
                )
                content_type = "application/pdf"
            elif format == "docx":
                file_bytes, filename = self.export_to_docx(html_content, title)
                content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            elif format == "txt":
                file_bytes, filename = self.export_to_txt(html_content, title)
                content_type = "text/plain"
            elif format == "html":
                file_bytes, filename = self.export_to_html(html_content, title)
                content_type = "text/html"
            else:
                raise ValueError(f"Unsupported format: {format}")

            # R2 key: exports/{user_id}/{document_id}/{format}/{filename}
            r2_key = f"exports/{user_id}/{document_id}/{format}/{filename}"

            # Upload to R2
            logger.info(f"📤 Uploading {format.upper()} to R2: {r2_key}")

            # Create temp file for upload
            with tempfile.NamedTemporaryFile(delete=False, suffix=f".{format}") as tmp:
                tmp.write(file_bytes)
                tmp_path = tmp.name

            try:
                # Upload to R2 with Content-Disposition header for download
                content_disposition = f'attachment; filename="{filename}"'
                upload_success = await self.r2_client.upload_file(
                    local_path=tmp_path,
                    remote_path=r2_key,
                    content_type=content_type,
                    content_disposition=content_disposition,
                )

                if not upload_success:
                    raise Exception("Failed to upload file to R2")

                # Generate presigned URL (1 hour expiry)
                download_url = await self.r2_client.generate_presigned_url(
                    remote_path=r2_key, expiration=3600, method="GET"
                )

                # Calculate expiry time
                expires_at = datetime.utcnow() + timedelta(seconds=3600)

                result = {
                    "download_url": download_url,
                    "filename": filename,
                    "file_size": len(file_bytes),
                    "format": format,
                    "expires_in": 3600,
                    "expires_at": expires_at.isoformat() + "Z",
                }

                logger.info(
                    f"✅ Export successful: {filename} ({len(file_bytes)} bytes)"
                )

                # Optional: Save export history to MongoDB
                # self.exports.insert_one({
                #     "export_id": f"exp_{uuid.uuid4().hex[:12]}",
                #     "user_id": user_id,
                #     "document_id": document_id,
                #     "format": format,
                #     "filename": filename,
                #     "file_size": len(file_bytes),
                #     "r2_key": r2_key,
                #     "created_at": datetime.utcnow()
                # })

                return result

            finally:
                # Clean up temp file
                if os.path.exists(tmp_path):
                    os.unlink(tmp_path)

        except Exception as e:
            logger.error(f"❌ Export and upload failed: {e}")
            raise Exception(f"Export failed: {str(e)}")

    def reconstruct_html_with_overlays(
        self, base_html: str, slide_elements: List[Dict]
    ) -> str:
        """
        Reconstruct HTML by injecting overlay elements into slides (ONLY for slide documents)

        Args:
            base_html: Background HTML (contentEditable content)
            slide_elements: Array of {slideIndex, elements[]}

        Returns:
            Full HTML with overlay elements injected
        """
        from bs4 import BeautifulSoup

        logger.info(
            f"🔧 Reconstructing HTML with overlay elements for SLIDE document..."
        )

        soup = BeautifulSoup(base_html, "html.parser")
        slides = soup.find_all(class_="slide")

        logger.info(f"📄 Found {len(slides)} slides in base HTML")

        # Create lookup map: slideIndex -> elements[]
        elements_map = {item["slideIndex"]: item["elements"] for item in slide_elements}

        # Inject overlays into each slide
        total_elements = 0
        for slide_idx, slide_tag in enumerate(slides):
            if slide_idx not in elements_map:
                continue

            elements = elements_map[slide_idx]
            logger.info(
                f"  📌 Slide {slide_idx + 1}: Injecting {len(elements)} overlay(s)"
            )

            for element in elements:
                # Convert JSON element to HTML string
                element_html_str = self._convert_element_to_html(element)

                if element_html_str:
                    # Parse and append to slide
                    element_soup = BeautifulSoup(element_html_str, "html.parser")
                    slide_tag.append(element_soup)
                    total_elements += 1

        logger.info(f"✅ Reconstruction complete: {total_elements} overlay(s) injected")

        return str(soup)

    def _convert_element_to_html(self, element: Dict) -> str:
        """
        Convert overlay element JSON to HTML string

        Supports: textbox, image, video, shape
        """
        element_type = element.get("type", "textbox")

        if element_type == "textbox":
            # Text align to justify content mapping
            text_align = element.get("textAlign", "left")
            justify_map = {
                "left": "flex-start",
                "center": "center",
                "right": "flex-end",
            }
            justify_content = justify_map.get(text_align, "flex-start")

            return f"""
        <div class="overlay-textbox" style="
            position: absolute;
            left: {element.get('x', 0)}px;
            top: {element.get('y', 0)}px;
            width: {element.get('width', 200)}px;
            height: {element.get('height', 100)}px;
            font-size: {element.get('fontSize', 16)}px;
            font-family: {element.get('fontFamily', 'Arial')};
            font-weight: {element.get('fontWeight', 'normal')};
            font-style: {element.get('fontStyle', 'normal')};
            text-decoration: {element.get('textDecoration', 'none')};
            color: {element.get('color', '#000000')};
            background-color: {element.get('backgroundColor', 'transparent')};
            border: {element.get('borderWidth', 0)}px {element.get('borderStyle', 'solid')} {element.get('borderColor', '#000')};
            border-radius: {element.get('borderRadius', 0)}px;
            padding: {element.get('padding', 8)}px;
            text-align: {text_align};
            transform: rotate({element.get('rotation', 0)}deg);
            transform-origin: center;
            z-index: {element.get('zIndex', 1)};
            opacity: {element.get('opacity', 1)};
            display: flex;
            align-items: center;
            justify-content: {justify_content};
            overflow: hidden;
            word-wrap: break-word;
            white-space: pre-wrap;
        ">{element.get('content', '')}</div>
        """

        elif element_type == "image":
            return f"""
        <img
            class="overlay-image"
            src="{element.get('src', '')}"
            alt="{element.get('alt', 'Image')}"
            style="
                position: absolute;
                left: {element.get('x', 0)}px;
                top: {element.get('y', 0)}px;
                width: {element.get('width', 200)}px;
                height: {element.get('height', 200)}px;
                object-fit: {element.get('objectFit', 'cover')};
                border-radius: {element.get('borderRadius', 0)}px;
                transform: rotate({element.get('rotation', 0)}deg);
                transform-origin: center;
                z-index: {element.get('zIndex', 1)};
                opacity: {element.get('opacity', 1)};
            "
        />
        """

        elif element_type == "video":
            # Enhanced video handling for PDF export
            video_id = element.get("videoId", "")
            width = element.get("width", 560)
            height = element.get("height", 315)
            x = element.get("x", 0)
            y = element.get("y", 0)
            z_index = element.get("zIndex", 1)

            if video_id:
                # YouTube video - show thumbnail with play button overlay
                # Use high quality thumbnail from YouTube
                thumbnail_url = (
                    f"https://img.youtube.com/vi/{video_id}/maxresdefault.jpg"
                )
                youtube_url = f"https://www.youtube.com/watch?v={video_id}"

                return f"""
        <div class="overlay-video-placeholder" style="
            position: absolute;
            left: {x}px;
            top: {y}px;
            width: {width}px;
            height: {height}px;
            background: #000;
            border-radius: 8px;
            overflow: hidden;
            z-index: {z_index};
        ">
            <img src="{thumbnail_url}"
                 alt="YouTube Video Thumbnail"
                 style="width: 100%; height: 100%; object-fit: cover;" />
            <div style="
                position: absolute;
                top: 50%;
                left: 50%;
                transform: translate(-50%, -50%);
                width: 80px;
                height: 80px;
                background: rgba(255, 0, 0, 0.9);
                border-radius: 50%;
                display: flex;
                align-items: center;
                justify-content: center;
            ">
                <div style="
                    width: 0;
                    height: 0;
                    border-left: 28px solid white;
                    border-top: 18px solid transparent;
                    border-bottom: 18px solid transparent;
                    margin-left: 6px;
                "></div>
            </div>
            <div style="
                position: absolute;
                bottom: 8px;
                left: 8px;
                right: 8px;
                background: rgba(0, 0, 0, 0.7);
                color: white;
                padding: 8px 12px;
                border-radius: 4px;
                font-size: 14px;
                font-family: Arial, sans-serif;
                text-align: center;
            ">
                🎥 YouTube: {youtube_url}
            </div>
        </div>
        """
            else:
                # Generic video placeholder
                return f"""
        <div class="overlay-video-placeholder" style="
            position: absolute;
            left: {x}px;
            top: {y}px;
            width: {width}px;
            height: {height}px;
            background-color: #1a1a1a;
            border-radius: 8px;
            display: flex;
            align-items: center;
            justify-content: center;
            flex-direction: column;
            color: white;
            z-index: {z_index};
        ">
            <div style="font-size: 64px; margin-bottom: 10px;">▶</div>
            <div style="font-size: 24px; font-weight: bold;">Video</div>
        </div>
        """

        elif element_type == "shape":
            # Enhanced shape rendering with full feature support
            shape_type = element.get("shapeType", "rectangle")

            # Determine border radius and special shapes
            if shape_type in ["circle"]:
                border_radius = "50%"
            elif shape_type == "ellipse":
                border_radius = "50%"
            elif shape_type == "sticky-note":
                border_radius = "8px"  # Rounded corners for sticky notes
            elif shape_type == "icon-circle-plus":
                border_radius = "50%"  # Circle shape
            else:
                border_radius = f"{element.get('borderRadius', 0)}px"

            # Fill (solid or gradient)
            fill_enabled = element.get("fillEnabled", True)
            fill_type = element.get("fillType", "solid")
            background_style = ""

            if fill_enabled:
                if fill_type == "gradient":
                    gradient = element.get("gradient", {})
                    gradient_type = gradient.get("type", "linear")
                    colors = gradient.get("colors", ["#667eea", "#764ba2"])

                    if gradient_type == "linear":
                        angle = gradient.get("angle", 135)
                        gradient_colors = ", ".join(colors)
                        background_style = f"background: linear-gradient({angle}deg, {gradient_colors});"
                    else:  # radial
                        gradient_colors = ", ".join(colors)
                        background_style = (
                            f"background: radial-gradient(circle, {gradient_colors});"
                        )
                else:  # solid
                    fill_color = element.get("fillColor", "#cccccc")
                    background_style = f"background-color: {fill_color};"
            else:
                background_style = "background: transparent;"

            # Border
            border_enabled = element.get("borderEnabled", True)
            border_style_css = ""

            if border_enabled:
                border_width = element.get("borderWidth", 2)
                border_color = element.get("borderColor", "#000000")
                border_style_type = element.get("borderStyle", "solid")
                border_style_css = (
                    f"border: {border_width}px {border_style_type} {border_color};"
                )
            else:
                border_style_css = "border: none;"

            # Shadow
            shadow_enabled = element.get("shadowEnabled", False)
            shadow_style = ""

            if shadow_enabled:
                shadow_x = element.get("shadowOffsetX", 0)
                shadow_y = element.get("shadowOffsetY", 4)
                shadow_blur = element.get("shadowBlur", 8)
                shadow_color = element.get("shadowColor", "rgba(0, 0, 0, 0.3)")
                shadow_style = f"box-shadow: {shadow_x}px {shadow_y}px {shadow_blur}px {shadow_color};"

            # Text overlay
            text_html = ""
            text_enabled = element.get("textEnabled", False)

            if text_enabled and element.get("text"):
                text_content = element.get("text", "")
                text_color = element.get("textColor", "#000000")
                text_size = element.get("textSize", 16)
                text_align = element.get("textAlign", "center")
                text_v_align = element.get("textVerticalAlign", "middle")
                text_bold = element.get("textBold", False)
                text_italic = element.get("textItalic", False)

                align_items_map = {
                    "top": "flex-start",
                    "middle": "center",
                    "bottom": "flex-end",
                }

                text_html = f"""
            <div style="
                position: absolute;
                top: 0;
                left: 0;
                width: 100%;
                height: 100%;
                display: flex;
                align-items: {align_items_map.get(text_v_align, 'center')};
                justify-content: {text_align};
                padding: 10px;
                color: {text_color};
                font-size: {text_size}px;
                font-weight: {'bold' if text_bold else 'normal'};
                font-style: {'italic' if text_italic else 'normal'};
                text-align: {text_align};
                pointer-events: none;
            ">{text_content}</div>
            """

            # Special shape rendering (icons, 3D shapes, etc.)
            special_shape_html = ""

            # Zigzag line pattern
            if shape_type == "line-zigzag":
                special_shape_html = """
            <svg width="100%" height="100%" style="position: absolute; top: 0; left: 0;">
                <path d="M 0,50 L 25,25 L 50,50 L 75,25 L 100,50" stroke="currentColor" stroke-width="2" fill="none" />
            </svg>
            """

            # 3D Cube
            elif shape_type == "cube-3d":
                special_shape_html = """
            <svg width="100%" height="100%" viewBox="0 0 100 100" style="position: absolute; top: 0; left: 0;">
                <polygon points="50,10 90,30 90,70 50,90 10,70 10,30" fill="currentColor" opacity="0.8"/>
                <polygon points="50,10 90,30 50,50 10,30" fill="currentColor" opacity="0.9"/>
                <polygon points="50,50 90,70 90,30 50,10" fill="currentColor" opacity="0.7"/>
            </svg>
            """

            # 3D Cylinder
            elif shape_type == "cylinder-3d":
                special_shape_html = """
            <svg width="100%" height="100%" viewBox="0 0 100 100" style="position: absolute; top: 0; left: 0;">
                <ellipse cx="50" cy="20" rx="40" ry="15" fill="currentColor" opacity="0.9"/>
                <rect x="10" y="20" width="80" height="60" fill="currentColor" opacity="0.8"/>
                <ellipse cx="50" cy="80" rx="40" ry="15" fill="currentColor" opacity="0.9"/>
            </svg>
            """

            # Sticky Note (with folded corner effect)
            elif shape_type == "sticky-note":
                sticky_bg = element.get("fillColor", "#fef08a")
                special_shape_html = f"""
            <div style="
                position: absolute;
                top: 0;
                left: 0;
                width: 100%;
                height: 100%;
                background: {sticky_bg};
                clip-path: polygon(0 0, 85% 0, 100% 15%, 100% 100%, 0 100%);
            "></div>
            <div style="
                position: absolute;
                top: 0;
                right: 0;
                width: 15%;
                height: 15%;
                background: linear-gradient(135deg, transparent 50%, rgba(0,0,0,0.1) 50%);
            "></div>
            """

            # Puzzle Piece
            elif shape_type == "puzzle-piece":
                special_shape_html = """
            <svg width="100%" height="100%" viewBox="0 0 100 100" style="position: absolute; top: 0; left: 0;">
                <path d="M 20,20 L 45,20 Q 50,10 55,20 L 80,20 L 80,45 Q 90,50 80,55 L 80,80 L 55,80 Q 50,90 45,80 L 20,80 L 20,55 Q 10,50 20,45 Z"
                      fill="currentColor" stroke="rgba(0,0,0,0.2)" stroke-width="1"/>
            </svg>
            """

            # Circle Plus Icon
            elif shape_type == "icon-circle-plus":
                icon_color = element.get("fillColor", "#10b981")
                special_shape_html = f"""
            <svg width="100%" height="100%" viewBox="0 0 100 100" style="position: absolute; top: 0; left: 0;">
                <circle cx="50" cy="50" r="45" fill="{icon_color}" stroke="rgba(0,0,0,0.2)" stroke-width="2"/>
                <line x1="50" y1="25" x2="50" y2="75" stroke="white" stroke-width="6" stroke-linecap="round"/>
                <line x1="25" y1="50" x2="75" y2="50" stroke="white" stroke-width="6" stroke-linecap="round"/>
            </svg>
            """

            return f"""
        <div class="overlay-shape shape-{shape_type}" style="
            position: absolute;
            left: {element.get('x', 0)}px;
            top: {element.get('y', 0)}px;
            width: {element.get('width', 100)}px;
            height: {element.get('height', 100)}px;
            {background_style if not special_shape_html else ''}
            {border_style_css if not special_shape_html else ''}
            border-radius: {border_radius};
            transform: rotate({element.get('rotation', 0)}deg);
            transform-origin: center;
            z-index: {element.get('zIndex', 1)};
            opacity: {element.get('opacity', 1)};
            {shadow_style}
            overflow: hidden;
            color: {element.get('fillColor', '#667eea')};
        ">{special_shape_html}{text_html}</div>
        """

        return ""  # Unknown type
