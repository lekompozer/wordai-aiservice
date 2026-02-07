"""
Document generation service layer
"""

import os
import uuid
from datetime import datetime, timedelta
from typing import Optional, List, Dict, Any
import asyncio
from pathlib import Path
import jinja2
from docx import Document
from docx.shared import Inches
import io
import base64

from ..models.document_generation_models import (
    DocumentTemplate,
    DocumentRequest,
    DocumentResponse,
    QuoteRequest,
    ContractRequest,
    AppendixRequest,
    CompanyInfo,
    CustomerInfo,
    ProductInfo,
    PaymentTerms,
)
from ..config.database import get_database
from ..config.ai_config import get_ai_client
from ..services.multi_ai_client import multi_ai_client
from .multi_ai_client import multi_ai_client


def format_currency(amount: float) -> str:
    """Format amount as Vietnamese currency"""
    return f"{amount:,.0f}".replace(",", ".")


def number_to_words_vietnamese(number: float) -> str:
    """Convert number to Vietnamese words (improved version)"""

    def convert_group_of_three(num):
        """Convert a group of three digits to words"""
        ones = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]

        if num == 0:
            return ""

        result = ""
        hundreds = num // 100
        remainder = num % 100

        # Handle hundreds
        if hundreds > 0:
            result += ones[hundreds] + " trăm"

        # Handle tens and ones
        if remainder > 0:
            if remainder < 10:
                if hundreds > 0:
                    result += " lẻ " + ones[remainder]
                else:
                    result += ones[remainder]
            elif remainder < 20:
                if remainder == 10:
                    result += " mười"
                elif remainder == 15:
                    result += " mười lăm"
                else:
                    result += " mười " + ones[remainder % 10]
            else:
                tens_digit = remainder // 10
                ones_digit = remainder % 10

                if tens_digit == 1:
                    result += " mười"
                else:
                    result += " " + ones[tens_digit] + " mươi"

                if ones_digit > 0:
                    if ones_digit == 5 and tens_digit > 1:
                        result += " lăm"
                    elif ones_digit == 1 and tens_digit > 1:
                        result += " một"
                    else:
                        result += " " + ones[ones_digit]

        return result.strip()

    number = int(number)
    if number == 0:
        return "không"

    if number < 0:
        return "âm " + number_to_words_vietnamese(-number)

    # Split number into groups of three digits
    groups = []
    while number > 0:
        groups.append(number % 1000)
        number //= 1000

    # Unit names for groups
    units = ["", "nghìn", "triệu", "tỷ"]

    result_parts = []
    for i in range(len(groups) - 1, -1, -1):
        if groups[i] > 0:
            group_words = convert_group_of_three(groups[i])
            if i > 0:
                group_words += " " + units[i]
            result_parts.append(group_words)

    return " ".join(result_parts)


class DocumentGenerationService:
    """Service xử lý tạo tài liệu"""

    def __init__(self):
        self.db = get_database()
        self.ai_client = get_ai_client()
        self.template_dir = Path("templates/documents")
        self.output_dir = Path("generated_documents")
        # Ensure parent directories are created as well
        self.template_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # Initialize Jinja2 environment
        self.jinja_env = jinja2.Environment(
            loader=jinja2.FileSystemLoader(str(self.template_dir)),
            autoescape=jinja2.select_autoescape(["html", "xml"]),
        )

    async def create_quote(
        self, request: QuoteRequest, user_id: Optional[str] = None
    ) -> DocumentResponse:
        """Tạo báo giá"""
        try:
            # Get template
            template = await self.get_template(request.template_id, "quote")
            if not template:
                raise ValueError(f"Template {request.template_id} không tồn tại")

            # Create document request
            doc_request = DocumentRequest(
                user_id=user_id,
                type="quote",
                template_id=template.id,
                company_info=request.company_info,
                customer_info=request.customer_info,
                products=request.products,
                payment_terms=request.payment_terms,
                additional_terms=request.additional_terms,
            )

            # Add quote-specific data
            quote_data = {
                "validity_period": getattr(request, "validity_period", "30 ngày"),
                "quote_number": f"BG-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}",
                "quote_date": datetime.now().strftime("%d/%m/%Y"),
            }

            # Generate document
            return await self.generate_document(doc_request, quote_data)

        except Exception as e:
            return DocumentResponse(
                request_id="", status="failed", message=f"Lỗi tạo báo giá: {str(e)}"
            )

    async def create_contract(
        self, request: ContractRequest, user_id: Optional[str] = None
    ) -> DocumentResponse:
        """Tạo hợp đồng"""
        try:
            # Get template
            template = await self.get_template(request.template_id, "contract")
            if not template:
                raise ValueError(f"Template {request.template_id} không tồn tại")

            # Create document request
            doc_request = DocumentRequest(
                user_id=user_id,
                type="contract",
                template_id=template.id,
                company_info=request.company_info,
                customer_info=request.customer_info,
                products=request.products,
                payment_terms=request.payment_terms,
                additional_terms=request.additional_terms,
            )

            # Add contract-specific data
            contract_data = {
                "contract_number": f"HD-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}",
                "contract_date": datetime.now().strftime("%d/%m/%Y"),
                "contract_duration": getattr(request, "contract_duration", "12 tháng"),
                "effective_date": (
                    getattr(request, "effective_date", datetime.now()).strftime(
                        "%d/%m/%Y"
                    )
                    if hasattr(request, "effective_date") and request.effective_date
                    else datetime.now().strftime("%d/%m/%Y")
                ),
            }

            # Generate document
            return await self.generate_document(doc_request, contract_data)

        except Exception as e:
            return DocumentResponse(
                request_id="", status="failed", message=f"Lỗi tạo hợp đồng: {str(e)}"
            )

    async def create_appendix(
        self, request: AppendixRequest, user_id: Optional[str] = None
    ) -> DocumentResponse:
        """Tạo phụ lục"""
        try:
            # Get template
            template = await self.get_template(request.template_id, "appendix")
            if not template:
                raise ValueError(f"Template {request.template_id} không tồn tại")

            # Create document request
            doc_request = DocumentRequest(
                user_id=user_id,
                type="appendix",
                template_id=template.id,
                company_info=request.company_info,
                customer_info=request.customer_info,
                products=request.products,
                additional_terms=request.additional_terms,
            )

            # Add appendix-specific data
            appendix_data = {
                "appendix_number": f"PL-{datetime.now().strftime('%Y%m%d')}-{str(uuid.uuid4())[:8].upper()}",
                "appendix_date": datetime.now().strftime("%d/%m/%Y"),
                "parent_contract_id": getattr(request, "parent_contract_id", ""),
                "changes_description": request.changes_description,
            }

            # Generate document
            return await self.generate_document(doc_request, appendix_data)

        except Exception as e:
            return DocumentResponse(
                request_id="", status="failed", message=f"Lỗi tạo phụ lục: {str(e)}"
            )

    async def generate_document(
        self, doc_request: DocumentRequest, extra_data: Dict[str, Any] = None
    ) -> DocumentResponse:
        """Generate document using AI and templates"""
        start_time = datetime.utcnow()

        try:
            # Save request to database
            result = self.db.document_requests.insert_one(
                doc_request.dict(by_alias=True)
            )
            doc_request.id = result.inserted_id
            request_id = str(doc_request.id)

            # Get template
            template = await self.get_template(
                str(doc_request.template_id), doc_request.type
            )

            # Prepare template variables
            template_vars = await self.prepare_template_variables(
                doc_request, extra_data or {}
            )

            # Generate content using AI
            ai_content = await self.generate_ai_content(
                doc_request, template, template_vars
            )

            # Generate Word document
            file_path = await self.generate_word_document(
                doc_request, ai_content, template_vars
            )

            # Calculate processing time
            processing_time = (datetime.utcnow() - start_time).total_seconds()

            # Update request with results
            expires_at = datetime.utcnow() + timedelta(
                hours=24
            )  # File expires in 24 hours
            file_url = f"/api/documents/download/{request_id}"

            self.db.document_requests.update_one(
                {"_id": doc_request.id},
                {
                    "$set": {
                        "generated_content": ai_content,
                        "ai_processing_time": processing_time,
                        "file_path": file_path,
                        "file_url": file_url,
                        "status": "completed",
                        "expires_at": expires_at,
                        "updated_at": datetime.utcnow(),
                    }
                },
            )

            return DocumentResponse(
                request_id=request_id,
                status="completed",
                message="Tài liệu đã được tạo thành công",
                file_url=file_url,
                expires_at=expires_at,
                processing_time=processing_time,
            )

        except Exception as e:
            # Update request with error
            if hasattr(doc_request, "id") and doc_request.id:
                self.db.document_requests.update_one(
                    {"_id": doc_request.id},
                    {
                        "$set": {
                            "status": "failed",
                            "error_message": str(e),
                            "updated_at": datetime.utcnow(),
                        }
                    },
                )

            return DocumentResponse(
                request_id=(
                    str(doc_request.id)
                    if hasattr(doc_request, "id") and doc_request.id
                    else ""
                ),
                status="failed",
                message=f"Lỗi tạo tài liệu: {str(e)}",
            )

    async def prepare_template_variables(
        self, doc_request: DocumentRequest, extra_data: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Prepare all variables for template rendering"""

        # Format products table
        products_data = []
        for i, product in enumerate(doc_request.products, 1):
            products_data.append(
                {
                    "stt": i,
                    "name": product.name,
                    "description": product.description,
                    "quantity": product.quantity,
                    "unit": product.unit,
                    "unit_price": format_currency(product.unit_price),
                    "total_price": format_currency(product.total_price),
                    "specifications": product.specifications or {},
                    "warranty_period": product.warranty_period
                    or "Theo quy định nhà sản xuất",
                    "delivery_time": product.delivery_time or "Theo thỏa thuận",
                }
            )

        template_vars = {
            # Company info
            "company": doc_request.company_info.dict(),
            # Customer info
            "customer": doc_request.customer_info.dict(),
            # Products
            "products": products_data,
            "total_amount": format_currency(doc_request.total_amount),
            "total_quantity": doc_request.total_quantity,
            # Payment terms
            "payment_terms": (
                doc_request.payment_terms.dict() if doc_request.payment_terms else {}
            ),
            # Additional terms
            "additional_terms": doc_request.additional_terms or "",
            # VAT information
            "vat_rate": extra_data.get("vat_rate", 10.0),
            # Notes
            "notes": extra_data.get("notes", ""),
            # Dates
            "current_date": datetime.now().strftime("%d/%m/%Y"),
            "current_time": datetime.now().strftime("%H:%M:%S"),
            # Document type
            "document_type": doc_request.type,
        }

        # Add extra data
        template_vars.update(extra_data)

        return template_vars

    async def generate_ai_content(
        self,
        doc_request: DocumentRequest,
        template: DocumentTemplate,
        template_vars: Dict[str, Any],
    ) -> str:
        """Generate content using AI"""

        # Create AI prompt
        prompt = f"""
        Bạn là chuyên gia pháp lý và kinh doanh. Hãy tạo nội dung {doc_request.type} chuyên nghiệp bằng tiếng Việt.

        Thông tin công ty: {doc_request.company_info.dict()}
        Thông tin khách hàng: {doc_request.customer_info.dict()}
        Danh sách sản phẩm/dịch vụ: {[p.dict() for p in doc_request.products]}
        Điều khoản thanh toán: {doc_request.payment_terms.dict() if doc_request.payment_terms else 'Chưa có'}
        Điều khoản bổ sung: {doc_request.additional_terms or 'Không có'}

        Template cơ bản: {template.template_content}

        Yêu cầu:
        1. Sử dụng ngôn ngữ chuyên nghiệp, trang trọng
        2. Đảm bảo đầy đủ thông tin pháp lý
        3. Có cấu trúc rõ ràng, dễ đọc
        4. Phù hợp với luật pháp Việt Nam
        5. Trả về nội dung định dạng HTML đẹp

        Hãy tạo nội dung hoàn chỉnh:
        """

        try:
            # Call AI service (assuming we have an AI client configured)
            response = await self.ai_client.generate_text(
                prompt=prompt, max_tokens=4000, temperature=0.3
            )

            return response.get("content", template.template_content)

        except Exception as e:
            print(f"AI generation error: {e}")
            # Fallback to template rendering
            template_obj = self.jinja_env.from_string(template.template_content)
            return template_obj.render(**template_vars)

    async def generate_word_document(
        self, doc_request: DocumentRequest, content: str, template_vars: Dict[str, Any]
    ) -> str:
        """Generate Word document with detailed quote format"""

        # Create new Word document
        doc = Document()

        # Set up document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.8)
            section.bottom_margin = Inches(0.8)
            section.left_margin = Inches(0.8)
            section.right_margin = Inches(0.8)

        # Company logo and header information
        if doc_request.company_info.logo:
            try:
                # If logo is provided, add it (assuming it's a file path or base64)
                # For now, we'll add a placeholder for logo implementation
                logo_para = doc.add_paragraph()
                logo_para.alignment = 0  # Left alignment
                # TODO: Implement actual logo insertion
                doc.add_paragraph(f"[LOGO: {doc_request.company_info.name}]")
            except:
                # If logo fails, continue without it
                pass

        # Company header information
        header_para = doc.add_paragraph()
        header_para.alignment = 0  # Left alignment

        company_header = f"""
{doc_request.company_info.name}
{doc_request.company_info.address}
"""
        if doc_request.company_info.city:
            company_header += f"{doc_request.company_info.city}\n"

        company_header += f"""
Điện thoại: {doc_request.company_info.phone}
"""

        if doc_request.company_info.fax:
            company_header += f"Fax: {doc_request.company_info.fax}\n"

        company_header += f"Email: {doc_request.company_info.email}\n"

        if doc_request.company_info.website:
            company_header += f"Website: {doc_request.company_info.website}\n"

        if doc_request.company_info.social_link:
            company_header += f"Social: {doc_request.company_info.social_link}\n"

        header_para.text = company_header.strip()

        # Add some space
        doc.add_paragraph()

        # Quote title
        title_para = doc.add_heading("BẢNG BÁO GIÁ", 0)
        title_para.alignment = 1  # Center alignment

        # Add space
        doc.add_paragraph()

        # "Kính gửi" section
        greeting_para = doc.add_paragraph()
        greeting_text = f"Kính gửi: {doc_request.customer_info.name}\n"
        greeting_text += f"{doc_request.company_info.name} xin gửi đến Quý khách hàng báo giá sản phẩm (dịch vụ) như sau:"
        greeting_para.text = greeting_text

        # Add space
        doc.add_paragraph()

        # Products table
        if doc_request.products:
            # Create table with 7 columns as specified
            table = doc.add_table(rows=1, cols=7)
            table.style = "Table Grid"

            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = "STT"
            hdr_cells[1].text = "Tên Hàng"
            hdr_cells[2].text = "Quy cách và đặc tính kỹ thuật"
            hdr_cells[3].text = "Đơn vị tính"
            hdr_cells[4].text = "Số lượng"
            hdr_cells[5].text = "Đơn giá (VND)"
            hdr_cells[6].text = "Thành tiền (VND)"

            # Make header bold
            for cell in hdr_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # Product rows
            subtotal = 0
            for i, product in enumerate(doc_request.products, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[1].text = product.name

                # Technical specifications
                specs_text = product.description
                if product.specifications:
                    specs_list = []
                    for key, value in product.specifications.items():
                        specs_list.append(f"{key}: {value}")
                    if specs_list:
                        specs_text += f"\n{', '.join(specs_list)}"

                row_cells[2].text = specs_text
                row_cells[3].text = product.unit
                row_cells[4].text = str(product.quantity)
                row_cells[5].text = f"{product.unit_price:,.0f}".replace(",", ".")
                row_cells[6].text = f"{product.total_price:,.0f}".replace(",", ".")
                subtotal += product.total_price

            # Subtotal row
            subtotal_row = table.add_row().cells
            subtotal_row[0].text = ""
            subtotal_row[1].text = ""
            subtotal_row[2].text = ""
            subtotal_row[3].text = ""
            subtotal_row[4].text = ""
            subtotal_row[5].text = "Cộng:"
            subtotal_row[6].text = f"{subtotal:,.0f}".replace(",", ".")

            # Make subtotal bold
            for cell in subtotal_row[5:]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

            # VAT calculation
            vat_rate = template_vars.get("vat_rate", 10.0)  # Default 10%
            vat_amount = subtotal * (vat_rate / 100)

            # VAT row
            vat_row = table.add_row().cells
            vat_row[0].text = ""
            vat_row[1].text = ""
            vat_row[2].text = ""
            vat_row[3].text = ""
            vat_row[4].text = ""
            vat_row[5].text = f"Thuế VAT ({vat_rate}%):"
            vat_row[6].text = f"{vat_amount:,.0f}".replace(",", ".")

            # Total row
            total_amount = subtotal + vat_amount
            total_row = table.add_row().cells
            total_row[0].text = ""
            total_row[1].text = ""
            total_row[2].text = ""
            total_row[3].text = ""
            total_row[4].text = ""
            total_row[5].text = "Tổng giá:"
            total_row[6].text = f"{total_amount:,.0f}".replace(",", ".")

            # Make total bold
            for cell in total_row[5:]:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True

        # Add space
        doc.add_paragraph()

        # Amount in words
        amount_words_para = doc.add_paragraph()
        total_amount_value = subtotal + (
            subtotal * (template_vars.get("vat_rate", 10.0) / 100)
        )
        amount_words_para.text = (
            f"Số tiền bằng chữ: {self.number_to_words(total_amount_value)} đồng"
        )

        # Add space
        doc.add_paragraph()

        # Notes section
        notes = template_vars.get("notes") or doc_request.additional_terms
        if notes:
            notes_para = doc.add_paragraph()
            notes_para.text = f"Ghi chú:\n{notes}"

        # Add space
        doc.add_paragraph()

        # Signature section (right aligned)
        signature_para = doc.add_paragraph()
        signature_para.alignment = 2  # Right alignment

        current_date = datetime.now()
        city = doc_request.company_info.city or "TP.HCM"
        signature_text = f"""{city}, Ngày {current_date.day} tháng {current_date.month} năm {current_date.year}

Đại diện {doc_request.company_info.name}


{doc_request.company_info.representative}
{doc_request.company_info.position or 'Giám đốc'}"""

        signature_para.text = signature_text

        # Save document
        filename = f"quote_{datetime.now().strftime('%Y%m%d_%H%M%S')}_{str(uuid.uuid4())[:8]}.docx"
        file_path = self.output_dir / filename

        doc.save(str(file_path))

        return str(file_path)

    def number_to_words(self, number: float) -> str:
        """Convert number to Vietnamese words"""

        def convert_group_of_three(num):
            """Convert a group of three digits to words"""
            ones = ["", "một", "hai", "ba", "bốn", "năm", "sáu", "bảy", "tám", "chín"]

            if num == 0:
                return ""

            result = ""
            hundreds = num // 100
            remainder = num % 100

            # Handle hundreds
            if hundreds > 0:
                result += ones[hundreds] + " trăm"

            # Handle tens and ones
            if remainder > 0:
                if remainder < 10:
                    if hundreds > 0:
                        result += " lẻ " + ones[remainder]
                    else:
                        result += ones[remainder]
                elif remainder < 20:
                    if remainder == 10:
                        result += " mười"
                    elif remainder == 15:
                        result += " mười lăm"
                    else:
                        result += " mười " + ones[remainder % 10]
                else:
                    tens_digit = remainder // 10
                    ones_digit = remainder % 10

                    if tens_digit == 1:
                        result += " mười"
                    else:
                        result += " " + ones[tens_digit] + " mươi"

                    if ones_digit > 0:
                        if ones_digit == 5 and tens_digit > 1:
                            result += " lăm"
                        elif ones_digit == 1 and tens_digit > 1:
                            result += " một"
                        else:
                            result += " " + ones[ones_digit]

            return result.strip()

        number = int(number)
        if number == 0:
            return "không"

        if number < 0:
            return "âm " + self.number_to_words(-number)

        # Split number into groups of three digits
        groups = []
        while number > 0:
            groups.append(number % 1000)
            number //= 1000

        # Unit names for groups
        units = ["", "nghìn", "triệu", "tỷ"]

        result_parts = []
        for i in range(len(groups) - 1, -1, -1):
            if groups[i] > 0:
                group_words = convert_group_of_three(groups[i])
                if i > 0:
                    group_words += " " + units[i]
                result_parts.append(group_words)

        return " ".join(result_parts)

    async def get_template(
        self, template_id: str, doc_type: str
    ) -> Optional[DocumentTemplate]:
        """Get template by ID and type"""
        try:
            from bson import ObjectId

            template_data = self.db.document_templates.find_one(
                {"_id": ObjectId(template_id), "type": doc_type, "is_active": True}
            )

            if template_data:
                return DocumentTemplate(**template_data)
            return None

        except Exception as e:
            print(f"Error getting template: {e}")
            return None

    async def list_templates(
        self, doc_type: Optional[str] = None
    ) -> List[DocumentTemplate]:
        """List available templates"""
        try:
            query = {"is_active": True}
            if doc_type:
                query["type"] = doc_type

            cursor = self.db.document_templates.find(query).sort("name", 1)
            templates = []

            for template_data in cursor:  # Use regular for loop, not async for
                templates.append(DocumentTemplate(**template_data))

            return templates

        except Exception as e:
            print(f"Error listing templates: {e}")
            return []

    async def get_document_status(self, request_id: str) -> Optional[DocumentRequest]:
        """Get document generation status"""
        try:
            from bson import ObjectId

            doc_data = self.db.document_requests.find_one({"_id": ObjectId(request_id)})

            if doc_data:
                return DocumentRequest(**doc_data)
            return None

        except Exception as e:
            print(f"Error getting document status: {e}")
            return None

    async def get_file_path(self, request_id: str) -> Optional[str]:
        """Get file path for download"""
        doc_request = await self.get_document_status(request_id)

        if doc_request and doc_request.status == "completed" and doc_request.file_path:
            # Check if file still exists and not expired
            if doc_request.expires_at and datetime.utcnow() > doc_request.expires_at:
                return None

            file_path = Path(doc_request.file_path)
            if file_path.exists():
                return str(file_path)

        return None

    async def generate_ai_template(self, request: Dict[str, Any]) -> Dict[str, Any]:
        """Generate template structure using AI"""
        try:
            # Build AI prompt based on request data
            prompt = self._build_ai_template_prompt(request)

            # Get selected AI model
            ai_model = request.get("ai_model", "deepseek")

            # Call AI to generate template structure
            ai_response = await multi_ai_client.generate_text(
                prompt=prompt, model=ai_model, max_tokens=2000, temperature=0.3
            )

            # Parse AI response to template structure
            template_structure = self._parse_ai_template_response(ai_response, request)

            return {
                "template_structure": template_structure,
                "placeholders": self._generate_placeholders(request),
                "sections": self._generate_sections(template_structure),
                "styling": self._get_default_styling(),
                "metadata": {
                    "generated_by": "ai",
                    "ai_model": ai_model,
                    "created_at": datetime.now().isoformat(),
                    "language": request.get("language", "vi"),
                    "style": request.get("style", "professional"),
                },
            }

        except Exception as e:
            print(f"Error generating AI template: {e}")
            # Fallback to default template structure
            fallback_structure = self._get_fallback_template_structure(request)
            return {
                "template_structure": fallback_structure,
                "placeholders": self._generate_placeholders(request),
                "sections": self._generate_sections(fallback_structure),
                "styling": self._get_default_styling(),
                "metadata": {
                    "generated_by": "fallback",
                    "created_at": datetime.now().isoformat(),
                    "language": request.get("language", "vi"),
                    "style": request.get("style", "professional"),
                },
            }

    def _build_ai_template_prompt(self, request: Dict[str, Any]) -> str:
        """Build prompt for AI template generation"""
        doc_type = request.get("document_type", "quote")
        template_option = request.get("template_option", "use_existing")
        language = request.get("language", "vi")

        prompt = f"""
Tạo cấu trúc template cho tài liệu {doc_type} bằng tiếng {language}.

Thông tin tài liệu:
- Loại: {doc_type}
- Ngôn ngữ: {language}
- Phong cách: {request.get('style', 'professional')}

"""

        # Add company info if available
        if request.get("company_info"):
            company = request["company_info"]
            prompt += f"""
Thông tin công ty:
- Tên: {company.get('name', '[Tên công ty]')}
- Địa chỉ: {company.get('address', '[Địa chỉ công ty]')}
- Phone: {company.get('phone', '[Số điện thoại]')}
- Email: {company.get('email', '[Email]')}
"""

        # Add custom requirements
        if request.get("custom_requirements"):
            prompt += f"\nYêu cầu đặc biệt: {request['custom_requirements']}\n"

        prompt += """
Hãy tạo cấu trúc JSON với các phần:
1. Header (tiêu đề, logo, thông tin công ty)
2. Document info (số tài liệu, ngày tháng)
3. Customer section (thông tin khách hàng)
4. Content (sản phẩm/dịch vụ, bảng giá)
5. Terms (điều khoản thanh toán, ghi chú)
6. Footer (chữ ký, thông tin liên hệ)

Trả về JSON structure cho Word document generation.
"""

        return prompt

    def _parse_ai_template_response(
        self, ai_response: str, request: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Parse AI response to template structure"""
        try:
            # Try to extract JSON from AI response
            import json
            import re

            # Find JSON in the response
            json_match = re.search(r"\{.*\}", ai_response, re.DOTALL)
            if json_match:
                return json.loads(json_match.group())
            else:
                # Fallback if no JSON found
                return self._get_fallback_template_structure(request)

        except Exception as e:
            print(f"Error parsing AI response: {e}")
            return self._get_fallback_template_structure(request)

    def _generate_placeholders(self, request: Dict[str, Any]) -> Dict[str, str]:
        """Generate placeholders with default values"""
        placeholders = {}

        # Company placeholders
        company = request.get("company_info", {})
        placeholders.update(
            {
                "company_name": company.get("name", "[Tên công ty]"),
                "company_address": company.get("address", "[Địa chỉ công ty]"),
                "company_phone": company.get("phone", "[Số điện thoại]"),
                "company_email": company.get("email", "[Email công ty]"),
                "company_tax_code": company.get("tax_code", "[Mã số thuế]"),
                "company_representative": company.get(
                    "representative", "[Người đại diện]"
                ),
            }
        )

        # Customer placeholders
        customer = request.get("customer_info", {})
        placeholders.update(
            {
                "customer_name": customer.get("name", "[Tên khách hàng]"),
                "customer_address": customer.get("address", "[Địa chỉ khách hàng]"),
                "customer_phone": customer.get("phone", "[Số điện thoại khách hàng]"),
                "customer_email": customer.get("email", "[Email khách hàng]"),
                "contact_person": customer.get("contact_person", "[Người liên hệ]"),
            }
        )

        # Document placeholders
        placeholders.update(
            {
                "document_number": "[Số tài liệu]",
                "document_date": datetime.now().strftime("%d/%m/%Y"),
                "total_amount": "[Tổng tiền]",
                "currency": "VNĐ",
            }
        )

        return placeholders

    def _generate_sections(
        self, template_structure: Dict[str, Any]
    ) -> List[Dict[str, Any]]:
        """Generate document sections"""
        return [
            {
                "name": "header",
                "type": "header",
                "content": {
                    "title": template_structure.get("title", "BÁO GIÁ"),
                    "logo_position": "left",
                    "company_info_position": "right",
                },
            },
            {
                "name": "document_info",
                "type": "info",
                "content": {"document_number": True, "date": True, "validity": True},
            },
            {
                "name": "parties",
                "type": "parties",
                "content": {"company_section": True, "customer_section": True},
            },
            {
                "name": "products",
                "type": "table",
                "content": {
                    "headers": [
                        "STT",
                        "Sản phẩm/Dịch vụ",
                        "Số lượng",
                        "Đơn vị",
                        "Đơn giá",
                        "Thành tiền",
                    ],
                    "show_total": True,
                },
            },
            {
                "name": "terms",
                "type": "terms",
                "content": {"payment_terms": True, "additional_notes": True},
            },
            {
                "name": "signatures",
                "type": "signatures",
                "content": {"company_signature": True, "customer_signature": False},
            },
        ]

    def _get_default_styling(self) -> Dict[str, Any]:
        """Get default document styling"""
        return {
            "font_family": "Times New Roman",
            "font_size": 12,
            "header_font_size": 16,
            "title_font_size": 18,
            "line_spacing": 1.15,
            "margins": {"top": 2.5, "bottom": 2.5, "left": 2.5, "right": 2.5},
            "colors": {
                "primary": "#000000",
                "secondary": "#666666",
                "accent": "#0066CC",
            },
        }

    def _get_fallback_template_structure(
        self, request: Dict[str, Any]
    ) -> Dict[str, Any]:
        """Get fallback template structure when AI fails"""
        doc_type = request.get("document_type", "quote")

        if doc_type == "quote":
            return {
                "title": "BÁO GIÁ",
                "sections": [
                    "header",
                    "document_info",
                    "parties",
                    "products",
                    "terms",
                    "signatures",
                ],
                "layout": "standard_quote",
            }
        elif doc_type == "contract":
            return {
                "title": "HỢP ĐỒNG",
                "sections": [
                    "header",
                    "document_info",
                    "parties",
                    "contract_terms",
                    "products",
                    "terms",
                    "signatures",
                ],
                "layout": "standard_contract",
            }
        else:
            return {
                "title": "PHỤ LỤC",
                "sections": [
                    "header",
                    "document_info",
                    "changes",
                    "products",
                    "terms",
                    "signatures",
                ],
                "layout": "standard_appendix",
            }
