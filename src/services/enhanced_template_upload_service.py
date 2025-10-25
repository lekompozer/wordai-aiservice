"""
Enhanced Template Upload Service
Xử lý upload và phân tích template DOCX với AI theo workflow đã thiết kế
"""

import os
import uuid
import base64
import json
from typing import Dict, Any, List, Optional
from datetime import datetime
from fastapi import UploadFile, HTTPException
from docx import Document
from pydantic import BaseModel, Field
from io import BytesIO

from src.utils.logger import setup_logger
from src.services.ai_service import get_ai_service
from src.config.database import get_async_database

logger = setup_logger(__name__)

# Create logs directory if not exists
LOGS_DIR = "/Users/user/Code/ai-chatbot-rag/logs"
os.makedirs(LOGS_DIR, exist_ok=True)

# Create temp files directory with full permissions
TEMP_FILES_DIR = "/Users/user/Code/ai-chatbot-rag/temp_files"
os.makedirs(TEMP_FILES_DIR, exist_ok=True)
os.chmod(TEMP_FILES_DIR, 0o755)


def safe_file_write(file_path: str, content, mode: str = "w", **kwargs) -> bool:
    """
    Safely write to file with macOS permission handling
    Returns True if successful, False otherwise
    """
    try:
        with open(file_path, mode, **kwargs) as f:
            if mode.startswith("w") and isinstance(content, dict):
                import json

                json.dump(content, f, indent=2, ensure_ascii=False)
            elif hasattr(content, "write"):
                f.write(content)
            elif isinstance(content, bytes):
                f.write(content)
            else:
                f.write(content)
        logger.debug(f"✅ Successfully wrote file: {file_path}")
        return True
    except PermissionError as e:
        logger.warning(f"⚠️ Permission denied writing to {file_path}: {e}")
        logger.info("💡 File write skipped due to macOS permissions")
        return False
    except Exception as e:
        logger.warning(f"⚠️ Failed to write file {file_path}: {e}")
        return False


class PlaceholderInfo(BaseModel):
    """Thông tin chi tiết về placeholder"""

    type: str  # text, number, date, currency, boolean, calculated, array
    description: str
    current_value: Any = ""  # Can be string, number, array, etc.
    default_value: Any = ""
    validation_rules: List[str] = []
    section: str
    auto_populate: bool = False
    calculation_formula: Optional[str] = None
    position: Optional[Dict[str, Any]] = None
    formatting: Optional[Dict[str, Any]] = None


class TemplateSection(BaseModel):
    """Thông tin về section trong template"""

    name: str
    description: str
    placeholders: List[str]
    order: int
    is_repeatable: bool = False
    required: bool = True
    table_structure: Optional[Dict[str, Any]] = None


class DocumentStructure(BaseModel):
    """Cấu trúc document"""

    total_pages: int = 1
    has_tables: bool = False
    table_locations: List[str] = []
    header_content: str = ""
    footer_content: str = ""
    visual_elements: List[str] = []


class TemplateAnalysisResult(BaseModel):
    """Kết quả phân tích template từ AI"""

    template_id: str
    placeholders: Dict[str, PlaceholderInfo]
    sections: List[TemplateSection]
    business_logic: Dict[str, Any]
    document_structure: DocumentStructure
    ai_analysis_score: float
    processing_time: float


class EnhancedTemplateUploadService:
    """Enhanced service với PDF processing cho Gemini"""

    def __init__(self):
        self.ai_service = get_ai_service()

    async def _get_database(self):
        """Get database instance"""
        return await get_async_database()

    async def validate_docx_file(self, file: UploadFile) -> Dict[str, Any]:
        """Validate uploaded DOCX file"""
        errors = []
        warnings = []

        # Check file type
        if not file.filename.endswith(".docx"):
            errors.append("Only DOCX files are supported")

        # Check file size (max 100MB)
        if file.size and file.size > 100 * 1024 * 1024:
            errors.append("File size must be less than 100MB")

        # Check if file is valid DOCX
        try:
            content = await file.read()
            await file.seek(0)  # Reset file pointer

            # Try to open with python-docx
            doc = Document(BytesIO(content))

            # Basic content validation
            if len(doc.paragraphs) < 3:
                warnings.append("Template seems too short, might be missing content")

        except Exception as e:
            errors.append(f"Invalid DOCX file format: {str(e)}")

        return {"is_valid": len(errors) == 0, "errors": errors, "warnings": warnings}

    async def convert_docx_to_pdf(
        self, docx_content: bytes, template_id: str
    ) -> Optional[str]:
        """Convert DOCX to PDF for Gemini Vision API"""
        try:
            import tempfile
            import os

            # Try LibreOffice first (headless, no TCC popup), then fallback to text-based PDF
            for method_name, method in [
                (
                    "libreoffice",
                    self._convert_with_libreoffice,
                ),  # Priority #1 - No TCC popup, production ready
                (
                    "text-to-pdf",
                    self._create_pdf_from_text,
                ),  # Priority #2 - Stable fallback, works everywhere
            ]:
                try:
                    logger.info(f"🔄 Attempting PDF conversion with {method_name}...")
                    pdf_content = await method(docx_content, template_id)
                    if pdf_content:
                        logger.info(f"✅ PDF conversion successful with {method_name}")
                        return pdf_content
                except PermissionError as e:
                    logger.warning(
                        f"⚠️ {method_name} permission denied (macOS dev environment): {e}"
                    )
                    logger.info(
                        "💡 Skipping PDF conversion due to macOS permissions - this is normal in development"
                    )
                    continue
                except Exception as e:
                    logger.warning(f"⚠️ {method_name} conversion failed: {e}")
                    continue

            logger.warning(
                "⚠️ All PDF conversion methods failed - continuing with DOCX-only analysis"
            )
            logger.info("💡 This is acceptable for development environment")
            return None

        except Exception as e:
            logger.error(f"❌ PDF conversion error: {e}")
            logger.info("💡 Continuing without PDF conversion")
            return None

    async def _convert_with_libreoffice(
        self, docx_content: bytes, template_id: str
    ) -> Optional[bytes]:
        """Convert using LibreOffice headless mode (no TCC popup)"""
        try:
            import subprocess
            import os

            # Create temporary files in project temp_files directory
            temp_id = str(uuid.uuid4())[:8]
            docx_temp_path = f"{TEMP_FILES_DIR}/libre_{template_id}_{temp_id}.docx"
            pdf_temp_path = f"{TEMP_FILES_DIR}/libre_{template_id}_{temp_id}.pdf"

            # Write DOCX content to temp file
            with open(docx_temp_path, "wb") as f:
                f.write(docx_content)

            # Detect LibreOffice path based on platform
            import platform
            import shutil

            # Try to find LibreOffice executable
            libreoffice_paths = [
                "/Applications/LibreOffice.app/Contents/MacOS/soffice",  # macOS
                "/usr/bin/libreoffice",  # Linux/Ubuntu (production)
                "/usr/local/bin/libreoffice",  # Alternative Linux path
                "libreoffice",  # System PATH
                "soffice",  # Alternative command name
            ]

            libreoffice_exe = None
            for path in libreoffice_paths:
                if path.startswith("/") and os.path.exists(path):
                    libreoffice_exe = path
                    break
                elif not path.startswith("/") and shutil.which(path):
                    libreoffice_exe = path
                    break

            if not libreoffice_exe:
                raise Exception(
                    "LibreOffice not found. Please install LibreOffice on the system."
                )

            # Use LibreOffice headless conversion (no GUI, no TCC popup)
            libreoffice_cmd = [
                libreoffice_exe,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                TEMP_FILES_DIR,
                docx_temp_path,
            ]

            logger.info(f"🔄 Running LibreOffice headless: {' '.join(libreoffice_cmd)}")

            # Run LibreOffice conversion
            result = subprocess.run(
                libreoffice_cmd,
                capture_output=True,
                text=True,
                timeout=30,  # 30 second timeout
            )

            if result.returncode == 0:
                # Check if PDF was created
                if os.path.exists(pdf_temp_path) and os.path.getsize(pdf_temp_path) > 0:
                    # Read PDF content
                    with open(pdf_temp_path, "rb") as pdf_file:
                        pdf_content = pdf_file.read()

                    # Clean up temp files
                    os.unlink(docx_temp_path)
                    os.unlink(pdf_temp_path)

                    logger.info(
                        f"✅ LibreOffice conversion successful: {len(pdf_content)} bytes"
                    )
                    return pdf_content
                else:
                    logger.error("❌ LibreOffice PDF output file not found or empty")
            else:
                logger.error(f"❌ LibreOffice conversion failed: {result.stderr}")

            # Clean up temp files on error
            if os.path.exists(docx_temp_path):
                os.unlink(docx_temp_path)
            if os.path.exists(pdf_temp_path):
                os.unlink(pdf_temp_path)

            return None

        except subprocess.TimeoutExpired:
            logger.error("❌ LibreOffice conversion timeout")
            return None
        except Exception as e:
            logger.error(f"❌ LibreOffice conversion error: {e}")
            return None

    async def _create_pdf_from_text(
        self, docx_content: bytes, template_id: str
    ) -> Optional[bytes]:
        """Fallback: Create PDF from extracted text using reportlab"""
        try:
            from reportlab.lib.pagesizes import letter, A4
            from reportlab.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from docx import Document
            import tempfile
            import os
            from io import BytesIO

            # Extract text from DOCX
            doc = Document(BytesIO(docx_content))
            text_content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Create PDF
            buffer = BytesIO()
            doc_pdf = SimpleDocTemplate(buffer, pagesize=A4)
            styles = getSampleStyleSheet()
            story = []

            # Add title
            title_style = ParagraphStyle(
                "CustomTitle",
                parent=styles["Heading1"],
                fontSize=16,
                spaceAfter=30,
            )
            story.append(Paragraph(f"Template: {template_id}", title_style))
            story.append(Spacer(1, 12))

            # Add content
            for text in text_content:
                if text.strip():
                    p = Paragraph(text, styles["Normal"])
                    story.append(p)
                    story.append(Spacer(1, 12))

            doc_pdf.build(story)
            pdf_content = buffer.getvalue()
            buffer.close()

            return pdf_content

        except Exception as e:
            logger.warning(f"Text-to-PDF conversion failed: {e}")
            return None

            # Upload PDF to R2 (placeholder - would need R2 service)
            pdf_url = await self._upload_to_r2(
                pdf_content, f"templates/{template_id}/preview.pdf"
            )

            return pdf_url

        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            return None

    async def analyze_template_with_gemini_vision(
        self, pdf_content: bytes, template_id: str
    ) -> TemplateAnalysisResult:
        """Phân tích template với Gemini Vision API sử dụng PDF"""
        try:
            start_time = datetime.now()
            timestamp = start_time.strftime("%Y%m%d_%H%M%S")

            # Save PDF content for debugging with permission handling
            pdf_debug_path = f"{LOGS_DIR}/pdf_input_{template_id}_{timestamp}.pdf"
            if safe_file_write(pdf_debug_path, pdf_content, "wb"):
                logger.info(f"📁 Saved PDF for debugging: {pdf_debug_path}")
            else:
                logger.info("💡 Debug PDF save skipped due to permissions")

            # Enhanced AI prompt for PDF analysis (Gemini 2.5 Pro can read PDF directly)
            analysis_prompt = f"""
            Phân tích file PDF template báo giá này và trả về JSON với cấu trúc chi tiết:

            NHIỆM VỤ: Convert template tĩnh này thành smart template với placeholders

            YÊU CẦU PHÂN TÍCH:
            1. PLACEHOLDER DETECTION:
               - Tìm tất cả dữ liệu có thể thay đổi (tên, địa chỉ, số tiền, ngày tháng, sản phẩm)
               - Thay thế bằng {{{{placeholder_name}}}} format
               - Phân loại type: text|currency|date|number|email|phone|calculated

            2. STRUCTURE ANALYSIS:
               - Xác định sections: header, customer_info, products, financial, footer
               - Detect table structures và repeatable rows
               - Map visual layout và formatting requirements

            3. BUSINESS LOGIC DETECTION:
               - Tìm calculation patterns (subtotal → VAT → total)
               - Xác định auto-generated fields (dates, reference numbers)
               - Detect conditional fields (discounts, special terms)

            TRẢ VỀ JSON FORMAT:
            {{
              "placeholders": {{
                "{{{{company_name}}}}": {{
                  "type": "text",
                  "description": "Tên công ty phát hành báo giá",
                  "current_value": "CÔNG TY TNHH ABC",
                  "section": "company_info",
                  "validation_rules": ["required"],
                  "auto_populate": false,
                  "position": {{"page": 1, "section": "header"}},
                  "formatting": {{"bold": true, "font_size": 14}}
                }},
                "{{{{total_amount}}}}": {{
                  "type": "currency",
                  "description": "Tổng tiền cuối cùng",
                  "current_value": "50,000,000",
                  "section": "financial",
                  "calculation_formula": "=subtotal + vat_amount - discount_amount",
                  "formatting": {{"bold": true, "color": "red", "alignment": "right"}}
                }}
              }},
              "sections": [
                {{
                  "name": "header",
                  "description": "Phần header với logo và thông tin công ty",
                  "placeholders": ["{{{{company_name}}}}", "{{{{company_address}}}}"],
                  "order": 1,
                  "is_repeatable": false,
                  "required": true
                }},
                {{
                  "name": "products_table",
                  "description": "Bảng danh sách sản phẩm/dịch vụ",
                  "placeholders": ["{{{{products[].name}}}}", "{{{{products[].price}}}}"],
                  "order": 3,
                  "is_repeatable": true,
                  "required": true,
                  "table_structure": {{
                    "columns": ["STT", "Sản phẩm", "Số lượng", "Đơn giá", "Thành tiền"],
                    "repeatable_rows": true,
                    "calculation_columns": ["Thành tiền"]
                  }}
                }}
              ],
              "business_logic": {{
                "calculation_fields": ["{{{{subtotal}}}}", "{{{{vat_amount}}}}", "{{{{total_amount}}}}"],
                "auto_fill_fields": ["{{{{current_date}}}}", "{{{{quote_number}}}}"],
                "conditional_fields": ["{{{{discount_amount}}}}"],
                "validation_rules": {{
                  "{{{{customer_email}}}}": ["email"],
                  "{{{{total_amount}}}}": ["numeric", "greater_than_zero"]
                }}
              }},
              "document_structure": {{
                "total_pages": 2,
                "has_tables": true,
                "table_locations": ["products_section"],
                "header_content": "Company logo and info",
                "footer_content": "Terms and signature line",
                "visual_elements": ["logo_placeholder", "signature_line"]
              }}
            }}

            QUAN TRỌNG:
            - Tìm mọi thông tin có thể thay đổi và convert thành {{{{placeholder}}}}
            - Phân tích table structures để support dynamic product lists
            - Detect calculation patterns và preserve formatting
            - Identify conditional sections (show/hide based on data)
            """

            # Save AI request payload for debugging
            ai_request_payload = {
                "template_id": template_id,
                "timestamp": timestamp,
                "prompt": analysis_prompt,
                "model": "gemini-2.5-pro",
                "max_tokens": 32000,
                "temperature": 1.0,
                "pdf_size_bytes": len(pdf_content),
                "pdf_file_mode": "direct_pdf_analysis",
            }

            payload_log_path = (
                f"{LOGS_DIR}/ai_request_payload_{template_id}_{timestamp}.json"
            )
            with open(payload_log_path, "w", encoding="utf-8") as f:
                json.dump(ai_request_payload, f, indent=2, ensure_ascii=False)
            logger.info(f"💾 Saved AI request payload: {payload_log_path}")

            # Call Gemini 2.5 Pro API with PDF file directly
            logger.info(f"🤖 Calling Gemini 2.5 Pro for PDF template analysis...")

            # Use new Google Gen AI SDK for direct PDF processing
            try:
                from google import genai
                from google.genai import types
                import os

                # Configure new Gemini client with API key
                api_key = os.getenv("GEMINI_API_KEY")
                if not api_key:
                    raise Exception("GEMINI_API_KEY not found in environment")

                client = genai.Client(api_key=api_key)

                # Inline PDF processing - pass PDF data directly
                logger.info(f"📤 Processing PDF with new Gemini SDK...")

                # Create PDF part from bytes
                pdf_part = types.Part.from_bytes(
                    data=pdf_content, mime_type="application/pdf"
                )

                # Generate response with inline PDF
                response = client.models.generate_content(
                    model="gemini-2.5-pro",
                    contents=[pdf_part, analysis_prompt],
                    config=types.GenerateContentConfig(
                        max_output_tokens=8000,
                        temperature=0.1,  # Low temperature for accurate analysis
                        response_mime_type="application/json",  # Ensure JSON response
                    ),
                )

                # Get response text
                if hasattr(response, "text") and response.text:
                    ai_response = response.text
                    logger.info(
                        f"✅ Gemini 2.5 Pro response received: {len(ai_response)} characters"
                    )
                else:
                    raise Exception("No text response from Gemini API")

                # Log usage metadata if available
                if hasattr(response, "usage_metadata") and response.usage_metadata:
                    logger.info(
                        f"📊 Token usage - Prompt: {response.usage_metadata.prompt_token_count}, "
                        f"Response: {response.usage_metadata.candidates_token_count}"
                    )

            except Exception as gemini_error:
                logger.error(f"❌ Gemini API error: {gemini_error}")
                # Fallback to text-based analysis
                ai_response = '{"placeholders": {}, "sections": [], "business_logic": {}, "document_structure": {"total_pages": 1, "has_tables": false}}'

            # Save full AI response for debugging
            ai_response_data = {
                "template_id": template_id,
                "timestamp": timestamp,
                "request_payload": ai_request_payload,
                "raw_response": ai_response,
                "response_length": len(ai_response),
                "processing_start": start_time.isoformat(),
                "response_received": datetime.now().isoformat(),
            }

            response_log_path = (
                f"{LOGS_DIR}/ai_response_raw_{template_id}_{timestamp}.json"
            )
            with open(response_log_path, "w", encoding="utf-8") as f:
                json.dump(ai_response_data, f, indent=2, ensure_ascii=False)
            logger.info(f"📝 Saved AI raw response: {response_log_path}")

            # Parse AI response
            try:
                analysis_data = json.loads(ai_response.strip())
                logger.info(f"✅ AI response parsed successfully")
            except json.JSONDecodeError as e:
                # Fallback to pattern-based analysis
                logger.warning(
                    f"❌ AI response is not valid JSON: {e}, falling back to pattern analysis"
                )
                analysis_data = await self._pattern_based_analysis_fallback(
                    pdf_content, template_id
                )

            # Save parsed analysis data
            parsed_log_path = (
                f"{LOGS_DIR}/ai_analysis_parsed_{template_id}_{timestamp}.json"
            )
            with open(parsed_log_path, "w", encoding="utf-8") as f:
                json.dump(analysis_data, f, indent=2, ensure_ascii=False)
            logger.info(f"📊 Saved parsed analysis: {parsed_log_path}")

            # Calculate processing time and confidence score
            processing_time = (datetime.now() - start_time).total_seconds()
            confidence_score = self._calculate_confidence_score(analysis_data)

            # Create TemplateAnalysisResult
            placeholders = {}
            for name, info in analysis_data.get("placeholders", {}).items():
                placeholders[name] = PlaceholderInfo(**info)

            sections = []
            for section_data in analysis_data.get("sections", []):
                sections.append(TemplateSection(**section_data))

            document_structure = DocumentStructure(
                **(analysis_data.get("document_structure", {}))
            )

            result = TemplateAnalysisResult(
                template_id=template_id,
                placeholders=placeholders,
                sections=sections,
                business_logic=analysis_data.get("business_logic", {}),
                document_structure=document_structure,
                ai_analysis_score=confidence_score,
                processing_time=processing_time,
            )

            # Save final result for debugging
            final_result_data = {
                "template_id": template_id,
                "timestamp": timestamp,
                "result": {
                    "placeholders_count": len(placeholders),
                    "sections_count": len(sections),
                    "confidence_score": confidence_score,
                    "processing_time": processing_time,
                    "business_logic_fields": len(
                        analysis_data.get("business_logic", {})
                    ),
                    "document_structure": analysis_data.get("document_structure", {}),
                },
                "full_analysis": analysis_data,
            }

            final_log_path = (
                f"{LOGS_DIR}/template_analysis_result_{template_id}_{timestamp}.json"
            )
            with open(final_log_path, "w", encoding="utf-8") as f:
                json.dump(final_result_data, f, indent=2, ensure_ascii=False)
            logger.info(f"🎯 Saved final analysis result: {final_log_path}")

            return result

        except Exception as e:
            logger.error(f"❌ Error in Gemini Vision analysis: {e}")
            # Save error log
            error_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            error_log_path = (
                f"{LOGS_DIR}/ai_analysis_error_{template_id}_{error_timestamp}.json"
            )
            error_data = {
                "template_id": template_id,
                "timestamp": error_timestamp,
                "error": str(e),
                "error_type": type(e).__name__,
                "pdf_size": len(pdf_content) if pdf_content else 0,
            }
            with open(error_log_path, "w", encoding="utf-8") as f:
                json.dump(error_data, f, indent=2, ensure_ascii=False)

            # Fallback to basic analysis
            return await self._create_fallback_analysis(template_id)

    async def _pattern_based_analysis_fallback(
        self, pdf_content: bytes, template_id: str
    ) -> Dict[str, Any]:
        """Fallback analysis when AI fails"""
        # Basic pattern-based analysis for fallback
        return {
            "placeholders": {
                "{{company_name}}": {
                    "type": "text",
                    "description": "Company name",
                    "current_value": "",
                    "section": "company_info",
                    "validation_rules": ["required"],
                    "auto_populate": False,
                }
            },
            "sections": [
                {
                    "name": "header",
                    "description": "Header section",
                    "placeholders": ["{{company_name}}"],
                    "order": 1,
                    "is_repeatable": False,
                    "required": True,
                }
            ],
            "business_logic": {
                "calculation_fields": [],
                "auto_fill_fields": [],
                "conditional_fields": [],
                "validation_rules": {},
            },
            "document_structure": {
                "total_pages": 1,
                "has_tables": False,
                "table_locations": [],
                "header_content": "",
                "footer_content": "",
                "visual_elements": [],
            },
        }

    async def _create_fallback_analysis(
        self, template_id: str
    ) -> TemplateAnalysisResult:
        """Create fallback analysis result"""
        return TemplateAnalysisResult(
            template_id=template_id,
            placeholders={},
            sections=[],
            business_logic={},
            document_structure=DocumentStructure(),
            ai_analysis_score=0.5,
            processing_time=1.0,
        )

    def _calculate_confidence_score(self, analysis_data: Dict[str, Any]) -> float:
        """Calculate confidence score for AI analysis"""
        score = 0.0

        # Check if placeholders were found
        placeholders = analysis_data.get("placeholders", {})
        if len(placeholders) > 0:
            score += 0.3

        # Check if sections were identified
        sections = analysis_data.get("sections", [])
        if len(sections) > 0:
            score += 0.3

        # Check if business logic was identified
        business_logic = analysis_data.get("business_logic", {})
        if business_logic:
            score += 0.2

        # Check if document structure was analyzed
        doc_structure = analysis_data.get("document_structure", {})
        if doc_structure:
            score += 0.2

        return min(score, 1.0)

    def create_templated_docx(
        self, original_file_content: bytes, placeholders: Dict[str, PlaceholderInfo]
    ) -> bytes:
        """Create a templated version of the DOCX with placeholders replaced by {{placeholder_name}}"""
        try:
            import io
            from docx import Document

            logger.info(
                f"Creating templated DOCX with {len(placeholders)} placeholders"
            )

            # Load the original document
            doc = Document(io.BytesIO(original_file_content))

            # Create mapping from current values to placeholder names with priority handling
            value_to_placeholder = {}
            # Track conflicts to resolve them with priority
            placeholder_priority = {
                "total_amount": 20,  # Highest priority for final total
                "company_name": 15,
                "customer_name": 14,
                "product_description": 13,  # High priority for product details
                "product_name": 12,
                "unit": 11,
                "subtotal": 10,
                "vat_amount": 9,
                "unit_price": 8,
                "line_total": 7,
                "total_amount_in_words": 6,
                "notes": 5,
                "signer_name": 4,
                "signer_title": 3,
                "quantity": 2,  # Higher priority than before
                # Sender info
                "sender_company_name": 16,
                "sender_company_address": 15,
                "sender_phone": 14,
                "sender_email": 14,
                "sender_website": 14,
                # Customer info
                "customer_company_name": 15,
                # Financial
                "vat_rate": 8,
                "validity_terms": 5,
                "issue_city": 4,
                "issue_date": 4,
                "sender_representative_name": 3,
                "sender_representative_position": 3,
            }

            # Skip generic values that cause conflicts, but with special handling
            skip_values = set()  # Don't skip "1" completely, handle with context
            quantity_specific_context = False  # Track if we're in quantity context

            for placeholder_name, info in placeholders.items():
                if info.current_value:
                    # Handle different data types
                    current_val = info.current_value

                    # Convert array-style placeholders to simple template placeholders
                    # e.g., "{{products[].name}}" -> "{{product_name}}"
                    template_name = placeholder_name
                    simple_name = placeholder_name
                    if "[].name" in placeholder_name:
                        template_name = placeholder_name.replace(
                            "{{products[].name}}", "{{product_name}}"
                        )
                        simple_name = "product_name"
                    elif "[].description" in placeholder_name:
                        template_name = placeholder_name.replace(
                            "{{products[].description}}", "{{product_description}}"
                        )
                        simple_name = "product_description"
                    elif "[].unit" in placeholder_name:
                        template_name = placeholder_name.replace(
                            "{{products[].unit}}", "{{unit}}"
                        )
                        simple_name = "unit"
                    elif "[].quantity" in placeholder_name:
                        template_name = placeholder_name.replace(
                            "{{products[].quantity}}", "{{quantity}}"
                        )
                        simple_name = "quantity"
                    elif "[].unit_price" in placeholder_name:
                        template_name = placeholder_name.replace(
                            "{{products[].unit_price}}", "{{unit_price}}"
                        )
                        simple_name = "unit_price"
                    elif "[].line_total" in placeholder_name:
                        template_name = placeholder_name.replace(
                            "{{products[].line_total}}", "{{line_total}}"
                        )
                        simple_name = "line_total"
                    else:
                        # Extract simple name from regular placeholders
                        simple_name = placeholder_name.replace("{{", "").replace(
                            "}}", ""
                        )

                    if isinstance(current_val, str) and current_val.strip():
                        # Skip generic values that cause too many conflicts
                        if current_val in skip_values:
                            logger.info(
                                f"🚫 Skipping generic value: '{current_val}' -> '{template_name}'"
                            )
                            continue

                        # Check for conflicts and prioritize
                        if current_val in value_to_placeholder:
                            existing_name = (
                                value_to_placeholder[current_val]
                                .replace("{{", "")
                                .replace("}}", "")
                            )
                            current_priority = placeholder_priority.get(simple_name, 0)
                            existing_priority = placeholder_priority.get(
                                existing_name, 0
                            )

                            if current_priority > existing_priority:
                                value_to_placeholder[current_val] = template_name
                                logger.info(
                                    f"🔄 Priority override: '{current_val}' -> '{template_name}' (was {existing_name})"
                                )
                            else:
                                logger.info(
                                    f"🔄 Skipping: '{current_val}' -> '{template_name}' (keeping {existing_name})"
                                )
                        else:
                            value_to_placeholder[current_val] = template_name
                            logger.info(
                                f"🔄 Mapping: '{current_val}' -> '{template_name}'"
                            )
                    elif isinstance(current_val, (int, float)):
                        str_val = str(current_val)
                        # Skip generic single digits
                        if str_val in skip_values:
                            logger.info(
                                f"🚫 Skipping generic value: '{str_val}' -> '{template_name}'"
                            )
                            continue

                        # Same conflict resolution for numbers
                        if str_val in value_to_placeholder:
                            existing_name = (
                                value_to_placeholder[str_val]
                                .replace("{{", "")
                                .replace("}}", "")
                            )
                            current_priority = placeholder_priority.get(simple_name, 0)
                            existing_priority = placeholder_priority.get(
                                existing_name, 0
                            )

                            if current_priority > existing_priority:
                                value_to_placeholder[str_val] = template_name
                                logger.info(
                                    f"🔄 Priority override: '{str_val}' -> '{template_name}' (was {existing_name})"
                                )
                            else:
                                logger.info(
                                    f"🔄 Skipping: '{str_val}' -> '{template_name}' (keeping {existing_name})"
                                )
                        else:
                            value_to_placeholder[str_val] = template_name
                            logger.info(f"🔄 Mapping: '{str_val}' -> '{template_name}'")

            logger.info(
                f"Value to placeholder mapping: {len(value_to_placeholder)} entries"
            )

            # Replace text in paragraphs
            replaced_count = 0
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    original_text = paragraph.text
                    updated_text = original_text

                    for current_value, placeholder_name in value_to_placeholder.items():
                        # Handle exact text replacement first
                        if current_value in updated_text:
                            updated_text = updated_text.replace(
                                current_value, placeholder_name
                            )
                            logger.info(
                                f"✅ Replaced in paragraph: '{current_value[:50]}...' -> '{placeholder_name}'"
                            )
                            replaced_count += 1
                        else:
                            # Handle multi-line text by normalizing whitespace for longer texts
                            normalized_original = " ".join(original_text.split())
                            normalized_value = " ".join(current_value.split())

                            if (
                                len(current_value) > 20
                                and normalized_value in normalized_original
                            ):
                                # For longer texts, try normalized replacement
                                updated_text = updated_text.replace(
                                    current_value, placeholder_name
                                )
                                logger.info(
                                    f"✅ Replaced normalized in paragraph: '{current_value[:50]}...' -> '{placeholder_name}'"
                                )
                                replaced_count += 1

                    if updated_text != original_text:
                        # Clear existing text and add updated text
                        paragraph.clear()
                        paragraph.add_run(updated_text)
                        logger.debug(
                            f"Updated paragraph: {original_text[:50]}... -> {updated_text[:50]}..."
                        )

            # Replace text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if paragraph.text:
                                original_text = paragraph.text
                                updated_text = original_text

                                for (
                                    current_value,
                                    placeholder_name,
                                ) in value_to_placeholder.items():
                                    # Handle exact text replacement first
                                    if current_value in updated_text:
                                        updated_text = updated_text.replace(
                                            current_value, placeholder_name
                                        )
                                        logger.info(
                                            f"✅ Replaced in table: '{current_value[:30]}...' -> '{placeholder_name}'"
                                        )
                                        replaced_count += 1
                                    else:
                                        # Handle multi-line text in table cells for longer texts
                                        normalized_original = " ".join(
                                            original_text.split()
                                        )
                                        normalized_value = " ".join(
                                            current_value.split()
                                        )

                                        if (
                                            len(current_value) > 20
                                            and normalized_value in normalized_original
                                        ):
                                            updated_text = updated_text.replace(
                                                current_value, placeholder_name
                                            )
                                            logger.info(
                                                f"✅ Replaced normalized in table: '{current_value[:30]}...' -> '{placeholder_name}'"
                                            )
                                            replaced_count += 1

                                if updated_text != original_text:
                                    paragraph.clear()
                                    paragraph.add_run(updated_text)
                                    logger.debug(
                                        f"Updated table cell: {original_text[:30]}... -> {updated_text[:30]}..."
                                    )

            logger.info(f"🎯 Total replacements made: {replaced_count}")

            # Save templated document to bytes
            templated_buffer = io.BytesIO()
            doc.save(templated_buffer)
            templated_content = templated_buffer.getvalue()

            logger.info(
                f"✅ Successfully created templated DOCX ({len(templated_content)} bytes)"
            )
            return templated_content

        except Exception as e:
            logger.error(f"❌ Failed to create templated DOCX: {e}")
            # Return original content as fallback
            return original_file_content

    async def _upload_to_r2(self, content: bytes, key: str) -> str:
        """Upload file to R2 storage using AIVungtau R2 config"""
        try:
            from src.config.r2_storage import AIVungtauR2StorageConfig

            r2_config = AIVungtauR2StorageConfig()

            if not r2_config.s3_client:
                logger.warning(
                    "R2 client not initialized, using fallback local storage"
                )
                return f"https://fallback.local/{key}"

            # Upload to R2
            r2_config.s3_client.put_object(
                Bucket=r2_config.bucket_name,
                Key=key,
                Body=content,
                ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

            # Return public URL
            public_url = f"https://{r2_config.public_domain}/{key}"
            logger.info(f"✅ File uploaded to R2: {public_url}")
            return public_url

        except Exception as e:
            logger.error(f"❌ R2 upload failed: {e}")
            # Return fallback URL
            return f"https://fallback.local/{key}"

    async def save_template_to_database(
        self,
        template_id: str,
        user_id: str,
        template_name: str,
        description: str,
        category: str,
        original_docx_url: str,
        template_docx_url: Optional[str],
        pdf_content: Optional[bytes],
        analysis_result: TemplateAnalysisResult,
        is_system_template: bool = False,
    ) -> Dict[str, Any]:
        """Save template metadata to database"""
        try:
            db = await self._get_database()

            # Determine template visibility and ownership
            if is_system_template:
                template_user_id = "system"
                is_public = True
                logger.info(f"📋 Creating system template: {template_name}")
            else:
                template_user_id = user_id
                is_public = False
                logger.info(
                    f"👤 Creating user template: {template_name} for user {user_id}"
                )

            template_doc = {
                "_id": template_id,
                "user_id": template_user_id,
                "name": template_name,
                "description": description,
                "category": category,
                "type": "quote",
                "subtype": "business",
                # Template visibility
                "is_public": is_public,
                "is_system_template": is_system_template,
                # File information
                "files": {
                    "original_docx_url": original_docx_url,
                    "template_docx_url": template_docx_url,
                    "pdf_conversion_success": pdf_content is not None,
                    "thumbnail_urls": [],
                },
                # AI Analysis results
                "ai_analysis": {
                    "placeholders": {
                        name: {
                            "type": info.type,
                            "description": info.description,
                            "current_value": info.current_value,
                            "default_value": info.default_value,
                            "validation_rules": info.validation_rules,
                            "section": info.section,
                            "auto_populate": info.auto_populate,
                            "calculation_formula": info.calculation_formula,
                            "position": info.position,
                            "formatting": info.formatting,
                        }
                        for name, info in analysis_result.placeholders.items()
                    },
                    "sections": [
                        {
                            "name": section.name,
                            "description": section.description,
                            "placeholders": section.placeholders,
                            "order": section.order,
                            "is_repeatable": section.is_repeatable,
                            "required": section.required,
                            "table_structure": section.table_structure,
                        }
                        for section in analysis_result.sections
                    ],
                    "business_logic": analysis_result.business_logic,
                    "document_structure": {
                        "total_pages": analysis_result.document_structure.total_pages,
                        "has_tables": analysis_result.document_structure.has_tables,
                        "table_locations": analysis_result.document_structure.table_locations,
                        "header_content": analysis_result.document_structure.header_content,
                        "footer_content": analysis_result.document_structure.footer_content,
                        "visual_elements": analysis_result.document_structure.visual_elements,
                    },
                    "confidence_score": analysis_result.ai_analysis_score,
                    "analysis_version": "2.0",
                },
                # Metadata
                "is_active": True,
                "is_public": False,
                "usage_count": 0,
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow(),
                # Template validation
                "validation": {"is_valid": True, "errors": [], "warnings": []},
            }

            result = await db.document_templates.insert_one(template_doc)

            logger.info(f"✅ Template saved to database: {template_id}")

            return {
                "success": True,
                "template_id": template_id,
                "database_id": str(result.inserted_id),
            }

        except Exception as e:
            logger.error(f"Error saving template to database: {e}")
            raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

    async def process_template_upload(
        self,
        file: UploadFile = None,
        file_content: bytes = None,
        filename: str = None,
        template_name: str = None,
        description: str = "",
        category: str = "standard",
        user_id: str = None,
        is_system_template: bool = False,
    ) -> Dict[str, Any]:
        """
        Main method to process template upload với PDF workflow

        Args:
            file: UploadFile object (for regular user uploads)
            file_content: Raw file bytes (for admin system uploads)
            filename: Original filename
            template_name: Name of the template
            description: Template description
            category: Template category
            user_id: User ID ("system" for system templates)
            is_system_template: Flag to indicate system template
        """
        try:
            # Generate template ID
            template_id = f"template_{uuid.uuid4().hex[:12]}"
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            logger.info(f"🔄 Processing template upload: {template_id}")
            logger.info(
                f"📋 Template type: {'System' if is_system_template else 'User'}"
            )

            # Handle both file upload and direct content
            if file is not None:
                actual_filename = file.filename
                actual_content_type = file.content_type
                actual_size = file.size if hasattr(file, "size") else "unknown"
            else:
                actual_filename = filename
                actual_content_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                actual_size = len(file_content) if file_content else 0

            # Log upload metadata
            upload_metadata = {
                "template_id": template_id,
                "timestamp": timestamp,
                "file_info": {
                    "filename": actual_filename,
                    "content_type": actual_content_type,
                    "size": actual_size,
                },
                "template_info": {
                    "name": template_name,
                    "description": description,
                    "category": category,
                    "user_id": user_id,
                    "is_system_template": is_system_template,
                },
            }

            metadata_log_path = (
                f"{LOGS_DIR}/upload_metadata_{template_id}_{timestamp}.json"
            )
            with open(metadata_log_path, "w", encoding="utf-8") as f:
                json.dump(upload_metadata, f, indent=2, ensure_ascii=False)
            logger.info(f"📋 Saved upload metadata: {metadata_log_path}")

            # Step 1: Validate file (skip for admin uploads with file_content)
            if file is not None:
                validation_result = await self.validate_docx_file(file)
                if not validation_result["is_valid"]:
                    error_log_path = (
                        f"{LOGS_DIR}/validation_error_{template_id}_{timestamp}.json"
                    )
                    with open(error_log_path, "w", encoding="utf-8") as f:
                        json.dump(
                            {
                                "template_id": template_id,
                                "timestamp": timestamp,
                                "validation_result": validation_result,
                            },
                            f,
                            indent=2,
                            ensure_ascii=False,
                        )

                    raise HTTPException(
                        status_code=400,
                        detail=f"File validation failed: {validation_result['errors']}",
                    )

            # Step 2: Get file content
            if file is not None:
                docx_content = await file.read()
                await file.seek(0)
            else:
                docx_content = file_content

            # Generate UUID for original file
            original_uuid = str(uuid.uuid4())

            # Save original DOCX for debugging with permission handling
            docx_debug_path = f"{LOGS_DIR}/docx_original_{template_id}_{timestamp}.docx"
            if safe_file_write(docx_debug_path, docx_content, "wb"):
                logger.info(f"📁 Saved original DOCX: {docx_debug_path}")
            else:
                logger.info("💡 Debug DOCX save skipped due to permissions")

            # Upload original file with UUID name
            original_docx_url = await self._upload_to_r2(
                docx_content, f"templates/{user_id}/{template_id}/{original_uuid}.docx"
            )
            logger.info(f"☁️ Original DOCX uploaded to R2: {original_docx_url}")

            # Step 3: Convert DOCX to PDF
            pdf_content = await self.convert_docx_to_pdf(docx_content, template_id)
            logger.info(
                f"📄 PDF conversion result: {'Success' if pdf_content else 'Failed'}"
            )

            # Step 4: AI Analysis với Gemini Vision
            if pdf_content:
                analysis_result = await self.analyze_template_with_gemini_vision(
                    pdf_content, template_id
                )
            else:
                # Fallback if PDF conversion failed
                logger.warning(f"⚠️ PDF conversion failed, using fallback analysis")
                analysis_result = await self._create_fallback_analysis(template_id)

            logger.info(
                f"🧠 AI analysis completed with score: {analysis_result.ai_analysis_score}"
            )

            # Step 5: Create templated DOCX with placeholders
            try:
                templated_docx_content = self.create_templated_docx(
                    docx_content, analysis_result.placeholders
                )

                # Upload templated version
                template_docx_url = await self._upload_to_r2(
                    templated_docx_content,
                    f"templates/{user_id}/{template_id}/template.docx",
                )
                logger.info(f"📝 Template DOCX uploaded to R2: {template_docx_url}")

            except Exception as e:
                logger.error(f"❌ Failed to create templated DOCX: {e}")
                template_docx_url = None

            # Step 6: Save to database
            db_result = await self.save_template_to_database(
                template_id=template_id,
                user_id=user_id,
                template_name=template_name,
                description=description,
                category=category,
                original_docx_url=original_docx_url,
                template_docx_url=template_docx_url,
                pdf_content=pdf_content,
                analysis_result=analysis_result,
                is_system_template=is_system_template,
            )

            # Prepare final response
            final_response = {
                "success": True,
                "template_id": template_id,
                "status": "completed",
                "message": "Template uploaded and analyzed successfully",
                "urls": {
                    "original_docx": original_docx_url,
                    "template_docx": template_docx_url,
                    "pdf_conversion_success": pdf_content is not None,
                    "status_check": f"/api/templates/{template_id}/status",
                },
                "analysis_summary": {
                    "placeholders_found": len(analysis_result.placeholders),
                    "sections_identified": len(analysis_result.sections),
                    "confidence_score": analysis_result.ai_analysis_score,
                    "processing_time": analysis_result.processing_time,
                    "has_tables": analysis_result.document_structure.has_tables,
                    "total_pages": analysis_result.document_structure.total_pages,
                },
                "warnings": validation_result.get("warnings", []),
            }

            # Save complete response log with all URLs and details
            complete_response_log = {
                "template_id": template_id,
                "timestamp": timestamp,
                "processing_completed": datetime.now().isoformat(),
                "file_urls": {
                    "original_docx": original_docx_url,
                    "template_docx": template_docx_url,
                    "pdf_conversion_success": pdf_content is not None,
                    "r2_bucket_info": "cloudflare-r2",
                    "access_method": "authenticated",
                },
                "database_info": {
                    "collection": "document_templates",
                    "inserted_id": str(db_result.get("database_id", "")),
                },
                "full_response": final_response,
                "processing_stats": {
                    "total_placeholders": len(analysis_result.placeholders),
                    "total_sections": len(analysis_result.sections),
                    "ai_confidence": analysis_result.ai_analysis_score,
                    "processing_time_seconds": analysis_result.processing_time,
                    "file_size_bytes": len(docx_content),
                    "pdf_conversion_success": pdf_content is not None,
                },
            }

            response_log_path = (
                f"{LOGS_DIR}/complete_response_{template_id}_{timestamp}.json"
            )
            with open(response_log_path, "w", encoding="utf-8") as f:
                json.dump(complete_response_log, f, indent=2, ensure_ascii=False)
            logger.info(f"🎯 Saved complete response log: {response_log_path}")

            return final_response

        except HTTPException:
            raise
        except Exception as e:
            # Save error log
            error_timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            # Use template_id if available, otherwise use fallback
            error_template_id = locals().get("template_id", "unknown_template")
            error_log_path = (
                f"{LOGS_DIR}/upload_error_{error_template_id}_{error_timestamp}.json"
            )
            error_data = {
                "template_id": error_template_id,
                "timestamp": error_timestamp,
                "error": str(e),
                "error_type": type(e).__name__,
                "file_info": {
                    "filename": file.filename,
                    "content_type": file.content_type,
                },
            }
            with open(error_log_path, "w", encoding="utf-8") as f:
                json.dump(error_data, f, indent=2, ensure_ascii=False)

            logger.error(f"❌ Template upload processing failed: {e}")
            raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")
