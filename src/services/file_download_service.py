"""
File Download and Parsing Service
Downloads files from R2 URLs, parses them to text, and manages temp files
"""

import os
import tempfile
import logging
import aiohttp
import asyncio
import boto3
from typing import Optional, Tuple
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

logger = logging.getLogger("chatbot")

# R2 Configuration
R2_ACCESS_KEY_ID = os.getenv("R2_ACCESS_KEY_ID")
R2_SECRET_ACCESS_KEY = os.getenv("R2_SECRET_ACCESS_KEY")
R2_BUCKET_NAME = os.getenv("R2_BUCKET_NAME", "wordai")
R2_ENDPOINT_URL = os.getenv("R2_ENDPOINT")

# Initialize R2 client with credentials
try:
    if all([R2_ACCESS_KEY_ID, R2_SECRET_ACCESS_KEY, R2_ENDPOINT_URL]):
        s3_client = boto3.client(
            "s3",
            endpoint_url=R2_ENDPOINT_URL,
            aws_access_key_id=R2_ACCESS_KEY_ID,
            aws_secret_access_key=R2_SECRET_ACCESS_KEY,
            region_name="auto",
        )
        logger.info("‚úÖ R2 S3 client initialized for file downloads")
    else:
        s3_client = None
        logger.warning("‚ö†Ô∏è R2 credentials missing - boto3 download will not work")
except Exception as e:
    s3_client = None
    logger.error(f"‚ùå Failed to initialize R2 client: {e}")


class FileDownloadService:
    """
    Service for downloading and parsing files from R2 storage

    Supports:
    - TXT: Direct text (no parsing needed)
    - DOCX: Extract text using python-docx
    - PDF: Extract text using PyPDF2

    Special Cases:
    - PDF + Gemini: Can use file path directly (no parsing needed)
    """

    # Track temp files per user for cleanup
    _user_temp_files: dict = {}

    @classmethod
    async def download_and_parse_file_from_r2(
        cls, r2_key: str, file_type: str, user_id: str
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Download file from R2 using boto3 (with credentials) and parse to text

        Args:
            r2_key: R2 object key (e.g. "files/user123/root/file_abc/doc.pdf")
            file_type: File extension (txt, docx, pdf)
            user_id: User ID for temp file tracking

        Returns:
            (text_content, temp_file_path) or (None, None) if failed
        """
        try:
            if s3_client is None:
                logger.error("‚ùå R2 S3 client not initialized")
                return None, None

            file_type = file_type.lower().replace(".", "")

            # Cleanup old temp files for this user
            await cls._cleanup_user_temp_files(user_id)

            # Download file from R2 using boto3
            logger.info(f"üì• Downloading {file_type.upper()} from R2 key: {r2_key}")
            temp_file_path = await cls._download_file_from_r2_with_boto3(
                r2_key, file_type
            )

            if not temp_file_path:
                logger.error(f"‚ùå Failed to download file from R2 key: {r2_key}")
                return None, None

            # Track temp file for cleanup
            if user_id not in cls._user_temp_files:
                cls._user_temp_files[user_id] = []
            cls._user_temp_files[user_id].append(temp_file_path)

            logger.info(f"‚úÖ Downloaded to: {temp_file_path}")

            # Parse file to text
            logger.info(f"üìù Parsing {file_type.upper()} to text...")
            text_content = await cls._parse_file(temp_file_path, file_type)

            if text_content:
                logger.info(
                    f"‚úÖ Parsed {len(text_content)} characters from {file_type.upper()}"
                )
                return text_content, temp_file_path
            else:
                logger.error(f"‚ùå Failed to parse {file_type.upper()} file")
                return None, None

        except Exception as e:
            logger.error(f"‚ùå Error downloading/parsing file from R2: {e}")
            return None, None

    @classmethod
    async def download_and_parse_file(
        cls, file_url: str, file_type: str, user_id: str, provider: str = None
    ) -> Tuple[Optional[str], Optional[str]]:
        """
        Download file from R2 and parse to text

        Args:
            file_url: R2 URL of the file
            file_type: File extension (txt, docx, pdf)
            user_id: User ID for temp file tracking
            provider: AI provider name (for special cases like PDF+Gemini)

        Returns:
            (text_content, temp_file_path) or (None, None) if failed

        Note:
            - For PDF + Gemini: Returns (None, temp_file_path) - Gemini uses file directly
            - For other cases: Returns (text_content, temp_file_path)
        """
        try:
            file_type = file_type.lower().replace(".", "")

            # Cleanup old temp files for this user
            await cls._cleanup_user_temp_files(user_id)

            # Download file to temp directory
            logger.info(
                f"üì• Downloading {file_type.upper()} from R2: {file_url[:50]}..."
            )
            temp_file_path = await cls._download_file(file_url, file_type)

            if not temp_file_path:
                logger.error(f"‚ùå Failed to download file from {file_url}")
                return None, None

            # Track temp file for cleanup
            if user_id not in cls._user_temp_files:
                cls._user_temp_files[user_id] = []
            cls._user_temp_files[user_id].append(temp_file_path)

            logger.info(f"‚úÖ Downloaded to: {temp_file_path}")

            # Special case: PDF + Gemini - return file path, no parsing
            if file_type == "pdf" and provider and provider.lower() == "gemini":
                logger.info(f"üî• PDF + Gemini: Returning file path (no parsing)")
                return None, temp_file_path

            # Parse file to text
            logger.info(f"üìù Parsing {file_type.upper()} to text...")
            text_content = await cls._parse_file(temp_file_path, file_type)

            if text_content:
                logger.info(
                    f"‚úÖ Parsed {len(text_content)} characters from {file_type.upper()}"
                )
                return text_content, temp_file_path
            else:
                logger.error(f"‚ùå Failed to parse {file_type.upper()} file")
                return None, None

        except Exception as e:
            logger.error(f"‚ùå Error downloading/parsing file: {e}")
            return None, None

    @classmethod
    async def _download_file_from_r2_with_boto3(
        cls, r2_key: str, file_type: str
    ) -> Optional[str]:
        """
        Download file from R2 using boto3 with credentials

        Args:
            r2_key: R2 object key
            file_type: File extension for temp file

        Returns:
            Path to downloaded temp file
        """
        try:
            if s3_client is None:
                logger.error("‚ùå R2 S3 client not available")
                return None

            # Create temp file
            temp_dir = tempfile.gettempdir()
            temp_file = tempfile.NamedTemporaryFile(
                delete=False, suffix=f".{file_type}", dir=temp_dir
            )
            temp_file_path = temp_file.name
            temp_file.close()

            # Download from R2 using boto3 (with credentials from env)
            logger.info(f"üîê Downloading from R2 bucket: {R2_BUCKET_NAME}")

            # Run boto3 download in thread (boto3 is synchronous)
            await asyncio.to_thread(
                s3_client.download_file,
                R2_BUCKET_NAME,
                r2_key,
                temp_file_path,
            )

            logger.info(f"‚úÖ Downloaded {os.path.getsize(temp_file_path)} bytes")
            return temp_file_path

        except Exception as e:
            logger.error(f"‚ùå Error downloading from R2 with boto3: {e}")
            return None

    @classmethod
    async def _download_file(cls, url: str, file_type: str) -> Optional[str]:
        """
        Download file from URL to temp directory

        Returns:
            Path to downloaded temp file
        """
        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(
                    url, timeout=aiohttp.ClientTimeout(total=60)
                ) as response:
                    if response.status != 200:
                        logger.error(f"‚ùå Failed to download: HTTP {response.status}")
                        return None

                    # Create temp file
                    temp_dir = tempfile.gettempdir()
                    temp_file = tempfile.NamedTemporaryFile(
                        delete=False, suffix=f".{file_type}", dir=temp_dir
                    )

                    # Write content
                    content = await response.read()
                    temp_file.write(content)
                    temp_file.close()

                    return temp_file.name

        except asyncio.TimeoutError:
            logger.error(f"‚ùå Timeout downloading file from {url}")
            return None
        except Exception as e:
            logger.error(f"‚ùå Error downloading file: {e}")
            return None

    @classmethod
    async def _parse_file(cls, file_path: str, file_type: str) -> Optional[str]:
        """
        Parse file to text based on file type

        Args:
            file_path: Path to temp file
            file_type: File extension (txt, docx, pdf)

        Returns:
            Extracted text content
        """
        try:
            if file_type == "txt":
                return await cls._parse_txt(file_path)
            elif file_type == "docx":
                return await cls._parse_docx(file_path)
            elif file_type == "pdf":
                return await cls._parse_pdf(file_path)
            else:
                logger.error(f"‚ùå Unsupported file type: {file_type}")
                return None

        except Exception as e:
            logger.error(f"‚ùå Error parsing {file_type} file: {e}")
            return None

    @classmethod
    async def _parse_txt(cls, file_path: str) -> Optional[str]:
        """Parse TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    @classmethod
    async def _parse_docx(cls, file_path: str) -> Optional[str]:
        """Parse DOCX file using python-docx with page breaks and font info"""
        try:
            from docx import Document
            from docx.oxml.text.paragraph import CT_P
            from docx.oxml.ns import qn

            doc = Document(file_path)

            # Extract text from paragraphs with page break detection
            text_parts = []
            current_page = []
            page_num = 1

            for paragraph in doc.paragraphs:
                # Check for page break
                if paragraph._element.xpath('.//w:br[@w:type="page"]'):
                    # End current page
                    if current_page:
                        text_parts.append(
                            f"[PAGE {page_num}]\n" + "\n".join(current_page)
                        )
                        page_num += 1
                        current_page = []

                if paragraph.text.strip():
                    # Try to get font size from first run
                    font_size = None
                    if paragraph.runs:
                        first_run = paragraph.runs[0]
                        if first_run.font.size:
                            # Font size in points
                            font_size = first_run.font.size.pt

                    # Format with font size if available
                    if font_size:
                        current_page.append(
                            f'<span style="font-size: {font_size}pt;">{paragraph.text}</span>'
                        )
                    else:
                        current_page.append(paragraph.text)

            # Add last page
            if current_page:
                text_parts.append(f"[PAGE {page_num}]\n" + "\n".join(current_page))

            # Extract text from tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            table_text.append(cell.text)

            if table_text:
                text_parts.append("\n".join(table_text))

            return (
                "\n\n[PAGE_BREAK]\n\n".join(text_parts)
                if len(text_parts) > 1
                else "\n".join(text_parts)
            )

        except ImportError:
            logger.error("‚ùå python-docx not installed. Run: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"‚ùå Error parsing DOCX: {e}")
            return None

    @classmethod
    async def _parse_pdf(cls, file_path: str) -> Optional[str]:
        """Parse PDF file using PyPDF2"""
        try:
            from PyPDF2 import PdfReader

            reader = PdfReader(file_path)

            # Extract text from all pages with page markers
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text.strip():
                    # Add page marker for document type formatting
                    text_parts.append(f"[PAGE {page_num}]\n{text}")

            return "\n\n[PAGE_BREAK]\n\n".join(text_parts)

        except ImportError:
            logger.error("‚ùå PyPDF2 not installed. Run: pip install PyPDF2")
            return None
        except Exception as e:
            logger.error(f"‚ùå Error parsing PDF: {e}")
            return None

    @classmethod
    async def _cleanup_user_temp_files(cls, user_id: str):
        """
        Delete old temp files for user when they switch to a new file

        Args:
            user_id: User ID
        """
        if user_id not in cls._user_temp_files:
            return

        old_files = cls._user_temp_files[user_id]
        deleted_count = 0

        for file_path in old_files:
            try:
                if os.path.exists(file_path):
                    os.remove(file_path)
                    deleted_count += 1
            except Exception as e:
                logger.warning(f"‚ö†Ô∏è Failed to delete temp file {file_path}: {e}")

        if deleted_count > 0:
            logger.info(f"üóëÔ∏è Cleaned up {deleted_count} temp files for user {user_id}")

        # Clear list
        cls._user_temp_files[user_id] = []

    @classmethod
    async def cleanup_all_temp_files(cls):
        """
        Cleanup all temp files (call on shutdown)
        """
        total_deleted = 0

        for user_id in list(cls._user_temp_files.keys()):
            for file_path in cls._user_temp_files[user_id]:
                try:
                    if os.path.exists(file_path):
                        os.remove(file_path)
                        total_deleted += 1
                except Exception as e:
                    logger.warning(f"‚ö†Ô∏è Failed to delete temp file {file_path}: {e}")

        cls._user_temp_files.clear()

        if total_deleted > 0:
            logger.info(f"üóëÔ∏è Cleaned up {total_deleted} temp files on shutdown")
