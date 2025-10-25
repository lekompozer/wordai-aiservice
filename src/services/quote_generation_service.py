"""
Service xử lý tạo và chỉnh sửa quote với AI
"""

import os
import json
import asyncio
import subprocess
import io
import uuid
from datetime import datetime, timedelta
from typing import Dict, Any, Optional, List
from pathlib import Path
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
import motor.motor_asyncio
from bson import ObjectId

from ..models.document_generation_models import (
    QuoteSettings,
    QuoteGenerationRequest,
    QuoteGenerationResponse,
    QuoteRecord,
    SaveQuoteSettingsRequest,
    GetUserQuoteDataResponse,
    DocumentTemplate,
    CompanyInfo,
    CustomerInfo,
    PaymentTerms,
)
from ..models.settings_models import UserQuoteSettings
from ..database.quote_db import get_database
from ..utils.file_utils import ensure_directory_exists
from ..utils.text_utils import number_to_words_vietnamese
from ..config.r2_storage import AIVungtauR2StorageConfig
from ..utils.logger import setup_logger

logger = setup_logger()


class QuoteGenerationService:
    """Service chính để xử lý tạo quote với AI"""

    def __init__(self):
        self.db = None
        self.gemini_client = None
        self.r2_storage = AIVungtauR2StorageConfig()
        self._initialize_gemini()

    async def initialize(self):
        """Initialize database connection"""
        if self.db is None:
            from ..database.quote_db import QuoteDatabase

            db_manager = QuoteDatabase()
            await db_manager.initialize()
            self.db = db_manager.db

    def _initialize_gemini(self):
        """Initialize Gemini Pro 2.5 client"""
        try:
            api_key = os.getenv("GEMINI_API_KEY")
            if not api_key:
                raise ValueError(
                    "GEMINI_API_KEY không được tìm thấy trong environment variables"
                )

            genai.configure(api_key=api_key)
            self.gemini_client = genai.GenerativeModel("gemini-2.5-pro")
            print("✅ Gemini Pro 2.5 client initialized successfully")

        except Exception as e:
            print(f"❌ Error initializing Gemini client: {str(e)}")
            self.gemini_client = None

    def _extract_response_text(self, response) -> str:
        """Safely extract text from Gemini response"""
        try:
            # Try simple text accessor first
            if hasattr(response, "text") and response.text:
                return response.text
        except Exception:
            pass

        try:
            # Use parts accessor for complex responses
            if response.candidates and len(response.candidates) > 0:
                candidate = response.candidates[0]
                if candidate.content and candidate.content.parts:
                    text_parts = []
                    for part in candidate.content.parts:
                        if hasattr(part, "text") and part.text:
                            text_parts.append(part.text)
                    if text_parts:
                        return "".join(text_parts)
        except Exception:
            pass

        # Fallback: try to get any text content
        try:
            if hasattr(response, "candidates") and response.candidates:
                return str(response.candidates[0])
        except Exception:
            pass

        raise Exception("Không thể trích xuất text từ Gemini response")

    async def save_quote_settings(
        self, user_id: str, firebase_uid: str, request: SaveQuoteSettingsRequest
    ) -> str:
        """Lưu settings ban đầu của user"""
        try:
            # Ensure database is initialized
            if self.db is None:
                await self.initialize()

            # Tạo quote settings
            settings = QuoteSettings(
                user_id=user_id,
                firebase_uid=firebase_uid,
                company_info=request.company_info,
                customer_info=request.customer_info,
                payment_terms=request.payment_terms,
                template_id=None,  # Sẽ được set khi user chọn template
                template_content=None,  # Không cần thiết vì sử dụng file DOCX
            )

            # Load template content nếu có template_id
            # if request.template_id:
            #     template = await self.db.document_templates.find_one(
            #         {"_id": ObjectId(request.template_id)}
            #     )
            #     if template:
            #         settings.template_content = template.get("template_content", "")

            # Lưu vào database
            result = await self.db.quote_settings.insert_one(
                settings.dict(by_alias=True)
            )

            return str(result.inserted_id)

        except Exception as e:
            raise Exception(f"Lỗi khi lưu quote settings: {str(e)}")

    async def get_user_quote_data(
        self, user_id: str, firebase_uid: str
    ) -> GetUserQuoteDataResponse:
        """Lấy dữ liệu quote gần nhất của user để frontend sử dụng"""
        try:
            # Lấy settings gần nhất
            latest_settings = await self.db.quote_settings.find_one(
                {"firebase_uid": firebase_uid, "is_active": True},
                sort=[("created_at", -1)],
            )

            settings_obj = None
            if latest_settings:
                settings_obj = QuoteSettings(**latest_settings)

            # Lấy quotes gần đây
            recent_quotes_cursor = self.db.quote_records.find(
                {"firebase_uid": firebase_uid, "status": "active"},
                sort=[("created_at", -1)],
                limit=10,
            )
            recent_quotes = [QuoteRecord(**doc) async for doc in recent_quotes_cursor]

            # Lấy available templates
            templates_cursor = self.db.user_upload_files.find(
                {"type": "quote", "is_active": True}, sort=[("created_at", -1)]
            )
            templates = [DocumentTemplate(**doc) async for doc in templates_cursor]

            return GetUserQuoteDataResponse(
                settings=settings_obj,
                recent_quotes=recent_quotes,
                available_templates=templates,
            )

        except Exception as e:
            raise Exception(f"Lỗi khi lấy dữ liệu user: {str(e)}")

    async def get_quote_history(
        self, user_id: str, firebase_uid: str
    ) -> List[Dict[str, Any]]:
        """Lấy lịch sử các quotes đã tạo của user"""
        try:
            # Lấy quote history từ database
            quotes_cursor = self.db.quote_records.find(
                {"firebase_uid": firebase_uid, "status": "active"},
                sort=[("created_at", -1)],
                limit=50,  # Giới hạn 50 quotes gần nhất
            )

            quotes = []
            async for doc in quotes_cursor:
                quote_data = {
                    "id": str(doc["_id"]),
                    "user_id": doc.get("user_id"),
                    "file_path": doc.get("file_path"),
                    "file_url": f"http://localhost:8000{doc.get('file_path', '')}",
                    "created_at": (
                        doc.get("created_at").isoformat()
                        if doc.get("created_at")
                        else None
                    ),
                    "total_amount": doc.get("total_amount", 0.0),
                    "currency": doc.get("currency", "VND"),
                    "customer_name": doc.get("customer_name", ""),
                }
                quotes.append(quote_data)

            return quotes

        except Exception as e:
            raise Exception(f"Lỗi khi lấy lịch sử quotes: {str(e)}")

    async def get_available_templates(self) -> List[Dict[str, Any]]:
        """Lấy danh sách templates có sẵn từ database"""
        try:
            # Ensure database is initialized
            if self.db is None:
                await self.initialize()

            templates = []
            cursor = self.db.user_upload_files.find({"is_active": True})

            async for doc in cursor:
                template_data = {
                    "template_id": doc.get("template_id"),
                    "name": doc.get("name"),
                    "description": doc.get("description", ""),
                    "category": doc.get("category", "general"),
                    "created_at": doc.get("created_at"),
                }
                templates.append(template_data)

            return templates

        except Exception as e:
            raise Exception(f"Lỗi khi lấy templates: {str(e)}")

    async def generate_quote_with_r2(
        self, user_id: str, firebase_uid: str, request: QuoteGenerationRequest
    ) -> Dict[str, Any]:
        """
        Wrapper method for generate_quote với R2 storage
        Returns simplified dict response for testing
        """
        try:
            response = await self.generate_quote(user_id, firebase_uid, request)

            # Convert to dict format
            return {
                "quote_id": response.quote_id,
                "status": response.status,
                "message": response.message,
                "file_key": response.file_key,
                "download_url": response.download_url,
                "file_size_bytes": response.file_size_bytes,
                "url_expires_at": (
                    response.url_expires_at.isoformat()
                    if response.url_expires_at
                    else None
                ),
                "generation_time_seconds": response.generation_time_seconds,
            }

        except Exception as e:
            raise Exception(f"Lỗi khi generate quote với R2: {str(e)}")

    async def generate_quote(
        self, user_id: str, firebase_uid: str, request: QuoteGenerationRequest
    ) -> QuoteGenerationResponse:
        """Tạo quote mới hoặc chỉnh sửa quote hiện có với R2 storage"""
        try:
            start_time = datetime.now()

            # Lấy settings
            settings = await self._get_quote_settings(request.settings_id)
            if not settings:
                raise Exception(
                    f"Không tìm thấy quote settings với ID: {request.settings_id}"
                )

            # Xây dựng prompt cho Gemini
            if request.generation_type == "new":
                # Set template_id vào settings trước khi call
                settings.template_id = request.template_id
                ai_content = await self._generate_new_quote(
                    settings, request.user_query, request.user_notes
                )
            else:
                ai_content = await self._edit_existing_quote(
                    settings, request.user_query, request.current_file_path
                )

            # Tạo file docx in-memory và upload lên R2
            r2_result = await self._create_and_upload_docx_file(
                ai_content, settings, user_id, request.user_notes
            )

            # Generate pre-signed download URL
            download_url = self.r2_storage.generate_presigned_download_url(
                r2_result["file_key"],
                expiry_minutes=30,
                filename=r2_result.get("file_name", "quote.docx"),
            )

            # Lưu quote record với R2 info
            quote_record = await self._save_quote_record_r2(
                user_id,
                firebase_uid,
                request,
                ai_content,
                r2_result,
                download_url,
                settings,
                start_time,
            )

            processing_time = (datetime.now() - start_time).total_seconds()

            return QuoteGenerationResponse(
                quote_id=str(quote_record.id),
                file_path="",  # Không còn lưu local file
                download_url=download_url,
                ai_generated_content=ai_content,
                processing_time=processing_time,
                status="success",
                message="Quote đã được tạo thành công",
            )

        except Exception as e:
            logger.error(f"❌ Error generating quote: {str(e)}")
            return QuoteGenerationResponse(
                quote_id="",
                file_path="",
                download_url="",
                ai_generated_content={},
                processing_time=0,
                status="error",
                message=f"Lỗi khi tạo quote: {str(e)}",
            )

    async def _get_quote_settings(self, settings_id: str) -> Optional[QuoteSettings]:
        """Lấy quote settings từ database"""
        try:
            # Tìm trong collection user_quote_settings trước (format mới)
            settings_doc = await self.db.user_quote_settings.find_one(
                {"_id": ObjectId(settings_id)}
            )
            if settings_doc:
                # Convert ObjectId to string before creating UserQuoteSettings
                settings_doc["_id"] = str(settings_doc["_id"])
                user_settings = UserQuoteSettings(**settings_doc)
                # Chuyển đổi sang format QuoteSettings
                return self._convert_user_settings_to_quote_settings(user_settings)

            # Nếu không tìm thấy, thử tìm trong collection quote_settings (format cũ)
            settings_doc = await self.db.quote_settings.find_one(
                {"_id": ObjectId(settings_id)}
            )
            if settings_doc:
                return QuoteSettings(**settings_doc)

            return None
        except Exception as e:
            return None

    def _convert_user_settings_to_quote_settings(
        self, user_settings: UserQuoteSettings
    ) -> QuoteSettings:
        """Chuyển đổi UserQuoteSettings sang QuoteSettings"""
        # Tạo customer_info và payment_terms mặc định
        default_customer_info = CustomerInfo(
            name="Khách hàng",
            address="Địa chỉ khách hàng",
            contact_person="Người liên hệ",
            phone="",
            email="",
        )

        default_payment_terms = PaymentTerms(
            payment_method="Chuyển khoản",
            payment_schedule="Thanh toán 100% trước khi giao hàng",
            currency="VND",
            advance_payment_percent=100.0,
            final_payment_terms=user_settings.quote_notes.payment_terms,
        )

        return QuoteSettings(
            user_id=user_settings.user_id,
            firebase_uid=user_settings.user_id,  # Giả sử user_id là firebase_uid
            company_info=user_settings.company_info.dict(),  # Convert to dict
            customer_info=default_customer_info.dict(),  # Convert to dict
            payment_terms=default_payment_terms.dict(),  # Convert to dict
            template_id=None,  # Sẽ được set từ request
            created_at=user_settings.created_at,
            updated_at=user_settings.updated_at,
        )

    async def _generate_new_quote(
        self, settings: QuoteSettings, user_query: str, user_notes: Optional[str] = None
    ) -> Dict[str, Any]:
        """Tạo quote mới với Gemini - gửi template DOCX + user info"""
        try:
            # 1. Load template DOCX từ database/filesystem
            template_file_path = await self._get_template_file(settings.template_id)

            # 2. Gửi template DOCX + user info cho Gemini
            prompt = self._build_quote_prompt_with_template(
                settings, user_query, template_file_path, user_notes
            )

            # 3. Gọi Gemini API - Gemini sẽ phân tích template và trả về structured text
            response = await self._call_gemini_api_with_template(
                prompt, template_file_path
            )

            # 4. Parse JSON response từ Gemini
            ai_content = self._parse_gemini_response(response)

            return ai_content

        except Exception as e:
            raise Exception(f"Lỗi khi tạo quote mới với AI: {str(e)}")

    async def _edit_existing_quote(
        self, settings: QuoteSettings, user_query: str, current_file_path: str
    ) -> Dict[str, Any]:
        """Chỉnh sửa quote hiện có với Gemini"""
        try:
            # Đọc nội dung file docx hiện tại
            current_content = await self._extract_docx_content(current_file_path)

            # Xây dựng prompt với file hiện tại
            prompt = self._build_edit_prompt(settings, user_query, current_content)

            # Gọi Gemini API
            response = await self._call_gemini_api(prompt)

            # Parse JSON response
            ai_content = self._parse_gemini_response(response)

            return ai_content

        except Exception as e:
            raise Exception(f"Lỗi khi chỉnh sửa quote với AI: {str(e)}")

    def _parse_gemini_response(self, response: str) -> Dict[str, Any]:
        """
        Parse Gemini response - handle both JSON and markdown formats
        """
        import re

        try:
            # Try direct JSON parsing first
            return json.loads(response)
        except json.JSONDecodeError:
            pass

        try:
            # Try to extract JSON from markdown code block
            # Pattern: ```json\n{...}\n```
            json_pattern = r"```json\s*\n(.*?)\n```"
            match = re.search(json_pattern, response, re.DOTALL)

            if match:
                json_content = match.group(1).strip()
                return json.loads(json_content)
        except json.JSONDecodeError:
            pass

        try:
            # Try to extract JSON from any code block
            # Pattern: ```\n{...}\n```
            code_pattern = r"```\s*\n(.*?)\n```"
            match = re.search(code_pattern, response, re.DOTALL)

            if match:
                content = match.group(1).strip()
                if content.startswith("{") and content.endswith("}"):
                    return json.loads(content)
        except json.JSONDecodeError:
            pass

        # If all parsing fails, create a basic structure
        logger.warning(f"⚠️ Could not parse Gemini response as JSON, using raw text")
        logger.warning(f"Raw response: {response[:500]}...")

        return {
            "quote_content": response,
            "summary": "Báo giá được tạo bởi AI",
            "metadata": {
                "parsing_method": "raw_text",
                "original_length": len(response),
            },
        }

    def _build_quote_prompt(
        self,
        settings: QuoteSettings,
        user_query: str,
        is_new: bool = True,
        user_notes: Optional[str] = None,
    ) -> str:
        """Xây dựng prompt chi tiết cho Gemini"""

        system_prompt = """
Bạn là chuyên gia tạo báo giá chuyên nghiệp cho doanh nghiệp Việt Nam.
Nhiệm vụ: Tạo ra báo giá chi tiết dựa trên thông tin được cung cấp và yêu cầu của người dùng.

QUAN TRỌNG: Trả về CHÍNH XÁC định dạng JSON sau:
{
    "header": {
        "title": "BÁO GIÁ SỐ ...",
        "company_name": "...",
        "date": "...",
        "quote_number": "..."
    },
    "company_section": {
        "supplier": {...},
        "customer": {...}
    },
    "products": [
        {
            "stt": 1,
            "name": "...",
            "description": "...",
            "quantity": 1,
            "unit": "...",
            "unit_price": 1000000,
            "total_price": 1000000,
            "specifications": "...",
            "warranty": "...",
            "delivery_time": "..."
        }
    ],
    "financial_summary": {
        "subtotal": 1000000,
        "vat_rate": 10,
        "vat_amount": 100000,
        "total_amount": 1100000,
        "total_in_words": "Một triệu một trăm nghìn đồng"
    },
    "payment_terms": {
        "method": "...",
        "schedule": "...",
        "advance_percent": 30,
        "final_terms": "..."
    },
    "additional_terms": [
        "...",
        "..."
    ],
    "validity": "30 ngày",
    "notes": "..."
}
"""

        # Thông tin input
        input_data = f"""
THÔNG TIN CÔNG TY:
- Tên: {settings.company_info.name or '[Tên công ty]'}
- Địa chỉ: {settings.company_info.address or '[Địa chỉ]'}
- MST: {settings.company_info.tax_code or '[MST]'}
- Người đại diện: {settings.company_info.representative or '[Người đại diện]'}
- Điện thoại: {settings.company_info.phone or '[SĐT]'}
- Email: {settings.company_info.email or '[Email]'}

THÔNG TIN KHÁCH HÀNG:
- Tên: {settings.customer_info.name or '[Tên khách hàng]'}
- Địa chỉ: {settings.customer_info.address or '[Địa chỉ KH]'}
- Người liên hệ: {settings.customer_info.contact_person or '[Người liên hệ]'}
- Điện thoại: {settings.customer_info.phone or '[SĐT KH]'}

ĐIỀU KHOẢN THANH TOÁN:
- Phương thức: {settings.payment_terms.payment_method or 'Chuyển khoản'}
- Lịch thanh toán: {settings.payment_terms.payment_schedule or '30% trước - 70% sau'}
- Tiền tệ: {settings.payment_terms.currency or 'VND'}

TEMPLATE CONTENT:
{settings.template_content or 'Sử dụng template báo giá chuẩn'}
"""

        user_request = f"""
YÊU CẦU CỦA NGƯỜI DÙNG:
{user_query}

{f'''
GHI CHÚ TỪ NGƯỜI DÙNG:
{user_notes}
''' if user_notes else ''}

Hãy tạo báo giá chuyên nghiệp theo yêu cầu, đảm bảo:
1. Sử dụng thông tin đã cung cấp
2. Tạo sản phẩm/dịch vụ phù hợp với yêu cầu
3. Tính toán chính xác giá tiền và VAT
4. Sử dụng ngôn ngữ tiếng Việt chuyên nghiệp hoặc tiếng Anh nếu người dùng yêu cầu
5. Trả về đúng định dạng JSON như yêu cầu
"""

        return f"{system_prompt}\n\n{input_data}\n\n{user_request}"

    def _build_edit_prompt(
        self, settings: QuoteSettings, user_query: str, current_content: str
    ) -> str:
        """Xây dựng prompt cho chỉnh sửa quote"""

        system_prompt = """
Bạn là chuyên gia chỉnh sửa báo giá chuyên nghiệp.
Nhiệm vụ: Chỉnh sửa báo giá hiện có theo yêu cầu của người dùng.

QUAN TRỌNG: Trả về CHÍNH XÁC định dạng JSON như template đã cho.
"""

        prompt = f"""
{system_prompt}

NỘI DUNG BÁO GIÁ HIỆN TẠI:
{current_content}

THÔNG TIN SETTINGS HIỆN TẠI:
- Công ty: {settings.company_info.name}
- Khách hàng: {settings.customer_info.name}

YÊU CẦU CHỈNH SỬA:
{user_query}

Hãy chỉnh sửa báo giá theo yêu cầu và trả về định dạng JSON hoàn chỉnh.
"""

        return prompt

    async def _call_gemini_api(self, prompt: str) -> str:
        """Gọi Gemini API"""
        try:
            if not self.gemini_client:
                raise Exception("Gemini client chưa được khởi tạo")

            response = await asyncio.to_thread(
                self.gemini_client.generate_content,
                prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,
                    max_output_tokens=4000,
                    response_mime_type="application/json",
                ),
            )

            return self._extract_response_text(response)

        except Exception as e:
            raise Exception(f"Lỗi khi gọi Gemini API: {str(e)}")

    async def _extract_docx_content(self, file_path: str) -> str:
        """Trích xuất nội dung từ file docx"""
        try:
            if not os.path.exists(file_path):
                raise Exception(f"File không tồn tại: {file_path}")

            doc = Document(file_path)
            content = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content.append(paragraph.text.strip())

            return "\n".join(content)

        except Exception as e:
            raise Exception(f"Lỗi khi đọc file docx: {str(e)}")

    async def _create_and_upload_docx_file(
        self,
        ai_content: Dict[str, Any],
        settings: QuoteSettings,
        user_id: str,
        user_notes: Optional[str] = None,
    ) -> Dict[str, Any]:
        """Tạo file docx in-memory và upload lên R2"""
        try:
            # Debug: Check ai_content type
            logger.info(f"🔍 DEBUG ai_content type: {type(ai_content)}")
            logger.info(f"🔍 DEBUG ai_content value: {ai_content}")

            # Validate ai_content is dict
            if not isinstance(ai_content, dict):
                logger.error(
                    f"❌ ai_content is not dict, got {type(ai_content)}: {ai_content}"
                )
                raise Exception(f"AI content must be dict, got {type(ai_content)}")

            # Load template file
            template_path = "templates/documents/quote_template_default.docx"
            doc = Document(template_path)

            # Lấy dữ liệu từ AI content và settings
            header = ai_content.get("header", {})
            company_section = ai_content.get("company_section", {})
            products = ai_content.get("products", [])
            summary = ai_content.get("summary", {})

            # Debug customer name extraction
            ai_customer_name = company_section.get("customer", {}).get("name")
            settings_customer_name = settings.customer_info.name
            logger.info(f"🔍 Customer name from AI: {ai_customer_name}")
            logger.info(f"🔍 Customer name from settings: {settings_customer_name}")

            # Tạo mapping để replace placeholders
            now = datetime.now()
            replacements = {
                # Company info
                "[COMPANY_NAME]": settings.company_info.name or "Tên công ty",
                "[COMPANY_ADDRESS]": settings.company_info.address or "Địa chỉ công ty",
                "[COMPANY_PHONE]": settings.company_info.phone or "0901234567",
                "[COMPANY_EMAIL]": settings.company_info.email or "email@company.com",
                "[COMPANY_WEBSITE]": settings.company_info.website
                or "https://company.com",
                # Customer info - Priority: AI response > settings > default
                "[CUSTOMER_NAME]": company_section.get("customer", {}).get("name")
                or settings.customer_info.name
                or "Tên khách hàng",
                # Date
                "[CITY]": "TP.HCM",
                "[DAY]": str(now.day),
                "[MONTH]": str(now.month),
                "[YEAR]": str(now.year),
                # Representative
                "[REPRESENTATIVE_NAME]": settings.company_info.representative
                or "Người đại diện",
                "[REPRESENTATIVE_POSITION]": settings.company_info.position
                or "Giám đốc",
                # Payment terms
                "[PAYMENT_TERMS]": settings.payment_terms.payment_schedule
                or "Báo giá có hiệu lực trong 30 ngày.",
                # Notes - Priority: AI response > user input > default settings
                "[NOTES]": (
                    ai_content.get("notes")
                    or user_notes
                    or getattr(settings, "quote_notes", {}).get("default_notes")
                    or "Cảm ơn quý khách đã quan tâm đến sản phẩm của chúng tôi."
                ),
                # Financial (sẽ được tính từ products)
                "[VAT_RATE]": str(summary.get("vat_rate", 10)),
                "[SUBTOTAL]": f"{summary.get('subtotal', 0):,.0f}",
                "[VAT_AMOUNT]": f"{summary.get('vat_amount', 0):,.0f}",
                "[TOTAL_AMOUNT]": f"{summary.get('total_amount', 0):,.0f}",
                "[TOTAL_AMOUNT_WORDS]": summary.get(
                    "total_amount_words",
                    number_to_words_vietnamese(summary.get("total_amount", 0)),
                ),
            }

            # Replace placeholders trong tất cả paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

            # Replace placeholders trong tables và điền product data
            self._update_table_with_products(doc, products, summary)

            # Tạo file buffer in-memory
            file_buffer = io.BytesIO()
            doc.save(file_buffer)
            file_buffer.seek(0)  # Reset pointer to beginning

            # Generate file key và file name
            file_key = self.r2_storage.generate_file_key(user_id, "quote")
            company_name = settings.company_info.name or "Company"

            # Create ASCII-only company name for metadata
            import unicodedata

            ascii_company_name = (
                unicodedata.normalize("NFKD", company_name)
                .encode("ascii", "ignore")
                .decode("ascii")
            )
            safe_company_name = "".join(
                c for c in ascii_company_name if c.isalnum() or c in (" ", "-", "_")
            ).rstrip()
            safe_company_name = safe_company_name.replace(" ", "_") or "Company"

            # Debug log
            logger.info(
                f"🔧 Company name conversion: '{company_name}' -> '{safe_company_name}'"
            )

            file_name = f"bao_gia_{safe_company_name}_{now.strftime('%Y%m%d')}.docx"

            # Upload to R2
            upload_result = await self.r2_storage.upload_file_from_buffer(
                file_buffer.getvalue(),
                file_key,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                metadata={
                    "user_id": user_id,
                    "company_name": safe_company_name,  # Use ASCII-safe version
                    "quote_type": "ai_generated",
                },
            )

            logger.info(f"✅ Quote file uploaded to R2: {file_key}")

            return {
                **upload_result,
                "file_name": file_name,
                "original_file_key": file_key,
            }

        except Exception as e:
            logger.error(f"❌ Error creating and uploading DOCX file: {str(e)}")
            raise Exception(f"Lỗi khi tạo và upload file docx: {str(e)}")

    async def _save_quote_record_r2(
        self,
        user_id: str,
        firebase_uid: str,
        request: QuoteGenerationRequest,
        ai_content: Dict[str, Any],
        r2_result: Dict[str, Any],
        download_url: str,
        settings: QuoteSettings,
        start_time: datetime,
    ) -> QuoteRecord:
        """Lưu quote record với thông tin R2"""
        try:
            # Tạo quote ID
            quote_id = f"quote_{int(datetime.now().timestamp())}"

            # Tính thời gian xử lý
            processing_time = (datetime.now() - start_time).total_seconds()

            # Tạo quote record
            quote_record = QuoteRecord(
                quote_id=quote_id,
                user_id=user_id,
                firebase_uid=firebase_uid,
                settings_id=request.settings_id,
                user_query=request.user_query,
                generation_type=request.generation_type,
                template_id=getattr(request, "template_id", None),
                ai_generated_content=ai_content,
                r2_file_key=r2_result["file_key"],
                file_name=r2_result["file_name"],
                download_url=download_url,
                url_expires_at=datetime.now() + timedelta(minutes=30),
                file_size_bytes=r2_result.get("size_bytes", 0),
                processing_time=processing_time,
                created_at=datetime.now(),
                status="completed",
            )

            # Lưu vào database
            result = await self.db.quote_records.insert_one(
                quote_record.dict(by_alias=True)
            )

            quote_record.id = str(result.inserted_id)
            logger.info(f"✅ Quote record saved: {quote_id}")

            return quote_record

        except Exception as e:
            logger.error(f"❌ Error saving quote record: {str(e)}")
            raise Exception(f"Lỗi khi lưu quote record: {str(e)}")

    async def _create_docx_file(
        self,
        ai_content: Dict[str, Any],
        settings: QuoteSettings,
        user_id: str,
        user_notes: Optional[str] = None,
    ) -> str:
        """Tạo file docx từ AI content sử dụng template"""
        try:
            # Load template file
            template_path = "templates/documents/quote_template_default.docx"
            doc = Document(template_path)

            # Sử dụng number_to_words_vietnamese từ utils (đã import ở đầu file)

            # Lấy dữ liệu từ AI content và settings
            header = ai_content.get("header", {})
            company_section = ai_content.get("company_section", {})
            products = ai_content.get("products", [])
            summary = ai_content.get("summary", {})

            # Debug customer name extraction
            ai_customer_name = company_section.get("customer", {}).get("name")
            settings_customer_name = settings.customer_info.name
            logger.info(
                f"🔍 [_create_docx_file] Customer name from AI: {ai_customer_name}"
            )
            logger.info(
                f"🔍 [_create_docx_file] Customer name from settings: {settings_customer_name}"
            )

            # Tạo mapping để replace placeholders
            now = datetime.now()
            replacements = {
                # Company info
                "[COMPANY_NAME]": settings.company_info.name or "Tên công ty",
                "[COMPANY_ADDRESS]": settings.company_info.address or "Địa chỉ công ty",
                "[COMPANY_PHONE]": settings.company_info.phone or "0901234567",
                "[COMPANY_EMAIL]": settings.company_info.email or "email@company.com",
                "[COMPANY_WEBSITE]": settings.company_info.website
                or "https://company.com",
                # Customer info - Priority: AI response > settings > default
                "[CUSTOMER_NAME]": company_section.get("customer", {}).get("name")
                or settings.customer_info.name
                or "Tên khách hàng",
                # Date
                "[CITY]": "TP.HCM",
                "[DAY]": str(now.day),
                "[MONTH]": str(now.month),
                "[YEAR]": str(now.year),
                # Representative
                "[REPRESENTATIVE_NAME]": settings.company_info.representative
                or "Người đại diện",
                "[REPRESENTATIVE_POSITION]": settings.company_info.position
                or "Giám đốc",
                # Payment terms
                "[PAYMENT_TERMS]": settings.payment_terms.payment_schedule
                or "Báo giá có hiệu lực trong 30 ngày.",
                # Notes - Priority: AI response > user input > default settings
                "[NOTES]": (
                    ai_content.get("notes")
                    or user_notes
                    or getattr(settings, "quote_notes", {}).get("default_notes")
                    or "Cảm ơn quý khách đã quan tâm đến sản phẩm của chúng tôi."
                ),
                # Financial (sẽ được tính từ products)
                "[VAT_RATE]": str(summary.get("vat_rate", 10)),
                "[SUBTOTAL]": f"{summary.get('subtotal', 0):,.0f}",
                "[VAT_AMOUNT]": f"{summary.get('vat_amount', 0):,.0f}",
                "[TOTAL_AMOUNT]": f"{summary.get('total_amount', 0):,.0f}",
                "[TOTAL_AMOUNT_WORDS]": summary.get(
                    "total_amount_words",
                    number_to_words_vietnamese(summary.get("total_amount", 0)),
                ),
            }

            # Replace placeholders trong tất cả paragraphs
            for paragraph in doc.paragraphs:
                for placeholder, value in replacements.items():
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, str(value))

            # Replace placeholders trong tables và điền product data
            self._update_table_with_products(doc, products, summary)

            # Save file
            file_name = (
                f"quote_{user_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )
            output_dir = Path("generated_files/quotes")
            ensure_directory_exists(str(output_dir))
            file_path = output_dir / file_name

            doc.save(str(file_path))
            return str(file_path)

        except Exception as e:
            raise Exception(f"Lỗi khi tạo file docx: {str(e)}")

    def _add_company_section(self, doc: Document, company_section: Dict[str, Any]):
        """Thêm phần thông tin công ty"""
        doc.add_heading("THÔNG TIN BÁO GIÁ", level=1)

        supplier = company_section.get("supplier", {})
        customer = company_section.get("customer", {})

        # Supplier info
        doc.add_paragraph(f"Đơn vị cung cấp: {supplier.get('name', '')}")
        doc.add_paragraph(f"Địa chỉ: {supplier.get('address', '')}")
        doc.add_paragraph(f"MST: {supplier.get('tax_code', '')}")
        doc.add_paragraph(f"Điện thoại: {supplier.get('phone', '')}")

        doc.add_paragraph("")  # Empty line

        # Customer info
        doc.add_paragraph(f"Khách hàng: {customer.get('name', '')}")
        doc.add_paragraph(f"Địa chỉ: {customer.get('address', '')}")
        doc.add_paragraph(f"Người liên hệ: {customer.get('contact_person', '')}")

    def _add_products_table(self, doc: Document, products: List[Dict[str, Any]]):
        """Thêm bảng sản phẩm"""
        doc.add_heading("CHI TIẾT SẢN PHẨM/DỊCH VỤ", level=1)

        # Create table
        table = doc.add_table(rows=1, cols=7)
        table.style = "Table Grid"

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "STT"
        hdr_cells[1].text = "Tên sản phẩm/dịch vụ"
        hdr_cells[2].text = "Số lượng"
        hdr_cells[3].text = "Đơn vị"
        hdr_cells[4].text = "Đơn giá"
        hdr_cells[5].text = "Thành tiền"
        hdr_cells[6].text = "Ghi chú"

        # Add product rows
        for product in products:
            row_cells = table.add_row().cells
            row_cells[0].text = str(product.get("stt", ""))
            row_cells[1].text = product.get("name", "")
            row_cells[2].text = str(product.get("quantity", ""))
            row_cells[3].text = product.get("unit", "")
            row_cells[4].text = f"{product.get('unit_price', 0):,} VND"
            row_cells[5].text = f"{product.get('total_price', 0):,} VND"
            row_cells[6].text = product.get("specifications", "")

    def _add_financial_summary(self, doc: Document, financial: Dict[str, Any]):
        """Thêm tổng kết tài chính"""
        doc.add_heading("TỔNG KẾT", level=1)

        doc.add_paragraph(f"Tổng tiền hàng: {financial.get('subtotal', 0):,} VND")
        doc.add_paragraph(
            f"VAT ({financial.get('vat_rate', 10)}%): {financial.get('vat_amount', 0):,} VND"
        )
        doc.add_paragraph(f"TỔNG CỘNG: {financial.get('total_amount', 0):,} VND")
        doc.add_paragraph(f"Bằng chữ: {financial.get('total_in_words', '')}")

    def _add_payment_terms(self, doc: Document, payment_terms: Dict[str, Any]):
        """Thêm điều khoản thanh toán"""
        doc.add_heading("ĐIỀU KHOẢN THANH TOÁN", level=1)

        doc.add_paragraph(f"Phương thức: {payment_terms.get('method', '')}")
        doc.add_paragraph(f"Lịch thanh toán: {payment_terms.get('schedule', '')}")
        if payment_terms.get("advance_percent"):
            doc.add_paragraph(f"Tạm ứng: {payment_terms.get('advance_percent')}%")

    def _add_additional_terms(self, doc: Document, terms: List[str]):
        """Thêm điều khoản bổ sung"""
        doc.add_heading("ĐIỀU KHOẢN BỔ SUNG", level=1)

        for i, term in enumerate(terms, 1):
            doc.add_paragraph(f"{i}. {term}")

    async def _save_quote_record(
        self,
        user_id: str,
        firebase_uid: str,
        request: QuoteGenerationRequest,
        ai_content: Dict[str, Any],
        file_path: str,
        settings: QuoteSettings,
        start_time: datetime,
    ) -> QuoteRecord:
        """Lưu quote record vào database"""
        try:
            processing_time = (datetime.now() - start_time).total_seconds()

            # Tìm version number
            version = 1
            if request.generation_type == "edit":
                # Tìm version cao nhất của settings này
                last_quote = await self.db.quote_records.find_one(
                    {"settings_id": request.settings_id}, sort=[("version", -1)]
                )
                if last_quote:
                    version = last_quote.get("version", 0) + 1

            quote_record = QuoteRecord(
                user_id=user_id,
                firebase_uid=firebase_uid,
                settings_id=request.settings_id,
                user_query=request.user_query,
                ai_generated_content=ai_content,
                file_path=file_path,
                file_size=(
                    os.path.getsize(file_path) if os.path.exists(file_path) else 0
                ),
                version=version,
                ai_model=request.ai_model,
                processing_time=processing_time,
                expires_at=datetime.now()
                + timedelta(days=30),  # File expires in 30 days
            )

            result = await self.db.quote_records.insert_one(
                quote_record.dict(by_alias=True)
            )
            quote_record.id = result.inserted_id

            return quote_record

        except Exception as e:
            raise Exception(f"Lỗi khi lưu quote record: {str(e)}")

    async def get_quote_file(self, quote_id: str, user_id: str) -> Optional[str]:
        """Lấy đường dẫn file quote để download"""
        try:
            quote = await self.db.quote_records.find_one(
                {"_id": ObjectId(quote_id), "user_id": user_id, "status": "active"}
            )

            if quote and os.path.exists(quote["file_path"]):
                # Update download count
                await self.db.quote_records.update_one(
                    {"_id": ObjectId(quote_id)}, {"$inc": {"download_count": 1}}
                )
                return quote["file_path"]

            return None

        except Exception:
            return None

    async def _convert_docx_to_pdf(self, docx_path: str) -> str:
        """Convert DOCX file to PDF for Gemini processing using docx2pdf"""
        try:
            from docx2pdf import convert

            pdf_path = docx_path.replace(".docx", "_for_gemini.pdf")

            # Check if PDF already exists and is newer than DOCX
            if os.path.exists(pdf_path):
                docx_time = os.path.getmtime(docx_path)
                pdf_time = os.path.getmtime(pdf_path)
                if pdf_time > docx_time:
                    return pdf_path

            # Use docx2pdf to convert DOCX to PDF
            # Create temporary file path for conversion
            temp_pdf = docx_path.replace(".docx", ".pdf")

            # Convert using docx2pdf
            convert(docx_path, temp_pdf)

            # Check if conversion was successful
            if os.path.exists(temp_pdf):
                # Rename to our naming convention
                if os.path.exists(pdf_path):
                    os.remove(pdf_path)
                os.rename(temp_pdf, pdf_path)
                return pdf_path
            else:
                raise Exception("PDF file was not created by docx2pdf")

        except ImportError:
            raise Exception(
                "docx2pdf not installed. Please install: pip install docx2pdf"
            )
        except Exception as e:
            raise Exception(f"Lỗi khi convert DOCX sang PDF: {str(e)}")

    async def _get_template_file(self, template_id: str) -> str:
        """Lấy đường dẫn file template PDF từ database (ưu tiên PDF, fallback convert DOCX)"""
        try:
            if not template_id:
                # Sử dụng template mặc định
                template_id = "template_quote_001"

            template = await self.db.user_upload_files.find_one(
                {"_id": template_id, "is_active": True}
            )

            template_path = None
            if template is not None and template.get("file_path"):
                template_path = template["file_path"]
                if os.path.exists(template_path):
                    # Nếu là PDF, sử dụng trực tiếp
                    if template_path.endswith(".pdf"):
                        return template_path

                    # Nếu là DOCX, check xem có PDF tương ứng không
                    if template_path.endswith(".docx"):
                        pdf_equivalent = template_path.replace(".docx", "_pdf.pdf")
                        if os.path.exists(pdf_equivalent):
                            return pdf_equivalent
                        # Nếu không có PDF, convert DOCX sang PDF
                        return await self._convert_docx_to_pdf(template_path)

            # Fallback: tìm template PDF mặc định trước
            default_pdf_path = "templates/documents/quote_template_default_pdf.pdf"
            if os.path.exists(default_pdf_path):
                return default_pdf_path

            # Fallback cuối: convert template DOCX mặc định
            default_docx_path = "templates/documents/quote_template_default.docx"
            if os.path.exists(default_docx_path):
                return await self._convert_docx_to_pdf(default_docx_path)

            raise FileNotFoundError(
                f"Không tìm thấy template file cho ID: {template_id}"
            )

        except Exception as e:
            raise Exception(f"Lỗi khi load template file: {str(e)}")

    def _build_quote_prompt_with_template(
        self,
        settings: QuoteSettings,
        user_query: str,
        template_file_path: str,
        user_notes: Optional[str] = None,
    ) -> str:
        """Xây dựng prompt cho Gemini với template PDF (có thể được convert từ DOCX)"""

        system_prompt = f"""
Bạn là chuyên gia tạo báo giá chuyên nghiệp cho doanh nghiệp Việt Nam.

NHIỆM VỤ:
1. Phân tích file template PDF được cung cấp (có thể được convert từ DOCX hoặc tạo trực tiếp)
2. Dựa trên thông tin từ user và yêu cầu, tạo nội dung báo giá hoàn chỉnh
3. Trả về JSON với đầy đủ thông tin để generate file DOCX

TEMPLATE FILE: {template_file_path}

THÔNG TIN CÔNG TY BÁN HÀNG:
- Tên: {settings.company_info.name or '[Tên công ty]'}
- Địa chỉ: {settings.company_info.address or '[Địa chỉ]'}
- MST: {settings.company_info.tax_code or '[MST]'}
- Đại diện: {settings.company_info.representative or '[Người đại diện]'}
- Điện thoại: {settings.company_info.phone or '[SĐT]'}
- Email: {settings.company_info.email or '[Email]'}

THÔNG TIN KHÁCH HÀNG:
- Tên: {settings.customer_info.name or '[Tên KH]'}
- Địa chỉ: {settings.customer_info.address or '[Địa chỉ KH]'}
- Người liên hệ: {settings.customer_info.contact_person or '[Người liên hệ]'}
- Điện thoại: {settings.customer_info.phone or '[SĐT KH]'}

ĐIỀU KHOẢN THANH TOÁN:
- Phương thức: {settings.payment_terms.payment_method or 'Chuyển khoản'}
- Lịch thanh toán: {settings.payment_terms.payment_schedule or '30% trước - 70% sau'}
- Tiền tệ: {settings.payment_terms.currency or 'VND'}

YÊU CẦU CỦA NGƯỜI DÙNG:
{user_query}

QUAN TRỌNG: Trả về CHÍNH XÁC định dạng JSON sau:
{{
    "header": {{
        "title": "BÁOGIÁ SỐ ...",
        "company_name": "...",
        "date": "...",
        "quote_number": "..."
    }},
    "company_section": {{
        "supplier": {{...}},
        "customer": {{...}}
    }},
    "products": [
        {{
            "stt": 1,
            "name": "...",
            "description": "...",
            "quantity": 1,
            "unit": "...",
            "unit_price": 1000000,
            "total_price": 1000000
        }}
    ],
    "summary": {{
        "subtotal": 0,
        "vat_rate": 10,
        "vat_amount": 0,
        "total_amount": 0,
        "total_amount_words": "..."
    }},
    "terms": [
        "Báo giá có hiệu lực trong 30 ngày",
        "Giá đã bao gồm VAT 10%"
    ],
    "footer": {{
        "signature_supplier": "...",
        "signature_customer": "..."
    }}
}}

CHỨC NĂNG ĐẶC BIỆT:
- Khi tính tổng tiền, hãy chuyển đổi số thành chữ tiếng Việt chính xác
- Ví dụ: 55.000.000 VND → "Năm mươi lăm triệu đồng"
"""

        # Thông tin input
        input_data = f"""
THÔNG TIN CÔNG TY:
- Tên: {settings.company_info.name or '[Tên công ty]'}
- Địa chỉ: {settings.company_info.address or '[Địa chỉ]'}
- MST: {settings.company_info.tax_code or '[MST]'}
- Người đại diện: {settings.company_info.representative or '[Người đại diện]'}
- Điện thoại: {settings.company_info.phone or '[SĐT]'}
- Email: {settings.company_info.email or '[Email]'}

THÔNG TIN KHÁCH HÀNG:
- Tên: {settings.customer_info.name or '[Tên khách hàng]'}
- Địa chỉ: {settings.customer_info.address or '[Địa chỉ KH]'}
- Người liên hệ: {settings.customer_info.contact_person or '[Người liên hệ]'}
- Điện thoại: {settings.customer_info.phone or '[SĐT KH]'}

ĐIỀU KHOẢN THANH TOÁN:
- Phương thức: {settings.payment_terms.payment_method or 'Chuyển khoản'}
- Lịch thanh toán: {settings.payment_terms.payment_schedule or '30% trước - 70% sau'}
- Tiền tệ: {settings.payment_terms.currency or 'VND'}

TEMPLATE CONTENT:
{settings.template_content or 'Sử dụng template báo giá chuẩn'}
"""

        user_request = f"""
YÊU CẦU CỦA NGƯỜI DÙNG:
{user_query}

{f'''
GHI CHÚ TỪ NGƯỜI DÙNG:
{user_notes}
''' if user_notes else ''}

Hãy tạo báo giá chuyên nghiệp theo yêu cầu, đảm bảo:
1. Sử dụng thông tin đã cung cấp
2. Tạo sản phẩm/dịch vụ phù hợp với yêu cầu
3. Tính toán chính xác giá tiền và VAT
4. Sử dụng ngôn ngữ tiếng Việt chuyên nghiệp hoặc tiếng Anh nếu người dùng yêu cầu
5. Trả về đúng định dạng JSON như yêu cầu
"""

        return f"{system_prompt}\n\n{input_data}\n\n{user_request}"

    async def _call_gemini_api_with_template(
        self, prompt: str, template_file_path: str
    ) -> str:
        """Gọi Gemini API - đơn giản hóa không upload file"""
        try:
            if not self.gemini_client:
                self._initialize_gemini()

            # Kiểm tra file template tồn tại
            if not os.path.exists(template_file_path):
                raise Exception(f"Template file không tồn tại: {template_file_path}")

            # Đọc thông tin template để đưa vào prompt (thay vì upload file)
            template_info = f"""
TEMPLATE INFO: Sử dụng template DOCX báo giá chuẩn với:
- Header: Số báo giá, ngày tháng, thông tin công ty
- Thông tin khách hàng
- Bảng sản phẩm/dịch vụ với các cột: STT, Tên, Số lượng, Đơn vị, Đơn giá, Thành tiền
- Tổng kết: Cộng, VAT (10%), Tổng cộng
- Điều khoản thanh toán
"""

            # Tạo prompt đầy đủ
            full_prompt = f"""
{template_info}

{prompt}

QUAN TRỌNG: Trả về ĐÚNG định dạng JSON, không có text khác. Đảm bảo:
1. Tất cả số tiền phải chính xác (đơn giá × số lượng = thành tiền)
2. VAT = 10% của subtotal
3. total_amount = subtotal + vat_amount
4. total_in_words phải chuyển đổi số tiền thành chữ tiếng Việt chính xác
5. Sử dụng tiếng Việt chuyên nghiệp
"""

            # Gọi Gemini - đơn giản chỉ với text
            response = await asyncio.to_thread(
                self.gemini_client.generate_content,
                full_prompt,
                generation_config=genai.types.GenerationConfig(
                    temperature=0.1,
                    max_output_tokens=4000,
                ),
            )

            response_text = self._extract_response_text(response)
            if response_text:
                return response_text.strip()
            else:
                raise Exception("Gemini không trả về response")

        except Exception as e:
            raise Exception(f"Lỗi khi gọi Gemini API với template: {str(e)}")

    def _update_table_with_products(
        self, doc: Document, products: List[Dict[str, Any]], summary: Dict[str, Any]
    ):
        """Cập nhật table với danh sách sản phẩm"""
        try:
            # Tìm table đầu tiên (table sản phẩm)
            if not doc.tables:
                raise Exception("Không tìm thấy table trong template")

            table = doc.tables[0]

            # Xóa row sản phẩm mẫu (row index 1)
            if len(table.rows) > 1:
                # Lưu lại các row summary (Cộng, VAT, Tổng) để add lại sau
                summary_rows_data = []
                for i in range(2, len(table.rows)):  # Từ row 2 trở đi là summary rows
                    row = table.rows[i]
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text)
                    summary_rows_data.append(row_data)

                # Xóa tất cả rows sau header (trừ header row 0)
                for i in range(len(table.rows) - 1, 0, -1):
                    table._element.remove(table.rows[i]._element)

            # Thêm các sản phẩm vào table
            for idx, product in enumerate(products, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = product.get("name", "")
                row_cells[2].text = product.get("description", "")
                row_cells[3].text = product.get("unit", "bộ")
                row_cells[4].text = str(product.get("quantity", 1))
                row_cells[5].text = f"{product.get('unit_price', 0):,.0f}"
                row_cells[6].text = f"{product.get('total_price', 0):,.0f}"

            # Thêm lại summary rows với merge cells
            # Row Cộng
            row_cells = table.add_row().cells
            merged_cell = row_cells[0]
            for i in range(1, 6):
                merged_cell.merge(row_cells[i])
            merged_cell.text = "Cộng:"
            row_cells[6].text = f"{summary.get('subtotal', 0):,.0f}"

            # Row VAT
            row_cells = table.add_row().cells
            merged_cell = row_cells[0]
            for i in range(1, 6):
                merged_cell.merge(row_cells[i])
            merged_cell.text = f"Thuế VAT ({summary.get('vat_rate', 10)}%):"
            row_cells[6].text = f"{summary.get('vat_amount', 0):,.0f}"

            # Row Tổng giá
            row_cells = table.add_row().cells
            merged_cell = row_cells[0]
            for i in range(1, 6):
                merged_cell.merge(row_cells[i])
            merged_cell.text = "Tổng giá:"
            row_cells[6].text = f"{summary.get('total_amount', 0):,.0f}"

        except Exception as e:
            raise Exception(f"Lỗi khi cập nhật table: {str(e)}")
