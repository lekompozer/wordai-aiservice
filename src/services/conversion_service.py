"""
HTML to DOCX Conversion Service
Service to convert HTML to DOCX with R2 upload support
"""

import io
import uuid
from typing import Optional, Dict, Any
from datetime import datetime

from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

from src.utils.logger import setup_logger
from src.config.r2_storage import AIVungtauR2StorageConfig

logger = setup_logger(__name__)


class ConversionService:
    """Service for HTML to DOCX conversion"""

    def __init__(self):
        self.r2_storage = AIVungtauR2StorageConfig()

    async def html_to_docx(
        self,
        html_content: str,
        file_name: Optional[str] = None,
        upload_to_r2: bool = True,
    ) -> Dict[str, Any]:
        """
        Convert HTML content to DOCX file

        Args:
            html_content: HTML content to convert
            file_name: Optional file name (without extension)
            upload_to_r2: Whether to upload to R2 storage

        Returns:
            Dict containing download URL and file info
        """
        try:
            logger.info(f"Starting HTML to DOCX conversion")

            # Create DOCX document
            doc = Document()

            # Parse HTML content
            soup = BeautifulSoup(html_content, "html.parser")

            # Convert HTML elements to DOCX
            await self._convert_html_elements(doc, soup)

            # Generate file name if not provided
            if not file_name:
                file_name = (
                    f"converted_document_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
                )

            # Ensure .docx extension
            if not file_name.endswith(".docx"):
                file_name += ".docx"

            # Save to bytes
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_content = docx_buffer.getvalue()

            logger.info(f"HTML converted to DOCX: {len(docx_content)} bytes")

            result = {
                "success": True,
                "file_name": file_name,
                "file_size": len(docx_content),
                "conversion_time": datetime.utcnow().isoformat(),
            }

            if upload_to_r2:
                # Upload to R2 storage
                file_key = (
                    f"conversions/{datetime.now().strftime('%Y/%m/%d')}/{file_name}"
                )
                upload_result = await self.r2_storage.upload_file_from_buffer(
                    docx_content,
                    file_key,
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )

                result.update(
                    {
                        "download_url": upload_result["public_url"],
                        "file_key": file_key,
                        "storage": "r2",
                    }
                )

                logger.info(f"File uploaded to R2: {upload_result['public_url']}")
            else:
                # Return file content directly
                result.update({"file_content": docx_content, "storage": "memory"})

            return result

        except Exception as e:
            logger.error(f"HTML to DOCX conversion failed: {e}")
            raise Exception(f"HTML to DOCX conversion failed: {str(e)}")

    async def _convert_html_elements(self, doc: Document, soup: BeautifulSoup):
        """Convert HTML elements to DOCX elements"""
        try:
            # Handle different HTML elements
            for element in soup.find_all(
                ["h1", "h2", "h3", "h4", "h5", "h6", "p", "div", "span", "table"]
            ):
                if element.name in ["h1", "h2", "h3", "h4", "h5", "h6"]:
                    # Headers
                    level = int(element.name[1])
                    paragraph = doc.add_heading(element.get_text(), level=level)

                elif element.name in ["p", "div", "span"]:
                    # Paragraphs and text containers
                    text = element.get_text()
                    if text.strip():
                        paragraph = doc.add_paragraph(text)

                        # Handle basic formatting
                        if element.find("strong") or element.find("b"):
                            # Bold text (simplified - would need more complex handling for partial bold)
                            run = (
                                paragraph.runs[0]
                                if paragraph.runs
                                else paragraph.add_run()
                            )
                            run.bold = True

                elif element.name == "table":
                    # Tables
                    await self._convert_table(doc, element)

        except Exception as e:
            logger.warning(f"Error converting HTML elements: {e}")
            # Add fallback - convert to plain text
            plain_text = soup.get_text()
            doc.add_paragraph(plain_text)

    async def _convert_table(self, doc: Document, table_element):
        """Convert HTML table to DOCX table"""
        try:
            rows = table_element.find_all("tr")
            if not rows:
                return

            # Get max columns
            max_cols = max(len(row.find_all(["td", "th"])) for row in rows)

            # Create DOCX table
            docx_table = doc.add_table(rows=len(rows), cols=max_cols)
            docx_table.style = "Table Grid"

            # Populate table
            for row_idx, row in enumerate(rows):
                cells = row.find_all(["td", "th"])
                for col_idx, cell in enumerate(cells):
                    if col_idx < max_cols:
                        docx_table.cell(row_idx, col_idx).text = cell.get_text().strip()

                        # Make header row bold
                        if cell.name == "th":
                            for paragraph in docx_table.cell(
                                row_idx, col_idx
                            ).paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True

        except Exception as e:
            logger.warning(f"Error converting table: {e}")

    async def docx_to_html(self, docx_content: bytes) -> str:
        """
        Convert DOCX to HTML (reverse conversion)

        Args:
            docx_content: DOCX file content as bytes

        Returns:
            HTML content as string
        """
        try:
            logger.info("Starting DOCX to HTML conversion")

            # Load DOCX
            doc = Document(io.BytesIO(docx_content))

            # Convert to HTML
            html_parts = ["<html><body>"]

            for paragraph in doc.paragraphs:
                if paragraph.style.name.startswith("Heading"):
                    level = (
                        paragraph.style.name[-1]
                        if paragraph.style.name[-1].isdigit()
                        else "1"
                    )
                    html_parts.append(f"<h{level}>{paragraph.text}</h{level}>")
                else:
                    html_parts.append(f"<p>{paragraph.text}</p>")

            # Handle tables
            for table in doc.tables:
                html_parts.append('<table border="1">')
                for row in table.rows:
                    html_parts.append("<tr>")
                    for cell in row.cells:
                        html_parts.append(f"<td>{cell.text}</td>")
                    html_parts.append("</tr>")
                html_parts.append("</table>")

            html_parts.append("</body></html>")
            html_content = "\n".join(html_parts)

            logger.info(f"DOCX converted to HTML: {len(html_content)} characters")
            return html_content

        except Exception as e:
            logger.error(f"DOCX to HTML conversion failed: {e}")
            raise Exception(f"DOCX to HTML conversion failed: {str(e)}")

    async def update_template_content(
        self, template_id: str, html_content: str, user_id: str
    ) -> Dict[str, Any]:
        """
        Update template with new HTML content and regenerate DOCX

        Args:
            template_id: Template ID to update
            html_content: New HTML content
            user_id: User ID for authorization

        Returns:
            Updated template info with new file URLs
        """
        try:
            from src.config.database import get_async_database

            logger.info(f"Updating template {template_id} with new content")

            # Convert HTML to DOCX
            conversion_result = await self.html_to_docx(
                html_content,
                file_name=f"template_{template_id}_updated",
                upload_to_r2=True,
            )

            # Update database
            db = await get_async_database()
            collection = db.templates

            update_data = {
                "files.docx_url": conversion_result["download_url"],
                "updated_at": datetime.utcnow(),
                "last_modified_by": user_id,
                "version": {"$inc": 1},  # Increment version
            }

            result = await collection.update_one(
                {"template_id": template_id, "user_id": user_id}, {"$set": update_data}
            )

            if result.modified_count == 0:
                raise Exception("Template not found or access denied")

            logger.info(f"Template {template_id} updated successfully")

            return {
                "success": True,
                "template_id": template_id,
                "new_docx_url": conversion_result["download_url"],
                "updated_at": update_data["updated_at"].isoformat(),
            }

        except Exception as e:
            logger.error(f"Template update failed: {e}")
            raise Exception(f"Template update failed: {str(e)}")
