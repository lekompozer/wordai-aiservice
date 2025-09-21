"""
AI Extraction Service for Products and Services Data
Service AI cho việc extract dữ liệu Products và Services

FLOW DESCRIPTION / MÔ TẢ LUỒNG XỬ LÝ:
1. Backend already uploads files to R2 and gets public URL
   Backend đã upload file lên R2 và lấy được public URL
2. Backend sends R2 URL + metadata to extraction API
   Backend gửi R2 URL + metadata đến extraction API
3. AI Service sends R2 URL (Company + Description + Appropriate Prompt) to AI Provider
   AI Service gửi R2 URL (Company + Description + Prompt phù hợp) đến AI Provider
4. Based on file type:
   - .txt files: Extract text by system, send to DeepSeek
   - All other files: Send to ChatGPT for extraction (images, PDFs, docs, etc.)
5. Save to Qdrant with proper tagging and metadata (can use worker)
   Lưu vào Qdrant với tag và metadata phù hợp (có thể dùng worker)
"""

import json
import re
import uuid
import asyncio
import aiohttp
import logging
import pandas as pd
import io
from docx import Document
from typing import Dict, Any, Optional, Union, List, AsyncGenerator
from datetime import datetime
from pathlib import Path
import uuid

from src.providers.ai_provider_manager import AIProviderManager
from src.services.extraction_templates.template_factory import ExtractionTemplateFactory
from config.config import DEEPSEEK_API_KEY, CHATGPT_API_KEY
import config.config as config
from src.models.unified_models import (
    QdrantDocumentChunk,
    IndustryDataType,
    Industry,
    Language,
)

logger = logging.getLogger(__name__)


class AIExtractionService:
    """
    AI Service for extracting Products and Services data from R2 URLs
    Service AI cho việc extract dữ liệu Products và Services từ R2 URLs
    """

    def __init__(self):
        """Initialize AI extraction service"""
        self.ai_manager = AIProviderManager(
            deepseek_api_key=DEEPSEEK_API_KEY,
            chatgpt_api_key=CHATGPT_API_KEY,
            gemini_api_key=config.GEMINI_API_KEY,
        )
        self.template_factory = ExtractionTemplateFactory()

        # Core extraction categories (AI will auto-categorize into these)
        self.extraction_categories = ["products", "services"]

        # Supported backend data types (for final mapping to Qdrant)
        self.supported_backend_types = [
            "company_info",
            "products",
            "services",
            "policies",
            "faq",
            "knowledge_base",
            "other",
        ]

        # File type mappings for AI provider selection
        self.text_file_extensions = [".txt", ".json", ".csv"]
        self.vision_file_extensions = [
            ".jpg",
            ".jpeg",
            ".png",
            ".webp",
            ".pdf",
            ".docx",
            ".doc",
            ".xlsx",
            ".xls",
        ]

        logger.info("🤖 AI Extraction Service initialized")
        logger.info(f"   📊 Extraction categories: {self.extraction_categories}")
        logger.info(f"   🗂️  Backend data types: {self.supported_backend_types}")
        logger.info(f"   🔑 ChatGPT configured: {bool(CHATGPT_API_KEY)} (for images)")
        logger.info(
            f"   🔑 Gemini configured: {bool(config.GEMINI_API_KEY)} (for text/docs)"
        )
        logger.info("   ✅ Using Gemini for most files, ChatGPT for images only")

    async def extract_from_r2_url(
        self,
        r2_url: str,
        metadata: Dict[str, Any],
        company_info: Optional[Dict[str, Any]] = None,
        target_categories: Optional[List[str]] = None,
    ) -> Dict[str, Any]:
        """
        Extract data from R2 URL with AI auto-categorization
        Extract dữ liệu từ R2 URL với AI tự động phân loại

        Args:
            r2_url: Public R2 URL of the file
            metadata: File metadata (filename, industry, etc.)
            company_info: Company information for context
            target_categories: Optional list to limit extraction to specific categories ["products", "services"]
                               If None, AI will extract both products and services automatically

        Returns:
            Extracted data with AI-categorized products and services
        """
        try:
            logger.info("🚀 [EXTRACTION START] Beginning extraction process...")

            # Default to extracting both categories if not specified
            if target_categories is None:
                target_categories = self.extraction_categories.copy()

            # Validate target categories
            invalid_categories = [
                cat
                for cat in target_categories
                if cat not in self.extraction_categories
            ]
            if invalid_categories:
                raise ValueError(
                    f"Invalid categories: {invalid_categories}. Supported: {self.extraction_categories}"
                )

            logger.info(f"🔗 Starting extraction from R2 URL: {r2_url}")
            logger.info(f"📊 Target categories: {target_categories}")
            logger.info(
                f"🏢 Company: {company_info.get('name', 'Unknown') if company_info else 'Unknown'}"
            )
            logger.info(f"🏭 Industry: {metadata.get('industry', 'Unknown')}")

            # Step 1: Determine AI provider based on file type
            logger.info("📌 [STEP 1] Determining AI provider...")

            # Try multiple sources for filename to get extension
            filename = (
                metadata.get("original_name")
                or metadata.get("filename")
                or (
                    metadata.get("file_metadata", {}).get("filename")
                    if isinstance(metadata.get("file_metadata"), dict)
                    else None
                )
                or Path(r2_url).name
                or ""
            )
            file_extension = Path(filename).suffix.lower()
            logger.info(
                f"🔍 Filename resolution: '{filename}' -> extension: '{file_extension}'"
            )
            ai_provider = self._select_ai_provider(file_extension)
            logger.info(
                f"🤖 Selected AI provider: {ai_provider} for extension: {file_extension}"
            )

            # Step 2: Get extraction template with fallback to generic
            logger.info("📌 [STEP 2] Getting extraction template...")
            template = self.template_factory.get_template_with_metadata(metadata)
            if template.__class__.__name__ == "GenericExtractionTemplate":
                logger.info(
                    f"📋 Using GenericExtractionTemplate (industry '{metadata.get('industry')}' not supported)"
                )
            else:
                logger.info(f"📋 Using template: {template.__class__.__name__}")

            # Step 3: Build extraction prompts for auto-categorization
            logger.info("📌 [STEP 3] Building extraction prompts...")
            system_prompt = self._build_auto_categorization_system_prompt(
                template, target_categories, company_info
            )
            user_prompt = self._build_auto_categorization_user_prompt(
                template, target_categories, metadata, company_info
            )
            logger.info(f"✅ System prompt length: {len(system_prompt)} chars")
            logger.info(f"✅ User prompt length: {len(user_prompt)} chars")

            # Step 4: Extract data using appropriate AI provider
            logger.info(f"📌 [STEP 4] Extracting with {ai_provider}...")

            if ai_provider == "chatgpt":
                # ChatGPT: Use for images (better vision capability)
                if file_extension in self.text_file_extensions:
                    logger.info("📥 Downloading text file for ChatGPT processing...")
                    file_content = await self._download_from_r2(r2_url)
                    logger.info(
                        f"📥 Downloaded {len(file_content)} bytes for text processing"
                    )

                    # Extract text content based on file type
                    logger.info("📄 Extracting text from file...")
                    text_content = await self._extract_text_from_file(
                        file_content, file_extension, metadata
                    )
                    logger.info(
                        f"📄 Extracted text content: {len(text_content)} characters"
                    )

                    # Use ChatGPT for text content processing
                    extracted_data = await self._extract_with_chatgpt_text(
                        text_content,
                        system_prompt,
                        user_prompt,
                        metadata,
                        target_categories,
                    )
                else:
                    # ChatGPT Vision: Send R2 URL directly (for images)
                    logger.info("🖼️ Using ChatGPT Vision - sending R2 URL directly...")
                    extracted_data = await self._extract_with_chatgpt_vision(
                        r2_url, system_prompt, user_prompt, metadata, target_categories
                    )

            elif ai_provider == "gemini":
                # Gemini: Use for all non-image files (text, docs, PDFs, etc.)
                logger.info("🔮 Using Gemini for file processing...")

                if file_extension in self.text_file_extensions:
                    # For text files, download and send as text to Gemini
                    logger.info(
                        "📥 Downloading text file for Gemini text processing..."
                    )
                    file_content = await self._download_from_r2(r2_url)
                    text_content = await self._extract_text_from_file(
                        file_content, file_extension, metadata
                    )
                    extracted_data = await self._extract_with_gemini_text(
                        text_content,
                        system_prompt,
                        user_prompt,
                        metadata,
                        target_categories,
                    )
                else:
                    # For other files (PDFs, docs, etc.), use Gemini file upload
                    logger.info(
                        "📄 Using Gemini file upload for document processing..."
                    )
                    file_content = await self._download_from_r2(r2_url)
                    filename = metadata.get(
                        "filename", metadata.get("original_name", "document")
                    )

                    # Try Gemini first, fallback to ChatGPT if fails
                    try:
                        extracted_data = await self._extract_with_gemini_file_upload(
                            file_content,
                            filename,
                            system_prompt,
                            user_prompt,
                            metadata,
                            target_categories,
                        )
                    except Exception as gemini_error:
                        logger.warning(f"⚠️ Gemini extraction failed: {gemini_error}")
                        logger.info("🔄 Falling back to ChatGPT...")

                        # Fallback to ChatGPT
                        extracted_data = await self._extract_with_chatgpt_file_upload(
                            file_content,
                            filename,
                            system_prompt,
                            user_prompt,
                            metadata,
                            target_categories,
                        )

            else:
                raise Exception(f"Unsupported AI provider: {ai_provider}")

            # Step 5: Post-process and validate result with auto-categorization
            logger.info("📌 [STEP 5] Post-processing results...")
            final_result = template.post_process_auto_categorized(
                extracted_data, target_categories
            )

            # Step 6: Count extracted items across all categories
            logger.info("📌 [STEP 6] Counting extracted items...")
            total_items = 0
            for category in target_categories:
                items = final_result.get(category, [])
                item_count = len(items) if isinstance(items, list) else 0
                total_items += item_count
                logger.info(f"   📊 {category.title()}: {item_count} items")

            # Step 7: Add extraction metadata
            logger.info("📌 [STEP 7] Adding extraction metadata...")
            final_result["extraction_metadata"] = {
                "r2_url": r2_url,
                "extraction_mode": "auto_categorization",
                "target_categories": target_categories,
                "ai_provider": ai_provider,
                "template_used": template.__class__.__name__,
                "industry": metadata.get("industry"),
                "company_name": company_info.get("name") if company_info else None,
                "file_extension": file_extension,
                "extraction_timestamp": datetime.now().isoformat(),
                "total_items": total_items,
                "source": "r2_extraction_service_v2",
                "categorization_summary": {
                    category: len(final_result.get(category, []))
                    for category in target_categories
                },
                # ✅ ADD FILE INFORMATION for catalog service
                "file_id": metadata.get("file_id"),
                "file_name": (
                    metadata.get("original_name")
                    or metadata.get("filename")
                    or filename
                ),
                "original_filename": filename,
            }

            logger.info(f"✅ Auto-categorization extraction completed successfully")
            logger.info(f"   📊 Total items: {total_items}")
            logger.info(f"   🤖 Provider: {ai_provider}")
            logger.info(f"   📋 Template: {template.__class__.__name__}")

            return final_result

        except Exception as e:
            logger.error(f"❌ R2 extraction failed: {str(e)}")
            raise Exception(f"R2 extraction failed: {str(e)}")

    async def extract_from_text_content(
        self,
        text_content: str,
        data_type: str,
        metadata: Dict[str, Any],
        company_info: Optional[Dict[str, Any]] = None,
    ) -> Dict[str, Any]:
        """
        Extract data from direct text content (for testing/debugging)
        Extract dữ liệu từ text content trực tiếp
        """
        try:
            # Validate data type
            if data_type not in self.extraction_categories:
                raise ValueError(f"Unsupported data type: {data_type}")

            logger.info(f"📝 Starting text extraction")
            logger.info(f"📊 Data type: {data_type}")
            logger.info(f"📄 Text length: {len(text_content)} characters")

            # Get extraction template
            template = self.template_factory.get_template_with_metadata(metadata)

            # Build prompts with company context
            system_prompt = self._build_auto_categorization_system_prompt(
                template, [data_type], company_info
            )
            user_prompt = self._build_auto_categorization_user_prompt(
                template, [data_type], metadata, company_info
            )

            # Add text content to prompt
            full_prompt = f"{user_prompt}\n\nDocument content:\n{text_content}"

            # Use ChatGPT for text content (stable and reliable)
            extracted_data = await self._call_chatgpt(system_prompt, full_prompt)

            # Post-process result
            final_result = template.post_process_auto_categorized(
                extracted_data, [data_type]
            )

            # Add metadata
            final_result["extraction_metadata"] = {
                "data_type": data_type,
                "ai_provider": "chatgpt",
                "template_used": template.__class__.__name__,
                "industry": metadata.get("industry"),
                "company_name": company_info.get("name") if company_info else None,
                "extraction_timestamp": datetime.now().isoformat(),
                "total_items": len(final_result.get(data_type, [])),
                "source": "text_extraction_service",
            }

            logger.info(
                f"✅ Text extraction completed: {len(final_result.get(data_type, []))} {data_type}"
            )
            return final_result

        except Exception as e:
            logger.error(f"❌ Text extraction failed: {str(e)}")
            raise Exception(f"Text extraction failed: {str(e)}")

    async def _download_from_r2(self, r2_url: str) -> bytes:
        """Download file content from R2 URL"""
        try:
            logger.info(f"🌐 [DOWNLOAD] Starting download from: {r2_url}")
            async with aiohttp.ClientSession() as session:
                async with session.get(r2_url) as response:
                    logger.info(f"📡 [DOWNLOAD] Response status: {response.status}")
                    if response.status == 200:
                        content = await response.read()
                        logger.info(f"✅ [DOWNLOAD] Downloaded {len(content)} bytes")
                        return content
                    else:
                        raise Exception(
                            f"Failed to download from R2: HTTP {response.status}"
                        )
        except Exception as e:
            logger.error(f"❌ R2 download failed: {str(e)}")
            raise Exception(f"R2 download failed: {str(e)}")

    def _select_ai_provider(self, file_extension: str) -> str:
        """
        Select AI provider based on file type
        Chọn AI provider dựa trên loại file

        Returns:
            "chatgpt" for images (jpg, jpeg, png, webp) - better vision capability
            "gemini" for all other files (text, docs, pdfs, etc.) - backup provider
        """
        # Use ChatGPT only for images (better vision capability)
        image_extensions = [".jpg", ".jpeg", ".png", ".webp"]

        if file_extension in image_extensions:
            logger.info(f"🖼️ Using ChatGPT for image file: {file_extension}")
            return "chatgpt"
        else:
            logger.info(f"📄 Using Gemini for file: {file_extension}")
            return "gemini"

    def _build_auto_categorization_system_prompt(
        self,
        template,
        target_categories: List[str],
        company_info: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Build system prompt for AI auto-categorization"""

        # Get JSON schemas for all target categories
        schemas = {}
        for category in target_categories:
            schemas[category] = template.get_extraction_schema(category)

        # Build comprehensive schema
        comprehensive_schema = {
            "raw_content": "Complete text content extracted from the document",
            "structured_data": {},
        }

        for category in target_categories:
            # ✅ FIXED: Provide the actual schema object structure for each item
            comprehensive_schema["structured_data"][category] = [
                {
                    "description": f"Each {category[:-1]} object must follow this exact schema",
                    "schema": schemas[category],
                }
            ]

        comprehensive_schema["structured_data"]["extraction_summary"] = {
            "total_products": "Number of products found (integer)",
            "total_services": "Number of services found (integer)",
            "data_quality": "Assessment of data quality: high, medium, or low",
            "categorization_notes": "Notes about how items were categorized",
            "industry_context": "Industry-specific observations",
        }

        schema_str = json.dumps(comprehensive_schema, indent=2, ensure_ascii=False)

        # ===== ADD TEMPLATE EXTRACTION INSTRUCTIONS =====
        template_instructions = ""
        for category in target_categories:
            instruction = template.get_system_prompt(category)
            template_instructions += (
                f"\n\nTEMPLATE INSTRUCTIONS FOR {category.upper()}:\n{instruction}"
            )

        base_prompt = f"""You are an AI data extraction specialist for the {template.__class__.__name__.replace('Template', '')} industry.

TASK / NHIỆM VỤ:
Extract and automatically categorize data into appropriate categories: {', '.join(target_categories)}

EXTRACTION + AI CATEGORIZATION INSTRUCTIONS / HƯỚNG DẪN EXTRACT + AI PHÂN LOẠI:
1. Extract ALL text content as raw_content
2. Analyze and categorize data intelligently:
   - PRODUCTS: Physical items, menu items, room types, loan products, courses, etc.
   - SERVICES: Services offered, amenities, support services, consulting, etc.
3. Some items may appear in both categories (e.g., "Room + Breakfast Package")
4. Use industry context to make smart categorization decisions
5. Preserve all important details (prices, descriptions, availability, etc.)

🎯 HYBRID SEARCH AI CATEGORIZATION (REQUIRED FOR ALL PRODUCTS/SERVICES):
For EACH product/service, you MUST include these categorization fields:

📂 **category**: Main category (snake_case Vietnamese, no accents)
   Examples: "bao_hiem", "dau_tu", "tiet_kiem", "tin_dung", "thanh_toan", "cong_nghe",
            "do_uong", "do_an", "phong_o", "dich_vu_luu_tru"

🔸 **sub_category**: Sub-category (snake_case Vietnamese, no accents)
   Examples: "bao_hiem_nhan_tho", "bao_hiem_suc_khoe", "mon_chinh", "do_uong_co_con",
            "phong_don", "phong_doi", "tu_van_tai_chinh"

🏷️ **tags**: Array of 5-8 relevant search tags (snake_case Vietnamese, no accents)
   Examples: ["gia_dinh", "tre_em", "nguoi_cao_tuoi", "cao_cap", "gia_re", "truyen_thong", "hien_dai"]

👥 **target_audience**: Array of target demographics (snake_case Vietnamese, no accents)
   Examples: ["gia_dinh_tre", "chuyen_gia", "doanh_nghiep_nho", "sinh_vien", "nguoi_lao_dong"]

�️ **coverage_type** (for insurance products): ["tai_nan", "benh_tat", "tu_vong", "thuong_tat"]
⚙️ **service_type** (for services): ["tu_van", "ho_tro", "dao_tao", "bao_tri", "thiet_ke"]

�🔥 CRITICAL TEMPLATE-SPECIFIC REQUIREMENTS:
{template_instructions}

CRITICAL DATA STRUCTURE REQUIREMENTS / YÊU CẦU CẤU TRÚC DỮ LIỆU QUAN TRỌNG:
- Each item in products/services arrays MUST be a complete object with AI categorization
- ALL products/services MUST have: category, sub_category, tags, target_audience
- Use snake_case for categorization fields (no spaces, no accents)
- Tags must be relevant and specific, avoid generic words
- For restaurant: "Vịt chiên Hong Kong 260K" should become:
  {{
    "name": "Vịt chiên Hong Kong",
    "price": 260000,
    "currency": "VND",
    "description": "Traditional Hong Kong style roasted duck",
    "category": "do_an",
    "sub_category": "mon_chinh",
    "tags": ["vit", "hong_kong", "chien", "truyen_thong", "cao_cap"],
    "target_audience": ["gia_dinh", "khach_vip", "yeu_thich_mon_a"]
  }}

RESPONSE SCHEMA / SCHEMA PHẢN HỒI:
```json
{schema_str}
```

CRITICAL REQUIREMENTS / YÊU CẦU QUAN TRỌNG:
- Return EXACTLY the schema structure above
- Each product/service MUST be a detailed object with all available fields
- Use Vietnamese for names/descriptions unless English specified
- Include both raw_content AND structured_data
- Categorize intelligently based on industry context
- Convert price notation (450K = 450000, 25K = 25000)
- Provide extraction_summary with categorization insights
- Use proper data types (numbers for prices, arrays for lists)
- For missing data, use empty arrays [] not null
- 🔥 CRITICAL: Follow ALL template-specific instructions above, especially content_for_embedding requirements"""

        if company_info:
            company_context = f"""

COMPANY CONTEXT / BỐI CẢNH CÔNG TY:
- Company Name: {company_info.get('name', 'Unknown')}
- Industry: {company_info.get('industry', 'Unknown')}
- Description: {company_info.get('description', 'Not provided')}

Use this company context to improve categorization accuracy.
Sử dụng bối cảnh công ty để cải thiện độ chính xác phân loại."""
            return f"{base_prompt}\n{company_context}"

        return base_prompt

    def _build_auto_categorization_user_prompt(
        self,
        template,
        target_categories: List[str],
        metadata: Dict[str, Any],
        company_info: Optional[Dict[str, Any]] = None,
    ) -> str:
        """Build user prompt for auto-categorization extraction"""

        extraction_instructions = f"""
AUTO-CATEGORIZATION EXTRACTION TASK / NHIỆM VỤ EXTRACT TỰ ĐỘNG PHÂN LOẠI:

TARGET CATEGORIES / DANH MỤC MUC TIÊU: {', '.join(target_categories)}
INDUSTRY / NGÀNH: {metadata.get('industry', 'generic')}
FILE / TẬP TIN: {metadata.get('original_name', 'document')}
LANGUAGE / NGÔN NGỮ: {metadata.get('language', 'vi')}

INTELLIGENT CATEGORIZATION GUIDELINES / HƯỚNG DẪN PHÂN LOẠI THÔNG MINH:

FOR PRODUCTS / CHO PRODUCTS:
- Physical items that can be purchased/ordered
- Menu items (food, drinks, dishes)
- Room types and accommodation
- Financial products (loans, accounts, cards)
- Courses and educational programs
- Physical goods or tangible offerings

FOR SERVICES / CHO SERVICES:
- Intangible services provided
- Support and customer service
- Consulting and advisory services
- Maintenance and repair services
- Delivery and transportation
- Professional services

SMART CATEGORIZATION EXAMPLES / VÍ DỤ PHÂN LOẠI THÔNG MINH:
- "Phở Bò + Giao hàng" → Products: Phở Bò, Services: Giao hàng
- "Deluxe Room with Spa Service" → Products: Deluxe Room, Services: Spa Service
- "Personal Loan with Advisory" → Products: Personal Loan, Services: Advisory

EXTRACTION INSTRUCTIONS / HƯỚNG DẪN EXTRACT:
1. Read and extract ALL content as raw_content
2. Identify and categorize items intelligently
3. Preserve complete information (prices, descriptions, features)
4. Use industry knowledge for accurate categorization
5. Provide insights in extraction_summary

Please analyze the document and extract all relevant information with intelligent auto-categorization.
Vui lòng phân tích tài liệu và extract tất cả thông tin liên quan với tự động phân loại thông minh."""

        return extraction_instructions

    async def _extract_with_chatgpt_vision(
        self,
        r2_url: str,
        system_prompt: str,
        user_prompt: str,
        metadata: Dict[str, Any],
        target_categories: List[str],
    ) -> Dict[str, Any]:
        """
        Extract using ChatGPT Vision with R2 URL (no download needed)
        Extract sử dụng ChatGPT Vision với R2 URL (không cần download)
        """
        try:
            logger.info("🖼️ Using ChatGPT Vision for file processing (sending R2 URL)")

            # Enhanced system prompt with vision and JSON instructions
            enhanced_system_prompt = f"""{system_prompt}

VISION PROCESSING INSTRUCTIONS:
- Analyze the image/document content carefully
- Extract all visible text, tables, and structured data
- Pay attention to layout, formatting, and visual hierarchy
- Include price information, product details, and categorization

RAW DATA + JSON REQUIREMENT / YÊU CẦU RAW DATA + JSON:
For ChatGPT Vision, you must return BOTH:
1. raw_content: All extracted text and content from the file
2. structured_data: JSON formatted according to the schema

Return format:
{{
  "raw_content": "All visible text, tables, and content extracted from the image/document...",
  "structured_data": {{
    "{metadata.get('data_type', 'products')}": [...],
    "extraction_summary": {{...}}
  }}
}}

CRITICAL JSON FORMATTING INSTRUCTIONS:
- Return ONLY valid JSON format
- Include both raw_content and structured_data sections
- Use proper JSON syntax with double quotes
- Ensure all numbers are numeric (not strings)
- For missing values, use null instead of empty strings
- Array fields should be empty arrays [] if no data found
- Do not wrap JSON in code blocks or markdown
- Start response directly with {{ and end with }}
"""

            # Enhanced user prompt for vision with R2 URL
            enhanced_user_prompt = f"""{user_prompt}

Please analyze the file at this URL and extract the requested information.
File URL: {r2_url}

IMPORTANT INSTRUCTIONS / HƯỚNG DẪN QUAN TRỌNG:
1. Analyze the content of the file from the provided URL
2. Extract ALL visible text as raw_content
3. Process and structure the data according to the JSON schema
4. Return BOTH raw content and structured data in the format specified
5. Return ONLY the JSON response, no additional text, explanations, or formatting

Expected response format:
{{
  "raw_content": "Complete text content from the file...",
  "structured_data": {{
    {', '.join([f'"{cat}": [...extracted {cat}...]' for cat in target_categories])},
    "extraction_summary": {{
      "total_items": number,
      "data_quality": "high|medium|low",
      "extraction_notes": "Any relevant notes..."
    }}
  }}
}}"""

            # TODO: Implement actual ChatGPT Vision API call with R2 URL
            # This would use messages with image_url parameter:
            messages = [
                {"role": "system", "content": enhanced_system_prompt},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": enhanced_user_prompt},
                        {"type": "image_url", "image_url": {"url": r2_url}},
                    ],
                },
            ]

            response = await self.ai_manager.chat_completion(
                messages=messages, provider="chatgpt"
            )

            return self._parse_ai_response(response)

        except Exception as e:
            logger.error(f"❌ ChatGPT Vision extraction failed: {str(e)}")
            raise e

    async def _extract_text_from_file(
        self, file_content: bytes, file_extension: str, metadata: Dict[str, Any]
    ) -> str:
        """
        Extract text content from different file types
        Extract text content từ các loại file khác nhau
        """
        try:
            file_name = metadata.get("original_name", "unknown")
            logger.info(f"📄 Extracting text from {file_extension} file: {file_name}")

            if file_extension == ".txt":
                # Plain text file
                text_content = file_content.decode("utf-8")
                logger.info(f"✅ TXT file decoded successfully")
                return text_content

            elif file_extension == ".csv":
                # CSV file - convert to readable text format
                try:
                    # Try to read as CSV with pandas
                    csv_content = file_content.decode("utf-8")
                    df = pd.read_csv(io.StringIO(csv_content))

                    # Convert DataFrame to readable text format
                    text_lines = []
                    text_lines.append(f"CSV FILE: {file_name}")
                    text_lines.append(f"Total Rows: {len(df)}")
                    text_lines.append(f"Columns: {', '.join(df.columns.tolist())}")
                    text_lines.append("")

                    # Add header
                    text_lines.append("HEADER:")
                    text_lines.append(" | ".join(df.columns.tolist()))
                    text_lines.append("")

                    # Add data rows (limit to prevent prompt overflow)
                    text_lines.append("DATA ROWS:")
                    max_rows = min(100, len(df))  # Limit to 100 rows
                    for idx, row in df.head(max_rows).iterrows():
                        row_text = " | ".join([str(val) for val in row.values])
                        text_lines.append(f"Row {idx + 1}: {row_text}")

                    if len(df) > max_rows:
                        text_lines.append(f"... and {len(df) - max_rows} more rows")

                    text_content = "\n".join(text_lines)
                    logger.info(
                        f"✅ CSV file processed: {len(df)} rows, {len(df.columns)} columns"
                    )
                    return text_content

                except Exception as csv_error:
                    logger.warning(
                        f"⚠️ CSV parsing failed: {csv_error}, treating as plain text"
                    )
                    text_content = file_content.decode("utf-8")
                    return text_content

            elif file_extension == ".docx":
                # Word document
                try:
                    doc = Document(io.BytesIO(file_content))
                    text_lines = []
                    text_lines.append(f"WORD DOCUMENT: {file_name}")
                    text_lines.append("")

                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_lines.append(paragraph.text)

                    # Extract text from tables
                    for table in doc.tables:
                        text_lines.append("\nTABLE:")
                        for row in table.rows:
                            row_text = " | ".join(
                                [cell.text.strip() for cell in row.cells]
                            )
                            if row_text.strip():
                                text_lines.append(row_text)

                    text_content = "\n".join(text_lines)
                    logger.info(
                        f"✅ DOCX file processed: {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables"
                    )
                    return text_content

                except Exception as docx_error:
                    logger.error(f"❌ DOCX parsing failed: {docx_error}")
                    raise Exception(f"Failed to process DOCX file: {docx_error}")

            elif file_extension == ".json":
                # JSON file - format as readable text
                try:
                    json_content = file_content.decode("utf-8")
                    json_data = json.loads(json_content)

                    text_lines = []
                    text_lines.append(f"JSON FILE: {file_name}")
                    text_lines.append("")
                    text_lines.append("FORMATTED JSON CONTENT:")
                    text_lines.append(
                        json.dumps(json_data, indent=2, ensure_ascii=False)
                    )

                    text_content = "\n".join(text_lines)
                    logger.info(f"✅ JSON file processed successfully")
                    return text_content

                except Exception as json_error:
                    logger.warning(
                        f"⚠️ JSON parsing failed: {json_error}, treating as plain text"
                    )
                    text_content = file_content.decode("utf-8")
                    return text_content

            else:
                # Unknown text file type, try UTF-8 decode
                try:
                    text_content = file_content.decode("utf-8")
                    logger.info(f"✅ File decoded as UTF-8 text")
                    return text_content
                except UnicodeDecodeError:
                    raise Exception(
                        f"Unsupported file type {file_extension} or not a text file"
                    )

        except Exception as e:
            logger.error(f"❌ Text extraction failed: {str(e)}")
            raise Exception(f"Text extraction failed for {file_extension}: {str(e)}")

    async def _extract_with_chatgpt_text(
        self,
        text_content: str,
        system_prompt: str,
        user_prompt: str,
        metadata: Dict[str, Any],
        target_categories: List[str],
    ) -> Dict[str, Any]:
        """
        Extract using ChatGPT for text content (enhanced version)
        Extract sử dụng ChatGPT cho text content (phiên bản nâng cao)
        """
        try:
            logger.info("📝 [CHATGPT_TEXT] Using ChatGPT for text file processing")
            logger.info(
                f"📄 [CHATGPT_TEXT] Text content length: {len(text_content)} characters"
            )

            # Enhanced user prompt with text content
            enhanced_user_prompt = f"""{user_prompt}

DOCUMENT CONTENT TO PROCESS:
{text_content}

RAW DATA + JSON REQUIREMENT / YÊU CẦU RAW DATA + JSON:
For ChatGPT text processing, you must return BOTH:
1. raw_content: The original text content as processed
2. structured_data: JSON formatted according to the schema

IMPORTANT INSTRUCTIONS / HƯỚNG DẪN QUAN TRỌNG:
1. Analyze the above text content carefully
2. Extract ALL text as raw_content (preserve original content)
3. Process and structure the data according to the JSON schema
4. Return BOTH raw content and structured data in the format specified
5. Return ONLY the JSON response, no additional text, explanations, or formatting

Expected response format:
{{
  "raw_content": "Complete original content of the file...",
  "structured_data": {{
    {', '.join([f'"{cat}": [...extracted {cat}...]' for cat in target_categories])},
    "extraction_summary": {{
      "total_items": number,
      "data_quality": "high|medium|low",
      "extraction_notes": "Any relevant notes..."
    }}
  }}
}}"""

            logger.info(
                f"✅ [CHATGPT_TEXT] Enhanced prompt prepared, length: {len(enhanced_user_prompt)} chars"
            )
            logger.info("🚀 [CHATGPT_TEXT] Calling ChatGPT API...")

            result = await self._call_chatgpt(system_prompt, enhanced_user_prompt)

            logger.info("✅ [CHATGPT_TEXT] ChatGPT API call completed successfully")
            return result

        except Exception as e:
            logger.error(f"❌ [CHATGPT_TEXT] ChatGPT text extraction failed: {str(e)}")
            raise e

    async def _extract_with_gemini_text(
        self,
        text_content: str,
        system_prompt: str,
        user_prompt: str,
        metadata: Dict[str, Any],
        target_categories: List[str],
    ) -> Dict[str, Any]:
        """
        Extract using Gemini for text content
        Extract sử dụng Gemini cho text content
        """
        try:
            logger.info("📝 [GEMINI_TEXT] Using Gemini for text file processing")
            logger.info(
                f"📄 [GEMINI_TEXT] Text content length: {len(text_content)} characters"
            )

            # Enhanced user prompt with text content for Gemini
            enhanced_user_prompt = f"""{user_prompt}

DOCUMENT CONTENT TO PROCESS:
{text_content}

RAW DATA + JSON REQUIREMENT / YÊU CẦU RAW DATA + JSON:
For Gemini text processing, you must return BOTH:
1. raw_content: The original text content as processed
2. structured_data: JSON formatted according to the schema

IMPORTANT INSTRUCTIONS / HƯỚNG DẪN QUAN TRỌNG:
1. Analyze the above text content carefully
2. Extract ALL text as raw_content (preserve original content)
3. Process and structure the data according to the JSON schema
4. Return BOTH raw content and structured_data in the format specified
5. Return ONLY the JSON response, no additional text, explanations, or formatting

Expected response format:
{{
  "raw_content": "Complete original content of the file...",
  "structured_data": {{
    {', '.join([f'"{cat}": [...extracted {cat}...]' for cat in target_categories])},
    "extraction_summary": {{
      "total_items": number,
      "data_quality": "high|medium|low",
      "extraction_notes": "Any relevant notes..."
    }}
  }}
}}"""

            logger.info(
                f"✅ [GEMINI_TEXT] Enhanced prompt prepared, length: {len(enhanced_user_prompt)} chars"
            )
            logger.info("🚀 [GEMINI_TEXT] Calling Gemini API...")

            result = await self._call_gemini(system_prompt, enhanced_user_prompt)

            logger.info("✅ [GEMINI_TEXT] Gemini API call completed successfully")
            return result

        except Exception as e:
            logger.error(f"❌ [GEMINI_TEXT] Gemini text extraction failed: {str(e)}")
            raise e

    async def _extract_with_gemini_file_upload(
        self,
        file_content: bytes,
        filename: str,
        system_prompt: str,
        user_prompt: str,
        metadata: Dict[str, Any],
        target_categories: List[str],
    ) -> Dict[str, Any]:
        """
        Extract using Gemini file upload for documents (PDFs, docs, etc.)
        Extract sử dụng Gemini file upload cho documents
        """
        try:
            logger.info(
                "📄 [GEMINI_FILE] Using Gemini file upload for document processing"
            )
            logger.info(
                f"📁 [GEMINI_FILE] File: {filename}, Size: {len(file_content)} bytes"
            )

            # Enhanced user prompt for file upload
            enhanced_user_prompt = f"""{user_prompt}

FILE UPLOAD PROCESSING INSTRUCTIONS / HƯỚNG DẪN XỬ LÝ FILE UPLOAD:
The file "{filename}" has been uploaded for analysis.

RAW DATA + JSON REQUIREMENT / YÊU CẦU RAW DATA + JSON:
For Gemini file processing, you must return BOTH:
1. raw_content: All text and content extracted from the uploaded file
2. structured_data: JSON formatted according to the schema

IMPORTANT INSTRUCTIONS / HƯỚNG DẪN QUAN TRỌNG:
1. Analyze the uploaded file content carefully
2. Extract ALL text and data as raw_content
3. Process and structure the data according to the JSON schema
4. Return BOTH raw content and structured data in the format specified
5. Return ONLY the JSON response, no additional text, explanations, or formatting

Expected response format:
{{
  "raw_content": "Complete content extracted from the file...",
  "structured_data": {{
    {', '.join([f'"{cat}": [...extracted {cat}...]' for cat in target_categories])},
    "extraction_summary": {{
      "total_items": number,
      "data_quality": "high|medium|low",
      "extraction_notes": "Any relevant notes...",
      "file_analysis": "Analysis of the uploaded file"
    }}
  }}
}}"""

            logger.info("🚀 [GEMINI_FILE] Calling Gemini file upload API...")

            # Use Gemini's chat_with_file_stream method
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": enhanced_user_prompt},
            ]

            # Collect streaming response
            response_chunks = []
            async for chunk in self.ai_manager.chat_with_file_stream(
                messages=messages,
                file_content=file_content,
                file_name=filename,
                provider="gemini",
            ):
                response_chunks.append(chunk)

            response = "".join(response_chunks)
            logger.info("✅ [GEMINI_FILE] Gemini file upload completed successfully")

            result = self._parse_ai_response(response)
            logger.info("✅ [GEMINI_FILE] Response parsed successfully")

            return result

        except Exception as e:
            logger.error(f"❌ [GEMINI_FILE] Gemini file upload failed: {str(e)}")
            raise e

    async def _extract_with_chatgpt_file_upload(
        self,
        file_content: bytes,
        filename: str,
        system_prompt: str,
        user_prompt: str,
        metadata: Dict[str, Any],
        target_categories: List[str],
    ) -> Dict[str, Any]:
        """
        Extract using ChatGPT with text-based approach (fallback from Gemini)
        Extract sử dụng ChatGPT với text-based approach (fallback từ Gemini)
        """
        try:
            logger.info(
                "📄 [CHATGPT_FILE] Using ChatGPT text-based processing (Gemini fallback)"
            )
            logger.info(
                f"📁 [CHATGPT_FILE] File: {filename}, Size: {len(file_content)} bytes"
            )

            # Extract text content from file
            extracted_text = await self._extract_text_content(file_content, filename)

            # Enhanced user prompt for text-based processing (same as Gemini)
            enhanced_user_prompt = f"""{user_prompt}

FILE CONTENT PROCESSING / XỬ LÝ NỘI DUNG FILE:
The file "{filename}" content has been extracted as text for analysis.

EXTRACTED CONTENT:
{extracted_text[:30000]}  # Limit to 30K chars to avoid token limits

RAW DATA + JSON REQUIREMENT / YÊU CẦU RAW DATA + JSON:
You must return BOTH:
1. raw_content: All text and content from the file
2. structured_data: JSON formatted according to the schema

IMPORTANT INSTRUCTIONS / HƯỚNG DẪN QUAN TRỌNG:
1. Analyze the extracted content carefully
2. Use ALL the extracted text as raw_content
3. Process and structure the data according to the JSON schema
4. Return BOTH raw content and structured data in the format specified
5. Return ONLY the JSON response, no additional text, explanations, or formatting

Expected response format:
{{
  "raw_content": "Complete content extracted from the file...",
  "structured_data": {{
    {', '.join([f'"{cat}": [...extracted {cat}...]' for cat in target_categories])},
    "extraction_summary": {{
      "total_items": number,
      "data_quality": "high|medium|low",
      "extraction_notes": "Any relevant notes...",
      "file_analysis": "Analysis of the file content"
    }}
  }}
}}"""

            logger.info("🚀 [CHATGPT_FILE] Calling ChatGPT text API...")

            # Use ChatGPT text method with extracted content
            result = await self._extract_with_chatgpt_text(
                extracted_text,
                system_prompt,
                enhanced_user_prompt,
                metadata,
                target_categories,
            )

            logger.info("✅ [CHATGPT_FILE] ChatGPT fallback completed successfully")
            return result

        except Exception as e:
            logger.error(f"❌ [CHATGPT_FILE] ChatGPT fallback failed: {str(e)}")
            raise e

    async def _extract_text_content(self, file_content: bytes, filename: str) -> str:
        """
        Extract text content from file for ChatGPT processing
        """
        try:
            import tempfile
            import os

            # Determine file type from filename
            file_ext = filename.lower().split(".")[-1]

            # Create temporary file
            with tempfile.NamedTemporaryFile(
                delete=False, suffix=f".{file_ext}"
            ) as tmp_file:
                tmp_file.write(file_content)
                tmp_file_path = tmp_file.name

            try:
                if file_ext in ["docx"]:
                    # Handle DOCX files
                    try:
                        import docx

                        doc = docx.Document(tmp_file_path)
                        text_content = []
                        for paragraph in doc.paragraphs:
                            if paragraph.text.strip():
                                text_content.append(paragraph.text)
                        return "\n".join(text_content)
                    except ImportError:
                        logger.warning("⚠️ python-docx not available, treating as text")
                        return file_content.decode("utf-8", errors="ignore")

                elif file_ext in ["pdf"]:
                    # Handle PDF files
                    try:
                        import PyPDF2

                        with open(tmp_file_path, "rb") as file:
                            pdf_reader = PyPDF2.PdfReader(file)
                            text_content = []
                            for page in pdf_reader.pages:
                                page_text = page.extract_text()
                                if page_text.strip():
                                    text_content.append(page_text)
                            return "\n".join(text_content)
                    except ImportError:
                        logger.warning("⚠️ PyPDF2 not available, treating as text")
                        return file_content.decode("utf-8", errors="ignore")

                else:
                    # Handle text files and others
                    return file_content.decode("utf-8", errors="ignore")

            finally:
                # Clean up temporary file
                if os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)

        except Exception as e:
            logger.error(f"❌ Text extraction failed: {e}")
            # Fallback: return raw content as text
            return file_content.decode("utf-8", errors="ignore")

    async def _call_gemini(
        self, system_prompt: str, user_prompt: str
    ) -> Dict[str, Any]:
        """Call Gemini API with JSON formatting instructions"""
        try:
            logger.info("🔮 [CALL_GEMINI] Starting Gemini API call...")
            logger.info(
                f"📝 [CALL_GEMINI] System prompt length: {len(system_prompt)} chars"
            )
            logger.info(
                f"📝 [CALL_GEMINI] User prompt length: {len(user_prompt)} chars"
            )

            # Enhanced system prompt with JSON formatting for Gemini
            enhanced_system_prompt = f"""{system_prompt}

CRITICAL JSON FORMATTING INSTRUCTIONS FOR GEMINI:
- Return ONLY valid JSON format
- Do not include any text before or after the JSON
- Use proper JSON syntax with double quotes
- Ensure all numbers are numeric (not strings)
- For missing values, use null instead of empty strings
- Array fields should be empty arrays [] if no data found
- Do not wrap JSON in code blocks or markdown
- Start response directly with {{ and end with }}
- Be precise and follow the schema exactly
"""

            # Enhanced user prompt
            enhanced_user_prompt = f"""{user_prompt}

IMPORTANT: Return ONLY the JSON response, no additional text, explanations, or formatting."""

            messages = [
                {"role": "system", "content": enhanced_system_prompt},
                {"role": "user", "content": enhanced_user_prompt},
            ]

            logger.info(
                f"📊 [CALL_GEMINI] Created messages array with {len(messages)} messages"
            )
            logger.info("🚀 [CALL_GEMINI] Calling Gemini API...")

            response = await self.ai_manager.chat_completion(
                messages=messages, provider="gemini"
            )

            logger.info("✅ [CALL_GEMINI] Gemini API call completed successfully")
            result = self._parse_ai_response(response)
            logger.info("✅ [CALL_GEMINI] Response parsed successfully")

            return result

        except Exception as e:
            logger.error(f"❌ [CALL_GEMINI] Gemini API call failed: {str(e)}")
            raise Exception(f"Gemini API call failed: {str(e)}")

    async def _call_chatgpt(
        self, system_prompt: str, user_prompt: str
    ) -> Dict[str, Any]:
        """Call ChatGPT API with JSON formatting instructions"""
        try:
            logger.info("🤖 [CALL_CHATGPT] Starting ChatGPT API call...")

            # Enhanced system prompt with JSON formatting for ChatGPT
            enhanced_system_prompt = f"""{system_prompt}

CRITICAL JSON FORMATTING INSTRUCTIONS FOR CHATGPT:
- Return ONLY valid JSON format
- Do not include any text before or after the JSON
- Use proper JSON syntax with double quotes
- Ensure all numbers are numeric (not strings)
- For missing values, use null instead of empty strings
- Array fields should be empty arrays [] if no data found
- Do not wrap JSON in code blocks or markdown
- Start response directly with {{ and end with }}
- Be precise and follow the schema exactly
"""

            # Enhanced user prompt
            enhanced_user_prompt = f"""{user_prompt}

IMPORTANT: Return ONLY the JSON response, no additional text, explanations, or formatting."""

            messages = [
                {"role": "system", "content": enhanced_system_prompt},
                {"role": "user", "content": enhanced_user_prompt},
            ]

            response = await self.ai_manager.chat_completion(
                messages=messages, provider="chatgpt"
            )

            logger.info("✅ [CALL_CHATGPT] ChatGPT API call completed successfully")
            result = self._parse_ai_response(response)
            logger.info("✅ [CALL_CHATGPT] Response parsed successfully")

            return result

        except Exception as e:
            logger.error(f"❌ [CALL_CHATGPT] ChatGPT API call failed: {str(e)}")
            raise Exception(f"ChatGPT API call failed: {str(e)}")

    async def test_connection(self) -> Dict[str, Any]:
        """Test connection to AI providers and services"""
        try:
            logger.info("🔍 Testing AI service connections...")
            results = {
                "overall_status": "healthy",
                "ai_providers": {},
                "template_system": {
                    "status": "healthy",
                    "available_templates": len(self.template_factory._templates),
                },
                "extraction_categories": self.extraction_categories,
                "supported_file_types": self.get_supported_file_types(),
            }

            # Test DeepSeek connection (for text files)
            try:
                test_response = await self.ai_manager.chat_completion(
                    messages=[{"role": "user", "content": "Test connection"}],
                    provider="deepseek",
                )
                results["ai_providers"]["deepseek"] = {
                    "status": "connected",
                    "response_received": len(str(test_response)) > 0,
                }
                logger.info("✅ DeepSeek connection test successful")
            except Exception as e:
                results["ai_providers"]["deepseek"] = {
                    "status": "error",
                    "error": str(e),
                }
                logger.warning(f"⚠️ DeepSeek connection test failed: {str(e)}")

            # Test ChatGPT connection (for vision files)
            try:
                test_response = await self.ai_manager.chat_completion(
                    messages=[{"role": "user", "content": "Test connection"}],
                    provider="chatgpt",
                )
                results["ai_providers"]["chatgpt"] = {
                    "status": "connected",
                    "response_received": len(str(test_response)) > 0,
                }
                logger.info("✅ ChatGPT connection test successful")
            except Exception as e:
                results["ai_providers"]["chatgpt"] = {
                    "status": "error",
                    "error": str(e),
                }
                logger.warning(f"⚠️ ChatGPT connection test failed: {str(e)}")

            # Check if any provider failed
            failed_providers = [
                name
                for name, info in results["ai_providers"].items()
                if info["status"] == "error"
            ]
            if failed_providers:
                results["overall_status"] = "unhealthy"
                logger.warning(f"⚠️ Failed providers: {failed_providers}")

            logger.info(
                f"🔍 Connection test completed. Status: {results['overall_status']}"
            )
            return results

        except Exception as e:
            logger.error(f"❌ Connection test failed: {str(e)}")
            return {
                "overall_status": "error",
                "error": str(e),
                "ai_providers": {},
                "template_system": {"status": "error"},
                "extraction_categories": [],
                "supported_file_types": [],
            }

    def get_supported_file_types(self) -> List[str]:
        """Get list of supported file types"""
        return [
            "text/plain",
            "application/pdf",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "text/csv",
            "application/json",
            "image/jpeg",
            "image/png",
            "image/webp",
        ]

    def get_service_info(self) -> Dict[str, Any]:
        """Get complete service information"""
        return {
            "service_name": "AI Extraction Service",
            "version": "2.0",
            "supported_industries": list(self.template_factory._templates.keys()),
            "supported_file_types": self.get_supported_file_types(),
            "extraction_categories": self.extraction_categories,
            "ai_providers": ["chatgpt", "gemini"],
            "features": [
                "Industry-specific templates",
                "Auto-categorization",
                "Optimized chunking",
                "Multi-language support",
                "Company context awareness",
            ],
        }

    def _parse_ai_response(self, response: str) -> Dict[str, Any]:
        """Parse AI response and return structured data"""
        try:
            logger.info("🔍 [PARSE_AI_RESPONSE] Parsing AI response...")

            # Remove any markdown code blocks if present
            if "```json" in response:
                # Extract JSON from markdown code blocks
                import re

                json_match = re.search(r"```json\s*(.*?)\s*```", response, re.DOTALL)
                if json_match:
                    response = json_match.group(1)
                    logger.info(
                        "✅ [PARSE_AI_RESPONSE] Extracted JSON from markdown blocks"
                    )

            # Remove any extra text before/after JSON
            response = response.strip()

            # Find JSON object boundaries
            start_idx = response.find("{")
            end_idx = response.rfind("}") + 1

            if start_idx != -1 and end_idx > start_idx:
                json_str = response[start_idx:end_idx]
                logger.info(
                    f"🔍 [PARSE_AI_RESPONSE] Extracted JSON string (length: {len(json_str)})"
                )
            else:
                json_str = response
                logger.warning(
                    "⚠️ [PARSE_AI_RESPONSE] Could not find JSON boundaries, using full response"
                )

            # Parse JSON
            try:
                parsed_data = json.loads(json_str)
                logger.info("✅ [PARSE_AI_RESPONSE] Successfully parsed JSON response")

                # Validate structure
                if isinstance(parsed_data, dict):
                    return parsed_data
                else:
                    logger.warning(
                        "⚠️ [PARSE_AI_RESPONSE] Response is not a dictionary, wrapping it"
                    )
                    return {"data": parsed_data}

            except json.JSONDecodeError as e:
                logger.error(f"❌ [PARSE_AI_RESPONSE] JSON decode error: {str(e)}")
                logger.error(
                    f"❌ [PARSE_AI_RESPONSE] Raw response preview: {json_str[:200]}..."
                )

                # Try to fix common JSON issues
                json_str = self._fix_common_json_issues(json_str)

                try:
                    parsed_data = json.loads(json_str)
                    logger.info(
                        "✅ [PARSE_AI_RESPONSE] Successfully parsed after fixing JSON issues"
                    )
                    return parsed_data
                except json.JSONDecodeError as e2:
                    logger.error(
                        f"❌ [PARSE_AI_RESPONSE] Still failed after JSON fixes: {str(e2)}"
                    )
                    raise Exception(f"Failed to parse AI response as JSON: {str(e2)}")

        except Exception as e:
            logger.error(f"❌ [PARSE_AI_RESPONSE] Error parsing AI response: {str(e)}")
            raise Exception(f"Failed to parse AI response: {str(e)}")

    def _fix_common_json_issues(self, json_str: str) -> str:
        """Fix common JSON formatting issues"""
        try:
            logger.info("🔧 [FIX_JSON] Attempting to fix common JSON issues...")

            # Remove trailing commas
            json_str = re.sub(r",\s*}", "}", json_str)
            json_str = re.sub(r",\s*]", "]", json_str)

            # Fix unescaped quotes in strings
            json_str = re.sub(r'(?<!\\)"(?![,\]\}:])', '\\"', json_str)

            # Fix single quotes to double quotes (basic case)
            json_str = re.sub(r"'([^']*)':", r'"\1":', json_str)

            logger.info("✅ [FIX_JSON] Applied common JSON fixes")
            return json_str

        except Exception as e:
            logger.warning(f"⚠️ [FIX_JSON] Error fixing JSON: {str(e)}")
            return json_str

    async def prepare_for_qdrant_ingestion(
        self,
        extraction_result: Dict[str, Any],
        user_id: str,
        document_id: str,
        company_id: str,
        industry: "Industry",
        language: "Language",
        callback_url: Optional[str] = None,
    ) -> Dict[str, Any]:
        """
        Prepare extraction results for Qdrant ingestion
        Chuẩn bị kết quả extraction cho Qdrant ingestion
        """
        try:
            logger.info(
                "🔄 [QDRANT_PREP] Preparing extraction results for Qdrant ingestion"
            )

            # Debug the extraction_result structure
            logger.info(
                f"🔍 [QDRANT_PREP] extraction_result keys: {list(extraction_result.keys())}"
            )

            from src.models.unified_models import IndustryDataType
            from datetime import datetime

            qdrant_chunks = []

            # --- ADAPTIVE DATA ACCESS - Support both nested and flat structures ---
            # Check if data is in 'structured_data' (sync /process endpoint) or top-level (async workflow)
            if "structured_data" in extraction_result:
                # Sync endpoint format: data nested in structured_data
                data_source = extraction_result["structured_data"]
                logger.info(
                    "✅ [QDRANT_PREP] Using nested 'structured_data' format (sync endpoint)"
                )
            else:
                # Async endpoint format: data at top level
                data_source = extraction_result
                logger.info("✅ [QDRANT_PREP] Using flat format (async endpoint)")

            # Process products from the correct data source
            products = data_source.get("products", [])
            logger.info(
                f"🔍 [QDRANT_PREP] Found {len(products)} products in data source"
            )

            for i, product in enumerate(products):
                chunk_id = str(uuid.uuid4())

                # Create optimized content for embedding
                content_for_embedding = product.get("content_for_embedding")
                if not content_for_embedding:
                    # Fallback: create content from product data
                    content_for_embedding = (
                        f"{product.get('name', '')}: {product.get('description', '')}"
                    )

                chunk_data = {
                    "chunk_id": chunk_id,
                    "company_id": company_id,
                    "file_id": document_id,
                    "content": json.dumps(product, ensure_ascii=False),
                    "content_for_embedding": content_for_embedding,
                    "content_type": "products",
                    "structured_data": product,
                    "language": (
                        language.value if hasattr(language, "value") else str(language)
                    ),
                    "industry": (
                        industry.value if hasattr(industry, "value") else str(industry)
                    ),
                    "location": {"country": "vietnam"},
                    "valid_from": None,
                    "valid_until": None,
                    "created_at": datetime.utcnow().isoformat(),
                    "updated_at": datetime.utcnow().isoformat(),
                }
                qdrant_chunks.append(chunk_data)

            # Process services from the correct data source
            services = data_source.get("services", [])
            logger.info(
                f"🔍 [QDRANT_PREP] Found {len(services)} services in data source"
            )

            for i, service in enumerate(services):
                chunk_id = str(uuid.uuid4())

                # Create optimized content for embedding
                content_for_embedding = service.get("content_for_embedding")
                if not content_for_embedding:
                    # Fallback: create content from service data
                    content_for_embedding = (
                        f"{service.get('name', '')}: {service.get('description', '')}"
                    )

                chunk_data = {
                    "chunk_id": chunk_id,
                    "company_id": company_id,
                    "file_id": document_id,
                    "content": json.dumps(service, ensure_ascii=False),
                    "content_for_embedding": content_for_embedding,
                    "content_type": "services",
                    "structured_data": service,
                    "language": (
                        language.value if hasattr(language, "value") else str(language)
                    ),
                    "industry": (
                        industry.value if hasattr(industry, "value") else str(industry)
                    ),
                    "location": {"country": "vietnam"},
                    "valid_from": None,
                    "valid_until": None,
                    "created_at": datetime.utcnow().isoformat(),
                    "updated_at": datetime.utcnow().isoformat(),
                }
                qdrant_chunks.append(chunk_data)

            ingestion_data = {
                "qdrant_chunks": qdrant_chunks,
                "ingestion_metadata": {
                    "total_chunks": len(qdrant_chunks),
                    "products_count": len(products),
                    "services_count": len(services),
                    "company_id": company_id,
                    "industry": (
                        industry.value if hasattr(industry, "value") else str(industry)
                    ),
                    "language": (
                        language.value if hasattr(language, "value") else str(language)
                    ),
                    "document_id": document_id,
                    "user_id": user_id,
                    "callback_url": callback_url,
                    "created_at": datetime.utcnow().isoformat(),
                },
            }

            logger.info(
                f"✅ [QDRANT_PREP] Prepared {len(qdrant_chunks)} chunks for Qdrant ingestion"
            )
            logger.info(f"   📦 Products: {len(products)}")
            logger.info(f"   🔧 Services: {len(services)}")

            return ingestion_data

        except Exception as e:
            logger.error(
                f"❌ [QDRANT_PREP] Failed to prepare for Qdrant ingestion: {str(e)}"
            )
            raise Exception(f"Failed to prepare for Qdrant ingestion: {str(e)}")

    async def ensure_qdrant_indexes(self, collection_name: str) -> bool:
        """
        Ensure required indexes exist in Qdrant collection
        Đảm bảo các indexes cần thiết tồn tại trong Qdrant collection

        Returns:
            bool: True if indexes are ready, False otherwise
        """
        try:
            from qdrant_client import QdrantClient
            from qdrant_client.models import PayloadSchemaType
            import os

            logger.info(f"🔧 Ensuring indexes for collection: {collection_name}")

            # Initialize Qdrant client
            qdrant_client = (
                QdrantClient(
                    url=os.getenv("QDRANT_URL"),
                    api_key=os.getenv("QDRANT_API_KEY"),
                )
                if os.getenv("QDRANT_URL")
                else None
            )

            if not qdrant_client:
                logger.error("❌ Qdrant client not available")
                return False

            # Check if collection exists
            collections = qdrant_client.get_collections().collections
            collection_exists = any(c.name == collection_name for c in collections)

            if not collection_exists:
                logger.warning(f"⚠️ Collection {collection_name} does not exist")
                return False

            # Create indexes if they don't exist
            indexes_created = []

            try:
                # Create file_id index
                qdrant_client.create_payload_index(
                    collection_name=collection_name,
                    field_name="file_id",
                    field_schema=PayloadSchemaType.KEYWORD,
                )
                indexes_created.append("file_id")
                logger.info(
                    f"✅ Created file_id index for collection: {collection_name}"
                )
            except Exception as e:
                if "already exists" in str(e).lower():
                    logger.info(
                        f"📋 file_id index already exists for collection: {collection_name}"
                    )
                else:
                    logger.warning(f"⚠️ Failed to create file_id index: {str(e)}")

            try:
                # Create company_id index
                qdrant_client.create_payload_index(
                    collection_name=collection_name,
                    field_name="company_id",
                    field_schema=PayloadSchemaType.KEYWORD,
                )
                indexes_created.append("company_id")
                logger.info(
                    f"✅ Created company_id index for collection: {collection_name}"
                )
            except Exception as e:
                if "already exists" in str(e).lower():
                    logger.info(
                        f"📋 company_id index already exists for collection: {collection_name}"
                    )
                else:
                    logger.warning(f"⚠️ Failed to create company_id index: {str(e)}")

            # Try to create additional useful indexes
            for field_name in ["content_type", "industry", "data_type"]:
                try:
                    qdrant_client.create_payload_index(
                        collection_name=collection_name,
                        field_name=field_name,
                        field_schema=PayloadSchemaType.KEYWORD,
                    )
                    indexes_created.append(field_name)
                    logger.info(
                        f"✅ Created {field_name} index for collection: {collection_name}"
                    )
                except Exception as e:
                    if "already exists" in str(e).lower():
                        logger.info(
                            f"📋 {field_name} index already exists for collection: {collection_name}"
                        )
                    else:
                        logger.warning(
                            f"⚠️ Failed to create {field_name} index: {str(e)}"
                        )

            if indexes_created:
                logger.info(
                    f"🔧 Indexes ready for collection {collection_name}: {indexes_created}"
                )
            else:
                logger.info(
                    f"📋 All indexes already exist for collection: {collection_name}"
                )

            return True

        except Exception as e:
            logger.error(
                f"❌ Failed to ensure indexes for collection {collection_name}: {str(e)}"
            )
            return False

    async def delete_file_from_qdrant(self, collection_name: str, file_id: str) -> int:
        """
        Delete all points associated with a file_id from Qdrant collection
        Xóa tất cả points liên quan đến file_id từ Qdrant collection

        Args:
            collection_name: Name of the Qdrant collection
            file_id: File ID to delete all associated points

        Returns:
            int: Number of points deleted
        """
        try:
            from qdrant_client import QdrantClient
            from qdrant_client.models import Filter, FieldCondition, MatchValue
            import os

            logger.info(
                f"🗑️ Deleting all points for file_id: {file_id} from collection: {collection_name}"
            )

            # Initialize Qdrant client
            qdrant_client = (
                QdrantClient(
                    url=os.getenv("QDRANT_URL"),
                    api_key=os.getenv("QDRANT_API_KEY"),
                )
                if os.getenv("QDRANT_URL")
                else None
            )

            if not qdrant_client:
                logger.error("❌ Qdrant client not available")
                return 0

            # Check if collection exists
            collections = qdrant_client.get_collections().collections
            collection_exists = any(c.name == collection_name for c in collections)

            if not collection_exists:
                logger.warning(f"⚠️ Collection {collection_name} does not exist")
                return 0

            # Ensure required indexes exist
            indexes_ready = await self.ensure_qdrant_indexes(collection_name)
            if not indexes_ready:
                logger.warning(
                    f"⚠️ Failed to ensure indexes for collection: {collection_name}"
                )
                # Continue anyway, might still work

            # First, count how many points will be deleted
            filter_condition = Filter(
                must=[
                    FieldCondition(
                        key="file_id",
                        match=MatchValue(value=file_id),
                    )
                ]
            )

            # Scroll to count points
            points_to_delete, _ = qdrant_client.scroll(
                collection_name=collection_name,
                scroll_filter=filter_condition,
                limit=10000,  # Large limit to get all points
                with_payload=False,
                with_vectors=False,
            )

            total_points = len(points_to_delete)
            logger.info(
                f"🔍 Found {total_points} points to delete for file_id: {file_id}"
            )

            if total_points == 0:
                logger.warning(f"📭 No points found for file_id: {file_id}")
                return 0

            # Delete points using filter
            result = qdrant_client.delete(
                collection_name=collection_name,
                points_selector=filter_condition,
                wait=True,
            )

            logger.info(
                f"✅ Successfully deleted {total_points} points for file_id: {file_id}"
            )
            logger.info(
                f"   🗑️ Delete operation ID: {getattr(result, 'operation_id', 'completed')}"
            )

            return total_points

        except Exception as e:
            logger.error(f"❌ Failed to delete file {file_id} from Qdrant: {str(e)}")
            return 0

    async def delete_file_with_company_filter(
        self, collection_name: str, company_id: str, file_id: str
    ) -> int:
        """
        Delete all points associated with company_id AND file_id from Qdrant collection
        Xóa tất cả points liên quan đến company_id VÀ file_id từ Qdrant collection

        This is used for additional cleanup in extraction deletion workflow
        Được sử dụng cho cleanup bổ sung trong workflow xóa extraction

        Args:
            collection_name: Name of the Qdrant collection
            company_id: Company ID filter
            file_id: File ID filter

        Returns:
            int: Number of additional points deleted
        """
        try:
            from qdrant_client import QdrantClient
            from qdrant_client.models import Filter, FieldCondition, MatchValue
            import os

            logger.info(
                f"🧹 Additional cleanup: deleting points for company_id: {company_id} AND file_id: {file_id}"
            )

            # Initialize Qdrant client
            qdrant_client = (
                QdrantClient(
                    url=os.getenv("QDRANT_URL"),
                    api_key=os.getenv("QDRANT_API_KEY"),
                )
                if os.getenv("QDRANT_URL")
                else None
            )

            if not qdrant_client:
                logger.error("❌ Qdrant client not available")
                return 0

            # Check if collection exists
            collections = qdrant_client.get_collections().collections
            collection_exists = any(c.name == collection_name for c in collections)

            if not collection_exists:
                logger.warning(f"⚠️ Collection {collection_name} does not exist")
                return 0

            # Ensure required indexes exist
            indexes_ready = await self.ensure_qdrant_indexes(collection_name)
            if not indexes_ready:
                logger.warning(
                    f"⚠️ Failed to ensure indexes for collection: {collection_name}"
                )
                # Continue anyway, might still work
                return 0

            # Check if collection exists
            collections = qdrant_client.get_collections().collections
            collection_exists = any(c.name == collection_name for c in collections)

            if not collection_exists:
                logger.warning(f"⚠️ Collection {collection_name} does not exist")
                return 0

            # Filter with both company_id AND file_id
            filter_condition = Filter(
                must=[
                    FieldCondition(
                        key="company_id",
                        match=MatchValue(value=company_id),
                    ),
                    FieldCondition(
                        key="file_id",
                        match=MatchValue(value=file_id),
                    ),
                ]
            )

            # Scroll to count points
            points_to_delete, _ = qdrant_client.scroll(
                collection_name=collection_name,
                scroll_filter=filter_condition,
                limit=10000,  # Large limit to get all points
                with_payload=False,
                with_vectors=False,
            )

            total_points = len(points_to_delete)
            logger.info(
                f"🔍 Found {total_points} additional points with company_id + file_id filter"
            )

            if total_points == 0:
                logger.info("✅ No additional points found with company+file filter")
                return 0

            # Delete points using filter
            result = qdrant_client.delete(
                collection_name=collection_name,
                points_selector=filter_condition,
                wait=True,
            )

            logger.info(
                f"✅ Additional cleanup: deleted {total_points} points with company+file filter"
            )
            logger.info(
                f"   🗑️ Delete operation ID: {getattr(result, 'operation_id', 'completed')}"
            )

            return total_points

        except Exception as e:
            logger.error(
                f"❌ Failed additional cleanup for company {company_id}, file {file_id}: {str(e)}"
            )
            return 0


# ===== SINGLETON INSTANCE =====

_ai_service_instance = None


def get_ai_service() -> AIExtractionService:
    """
    Get singleton instance of AIExtractionService.

    Returns:
        AIExtractionService: Singleton instance
    """
    global _ai_service_instance

    if _ai_service_instance is None:
        _ai_service_instance = AIExtractionService()

    return _ai_service_instance
