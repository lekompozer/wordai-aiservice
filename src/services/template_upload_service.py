"""
Template Upload Service
Xử lý upload và phân tích template DOCX với AI
"""

import os
import uuid
import aiofiles
from typing import Dict, Any, List, Optional
from datetime import datetime
from fastapi import UploadFile, HTTPException
from docx import Document
from pydantic import BaseModel, Field

from src.utils.logger import setup_logger
from src.services.ai_service import get_ai_service
from src.config.database import get_async_database

logger = setup_logger(__name__)


class PlaceholderInfo(BaseModel):
    """Thông tin chi tiết về placeholder"""

    type: str  # text, number, date, currency, boolean
    description: str
    default_value: Any = ""
    validation_rules: List[str] = []
    section: str
    auto_populate: bool = False
    calculation_formula: Optional[str] = None


class TemplateSection(BaseModel):
    """Thông tin về section trong template"""

    name: str
    description: str
    placeholders: List[str]
    order: int
    is_repeatable: bool = False
    required: bool = True


class TemplateAnalysisResult(BaseModel):
    """Kết quả phân tích template từ AI"""

    template_id: str
    placeholders: Dict[str, PlaceholderInfo]
    sections: List[TemplateSection]
    business_logic: Dict[str, Any]
    ai_analysis_score: float
    processing_time: float


class TemplateUploadService:
    """Service xử lý upload và phân tích template"""

    def __init__(self):
        self.ai_service = get_ai_service()

    async def _get_database(self):
        """Get database instance"""
        return await get_async_database()

    async def validate_docx_file(self, file: UploadFile) -> Dict[str, Any]:
        """Validate uploaded DOCX file"""
        errors = []
        warnings = []

        # Check file type
        if not file.filename.endswith(".docx"):
            errors.append("Only DOCX files are supported")

        # Check file size (max 10MB)
        if file.size and file.size > 10 * 1024 * 1024:
            errors.append("File size must be less than 10MB")

        # Check if file is valid DOCX
        try:
            content = await file.read()
            await file.seek(0)  # Reset file pointer

            # Try to open with python-docx
            from io import BytesIO

            doc = Document(BytesIO(content))

            # Basic content validation
            if len(doc.paragraphs) < 3:
                warnings.append("Template seems too short, might be missing content")

        except Exception as e:
            errors.append(f"Invalid DOCX file format: {str(e)}")

        return {"is_valid": len(errors) == 0, "errors": errors, "warnings": warnings}

    async def extract_text_content(self, file: UploadFile) -> str:
        """Extract text content from DOCX file"""
        try:
            content = await file.read()
            await file.seek(0)  # Reset file pointer

            from io import BytesIO

            doc = Document(BytesIO(content))

            text_content = []

            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())

            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))

            return "\n".join(text_content)

        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            raise HTTPException(
                status_code=400, detail=f"Cannot extract text from file: {str(e)}"
            )

    async def convert_to_pdf(self, docx_content: bytes, template_id: str) -> str:
        """Convert DOCX to PDF for preview"""
        try:
            # Import docx2pdf - handle if not available
            try:
                from docx2pdf import convert
            except ImportError:
                logger.warning("docx2pdf not available, skipping PDF conversion")
                return None
            import tempfile

            # Create temporary files
            with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as docx_temp:
                docx_temp.write(docx_content)
                docx_temp_path = docx_temp.name

            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as pdf_temp:
                pdf_temp_path = pdf_temp.name

            # Convert DOCX to PDF
            convert(docx_temp_path, pdf_temp_path)

            # Read PDF content
            with open(pdf_temp_path, "rb") as pdf_file:
                pdf_content = pdf_file.read()

            # Clean up temp files
            os.unlink(docx_temp_path)
            os.unlink(pdf_temp_path)

            # Upload PDF to R2 (placeholder - would need R2 service)
            pdf_url = await self._upload_to_r2(
                pdf_content, f"templates/{template_id}/preview.pdf"
            )

            return pdf_url

        except Exception as e:
            logger.error(f"Error converting DOCX to PDF: {e}")
            # Return None if conversion fails - not critical
            return None

    async def analyze_template_with_ai(
        self, text_content: str, template_id: str
    ) -> TemplateAnalysisResult:
        """Phân tích template với AI"""
        try:
            start_time = datetime.now()

            # Prepare AI prompt for template analysis
            analysis_prompt = f"""
            Phân tích template document sau đây và trả về JSON với cấu trúc sau:

            {{
                "placeholders": {{
                    "placeholder_name": {{
                        "type": "text|number|date|currency|boolean",
                        "description": "Mô tả field",
                        "default_value": "Giá trị mặc định",
                        "validation_rules": ["required", "email", "phone", ...],
                        "section": "company_info|customer_info|products|financial|terms",
                        "auto_populate": false,
                        "calculation_formula": "optional formula for calculated fields"
                    }}
                }},
                "sections": [
                    {{
                        "name": "section_name",
                        "description": "Mô tả section",
                        "placeholders": ["list of placeholders in this section"],
                        "order": 1,
                        "is_repeatable": false,
                        "required": true
                    }}
                ],
                "business_logic": {{
                    "calculation_fields": ["fields that need calculation"],
                    "auto_fill_fields": ["fields to auto-populate"],
                    "conditional_fields": ["fields with conditions"],
                    "validation_rules": {{"field_name": ["rules"]}}
                }}
            }}

            Tìm tất cả các placeholder patterns như: {{{{field_name}}}}, [field_name], __field_name__, {{field_name}}

            Template content:
            {text_content}

            Chỉ trả về JSON, không có text thêm.
            """

            # Call AI service
            ai_response = await self.ai_service.generate_response(
                prompt=analysis_prompt,
                model="gemini-2.5-flash",
                max_tokens=4000,
                temperature=0.1,
            )

            # Parse AI response
            import json

            try:
                analysis_data = json.loads(ai_response.strip())
            except json.JSONDecodeError:
                # Fallback to pattern-based analysis
                logger.warning(
                    "AI response is not valid JSON, falling back to pattern analysis"
                )
                analysis_data = await self._pattern_based_analysis(text_content)

            # Calculate processing time and confidence score
            processing_time = (datetime.now() - start_time).total_seconds()
            confidence_score = self._calculate_confidence_score(
                analysis_data, text_content
            )

            # Create TemplateAnalysisResult
            placeholders = {}
            for name, info in analysis_data.get("placeholders", {}).items():
                placeholders[name] = PlaceholderInfo(**info)

            sections = []
            for section_data in analysis_data.get("sections", []):
                sections.append(TemplateSection(**section_data))

            result = TemplateAnalysisResult(
                template_id=template_id,
                placeholders=placeholders,
                sections=sections,
                business_logic=analysis_data.get("business_logic", {}),
                ai_analysis_score=confidence_score,
                processing_time=processing_time,
            )

            return result

        except Exception as e:
            logger.error(f"Error in AI template analysis: {e}")
            # Fallback to basic pattern analysis
            return await self._pattern_based_analysis_fallback(
                text_content, template_id
            )

    async def _pattern_based_analysis(self, text_content: str) -> Dict[str, Any]:
        """Pattern-based analysis fallback"""
        import re

        # Find placeholder patterns
        patterns = [
            r"\{\{([^}]+)\}\}",  # {{placeholder}}
            r"\[([^\]]+)\]",  # [placeholder]
            r"\{([^}]+)\}",  # {placeholder}
            r"__([^_]+)__",  # __placeholder__
        ]

        placeholders = {}
        found_placeholders = set()

        for pattern in patterns:
            matches = re.findall(pattern, text_content)
            for match in matches:
                field_name = match.strip().lower()
                if field_name and field_name not in found_placeholders:
                    found_placeholders.add(field_name)

                    # Determine field type and section based on name
                    placeholder_info = self._infer_placeholder_info(field_name)
                    placeholders[f"{{{{{field_name}}}}}"] = placeholder_info

        # Create basic sections
        sections = [
            {
                "name": "header",
                "description": "Header section with company info",
                "placeholders": [
                    p for p in placeholders.keys() if "company" in p or "logo" in p
                ],
                "order": 1,
                "is_repeatable": False,
                "required": True,
            },
            {
                "name": "customer_info",
                "description": "Customer information section",
                "placeholders": [
                    p for p in placeholders.keys() if "customer" in p or "client" in p
                ],
                "order": 2,
                "is_repeatable": False,
                "required": True,
            },
            {
                "name": "products",
                "description": "Products/services section",
                "placeholders": [
                    p
                    for p in placeholders.keys()
                    if "product" in p or "service" in p or "item" in p
                ],
                "order": 3,
                "is_repeatable": True,
                "required": True,
            },
            {
                "name": "financial",
                "description": "Financial calculations",
                "placeholders": [
                    p
                    for p in placeholders.keys()
                    if "total" in p or "amount" in p or "price" in p
                ],
                "order": 4,
                "is_repeatable": False,
                "required": True,
            },
        ]

        return {
            "placeholders": placeholders,
            "sections": sections,
            "business_logic": {
                "calculation_fields": [
                    p for p in placeholders.keys() if "total" in p or "amount" in p
                ],
                "auto_fill_fields": [
                    p for p in placeholders.keys() if "date" in p or "number" in p
                ],
                "conditional_fields": [],
                "validation_rules": {},
            },
        }

    def _infer_placeholder_info(self, field_name: str) -> Dict[str, Any]:
        """Infer placeholder information from field name"""
        field_name_lower = field_name.lower()

        # Determine type
        if "date" in field_name_lower:
            field_type = "date"
        elif (
            "amount" in field_name_lower
            or "price" in field_name_lower
            or "total" in field_name_lower
        ):
            field_type = "currency"
        elif "quantity" in field_name_lower or "count" in field_name_lower:
            field_type = "number"
        elif "email" in field_name_lower:
            field_type = "email"
        elif "phone" in field_name_lower:
            field_type = "phone"
        else:
            field_type = "text"

        # Determine section
        if "company" in field_name_lower:
            section = "company_info"
        elif "customer" in field_name_lower or "client" in field_name_lower:
            section = "customer_info"
        elif (
            "product" in field_name_lower
            or "service" in field_name_lower
            or "item" in field_name_lower
        ):
            section = "products"
        elif (
            "total" in field_name_lower
            or "amount" in field_name_lower
            or "price" in field_name_lower
        ):
            section = "financial"
        else:
            section = "general"

        # Set validation rules
        validation_rules = []
        if field_type == "email":
            validation_rules.append("email")
        elif field_type == "phone":
            validation_rules.append("phone")
        elif field_type in ["currency", "number"]:
            validation_rules.append("numeric")

        return {
            "type": field_type,
            "description": f"Field for {field_name.replace('_', ' ')}",
            "default_value": "",
            "validation_rules": validation_rules,
            "section": section,
            "auto_populate": "date" in field_name_lower or "number" in field_name_lower,
            "calculation_formula": None,
        }

    async def _pattern_based_analysis_fallback(
        self, text_content: str, template_id: str
    ) -> TemplateAnalysisResult:
        """Fallback pattern analysis when AI fails"""
        analysis_data = await self._pattern_based_analysis(text_content)

        placeholders = {}
        for name, info in analysis_data["placeholders"].items():
            placeholders[name] = PlaceholderInfo(**info)

        sections = []
        for section_data in analysis_data["sections"]:
            sections.append(TemplateSection(**section_data))

        return TemplateAnalysisResult(
            template_id=template_id,
            placeholders=placeholders,
            sections=sections,
            business_logic=analysis_data["business_logic"],
            ai_analysis_score=0.7,  # Lower confidence for fallback
            processing_time=0.5,
        )

    def _calculate_confidence_score(
        self, analysis_data: Dict[str, Any], text_content: str
    ) -> float:
        """Calculate confidence score for AI analysis"""
        score = 0.0

        # Check if placeholders were found
        placeholders = analysis_data.get("placeholders", {})
        if len(placeholders) > 0:
            score += 0.3

        # Check if sections were identified
        sections = analysis_data.get("sections", [])
        if len(sections) > 0:
            score += 0.3

        # Check if business logic was identified
        business_logic = analysis_data.get("business_logic", {})
        if business_logic:
            score += 0.2

        # Check if placeholder types are reasonable
        for placeholder in placeholders.values():
            if placeholder.get("type") in [
                "text",
                "number",
                "date",
                "currency",
                "boolean",
            ]:
                score += 0.02

        # Maximum score is 1.0
        return min(score, 1.0)

    async def _upload_to_r2(self, content: bytes, key: str) -> str:
        """Upload file to R2 storage - placeholder implementation"""
        # TODO: Implement actual R2 upload
        # For now, return a placeholder URL
        return f"https://r2.dev/{key}"

    async def save_template_to_database(
        self,
        template_id: str,
        user_id: str,
        template_name: str,
        description: str,
        category: str,
        docx_url: str,
        pdf_url: Optional[str],
        analysis_result: TemplateAnalysisResult,
    ) -> Dict[str, Any]:
        """Save template metadata to database"""
        try:
            db = await self._get_database()

            template_doc = {
                "_id": template_id,
                "user_id": user_id,
                "name": template_name,
                "description": description,
                "category": category,
                "type": "quote",
                "subtype": "business",
                # File information
                "files": {
                    "docx_url": docx_url,
                    "pdf_url": pdf_url,
                    "thumbnail_urls": [],
                },
                # AI Analysis results
                "ai_analysis": {
                    "placeholders": {
                        name: {
                            "type": info.type,
                            "description": info.description,
                            "default_value": info.default_value,
                            "validation_rules": info.validation_rules,
                            "section": info.section,
                            "auto_populate": info.auto_populate,
                            "calculation_formula": info.calculation_formula,
                        }
                        for name, info in analysis_result.placeholders.items()
                    },
                    "sections": [
                        {
                            "name": section.name,
                            "description": section.description,
                            "placeholders": section.placeholders,
                            "order": section.order,
                            "is_repeatable": section.is_repeatable,
                            "required": section.required,
                        }
                        for section in analysis_result.sections
                    ],
                    "business_logic": analysis_result.business_logic,
                    "confidence_score": analysis_result.ai_analysis_score,
                    "analysis_version": "1.0",
                },
                # Metadata
                "is_active": True,
                "is_public": False,
                "usage_count": 0,
                "created_at": datetime.utcnow(),
                "updated_at": datetime.utcnow(),
                # Template validation
                "validation": {"is_valid": True, "errors": [], "warnings": []},
            }

            result = await db.document_templates.insert_one(template_doc)

            logger.info(f"✅ Template saved to database: {template_id}")

            return {
                "success": True,
                "template_id": template_id,
                "database_id": str(result.inserted_id),
            }

        except Exception as e:
            logger.error(f"Error saving template to database: {e}")
            raise HTTPException(status_code=500, detail=f"Database error: {str(e)}")

    async def process_template_upload(
        self,
        file: UploadFile,
        template_name: str,
        description: str,
        category: str,
        user_id: str,
    ) -> Dict[str, Any]:
        """Main method to process template upload"""
        try:
            # Generate template ID
            template_id = f"template_{uuid.uuid4().hex[:12]}"

            logger.info(f"🔄 Processing template upload: {template_id}")

            # Step 1: Validate file
            validation_result = await self.validate_docx_file(file)
            if not validation_result["is_valid"]:
                raise HTTPException(
                    status_code=400,
                    detail=f"File validation failed: {validation_result['errors']}",
                )

            # Step 2: Extract text content
            text_content = await self.extract_text_content(file)
            logger.info(f"📄 Extracted {len(text_content)} characters from template")

            # Step 3: Upload DOCX to R2
            docx_content = await file.read()
            await file.seek(0)
            docx_url = await self._upload_to_r2(
                docx_content, f"templates/{user_id}/{template_id}/original.docx"
            )

            # Step 4: Convert to PDF (optional)
            pdf_url = await self.convert_to_pdf(docx_content, template_id)

            # Step 5: AI Analysis
            analysis_result = await self.analyze_template_with_ai(
                text_content, template_id
            )
            logger.info(
                f"🧠 AI analysis completed with score: {analysis_result.ai_analysis_score}"
            )

            # Step 6: Save to database
            db_result = await self.save_template_to_database(
                template_id=template_id,
                user_id=user_id,
                template_name=template_name,
                description=description,
                category=category,
                docx_url=docx_url,
                pdf_url=pdf_url,
                analysis_result=analysis_result,
            )

            return {
                "success": True,
                "template_id": template_id,
                "status": "completed",
                "message": "Template uploaded and analyzed successfully",
                "urls": {
                    "docx": docx_url,
                    "pdf_preview": pdf_url,
                    "status_check": f"/api/templates/{template_id}/status",
                },
                "analysis_summary": {
                    "placeholders_found": len(analysis_result.placeholders),
                    "sections_identified": len(analysis_result.sections),
                    "confidence_score": analysis_result.ai_analysis_score,
                    "processing_time": analysis_result.processing_time,
                },
                "warnings": validation_result.get("warnings", []),
            }

        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"❌ Template upload processing failed: {e}")
            raise HTTPException(status_code=500, detail=f"Processing failed: {str(e)}")
