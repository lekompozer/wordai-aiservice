"""
Document Processing Worker for handling DocumentProcessingTask from Redis queue.
Handles file uploads, AI extraction, and Qdrant upload in the background.
"""

import os
import sys
import asyncio
import logging
import signal
import time
import traceback
import json
import uuid
import re
import aiohttp
import redis
import io
from typing import Optional, List
from datetime import datetime
from dotenv import load_dotenv

# Document processing imports
try:
    from docx import Document
except ImportError:
    Document = None

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# Add src to path to enable imports
sys.path.append(os.path.join(os.path.dirname(__file__), "..", ".."))

# Load environment variables
env_var = os.getenv("ENVIRONMENT", os.getenv("ENV", "production"))
env_file = "development.env" if env_var == "development" else ".env"
load_dotenv(env_file)

from src.queue.queue_dependencies import get_document_queue
from src.queue.task_models import DocumentProcessingTask
from src.services.ai_extraction_service import get_ai_service
from src.services.qdrant_company_service import QdrantCompanyDataService
from src.providers.ai_provider_manager import AIProviderManager
from src.models.unified_models import Industry, Language
from src.utils.logger import setup_logger
import config.config as config

# ==============================================================================
# WORKER CONFIGURATION
# ==============================================================================

logger = setup_logger(__name__)


class DocumentProcessingWorker:
    """
    Worker that processes DocumentProcessingTask from Redis queue.
    Handles file extraction and Qdrant upload in background.
    """

    def __init__(
        self,
        worker_id: str = None,
        redis_url: str = None,
        batch_size: int = 1,
        max_retries: int = 3,
        poll_interval: float = 1.0,
    ):
        self.worker_id = worker_id or f"doc_worker_{int(time.time())}_{os.getpid()}"
        self.redis_url = redis_url or os.getenv(
            "REDIS_URL", "redis://redis-server:6379"
        )
        self.batch_size = batch_size
        self.max_retries = max_retries
        self.poll_interval = poll_interval
        self.running = False

        # Initialize components
        self.queue_manager = None
        self.ai_service = None
        self.qdrant_service = None
        self.ai_manager = None

        # Webhook configuration
        self.webhook_secret = os.getenv(
            "WEBHOOK_SECRET", "webhook-secret-for-signature"
        )

        logger.info(f"🔧 Document Worker {self.worker_id} initialized")
        logger.info(f"   📡 Redis: {self.redis_url}")
        logger.info(f"   📦 Batch size: {self.batch_size}")
        logger.info(f"   ⏱️ Poll interval: {self.poll_interval}s")

    async def initialize(self):
        """Initialize worker components with Redis retry logic"""
        max_retries = 5
        retry_delay = 2

        for attempt in range(max_retries):
            try:
                # Get queue manager with retry logic
                logger.info(
                    f"🔄 Worker {self.worker_id}: Connecting to Redis (attempt {attempt + 1}/{max_retries})"
                )
                self.queue_manager = await get_document_queue()
                logger.info(f"✅ Worker {self.worker_id}: Connected to Redis queue")
                break

            except Exception as redis_error:
                logger.error(
                    f"❌ Worker {self.worker_id}: Redis connection failed (attempt {attempt + 1}): {redis_error}"
                )

                if attempt == max_retries - 1:
                    logger.error(
                        f"💀 Worker {self.worker_id}: Failed to connect to Redis after {max_retries} attempts"
                    )
                    raise redis_error

                logger.info(f"⏳ Retrying in {retry_delay}s...")
                await asyncio.sleep(retry_delay)
                retry_delay = min(30, retry_delay * 2)  # Exponential backoff

        try:
            # Initialize AI service
            self.ai_service = get_ai_service()
            logger.info(f"✅ Worker {self.worker_id}: AI service ready")

            # Initialize AI provider manager for embeddings
            self.ai_manager = AIProviderManager(
                deepseek_api_key=os.getenv("DEEPSEEK_API_KEY", ""),
                chatgpt_api_key=os.getenv("CHATGPT_API_KEY", ""),
                gemini_api_key=os.getenv("GEMINI_API_KEY", ""),
            )
            logger.info(f"✅ Worker {self.worker_id}: AI provider manager ready")

            # Initialize Qdrant service
            self.qdrant_service = QdrantCompanyDataService(
                qdrant_url=os.getenv("QDRANT_URL"),
                qdrant_api_key=os.getenv("QDRANT_API_KEY"),
            )
            logger.info(f"✅ Worker {self.worker_id}: Qdrant service ready")

        except Exception as e:
            logger.error(f"❌ Worker {self.worker_id}: Initialization failed: {e}")
            raise

    async def shutdown(self):
        """Gracefully shutdown worker"""
        logger.info(f"🛑 Worker {self.worker_id}: Shutting down...")
        self.running = False
        logger.info(f"✅ Worker {self.worker_id}: Shutdown complete")

    async def run(self):
        """Main worker loop - continuously poll for tasks from Redis queue"""
        logger.info(f"🚀 Starting document processing worker {self.worker_id}")
        self.running = True
        consecutive_redis_errors = 0
        max_consecutive_errors = 10

        while self.running:
            try:
                # Dequeue single task from Redis
                task_data = await self.queue_manager.dequeue_generic_task(
                    worker_id=self.worker_id, timeout=2  # 2 second timeout for polling
                )

                # Reset error counter on successful Redis operation
                consecutive_redis_errors = 0

                if task_data:
                    logger.info(
                        f"📋 Processing document task: {task_data.get('task_id')}"
                    )

                    try:
                        # DEBUG: Log task data structure
                        logger.info(
                            f"🔍 [WORKER] Task data keys: {list(task_data.keys())}"
                        )

                        # DocumentProcessingWorker ONLY handles DocumentProcessingTask
                        if "processing_metadata" in task_data:
                            # This is an ExtractionProcessingTask - skip it, let ExtractionProcessingWorker handle it
                            logger.info(
                                f"🎯 [WORKER] Detected ExtractionProcessingTask - skipping (ExtractionProcessingWorker will handle)"
                            )
                            continue
                        else:
                            # This is a DocumentProcessingTask - we handle this
                            logger.info(
                                f"📄 [WORKER] Detected DocumentProcessingTask - calling process_task()"
                            )
                            task = DocumentProcessingTask(**task_data)
                            await self.process_task(task)

                    except Exception as task_error:
                        logger.error(f"❌ Task processing error: {task_error}")
                        logger.error(f"🔍 Task data: {task_data}")

                else:
                    # No tasks, wait briefly before next poll
                    await asyncio.sleep(0.5)

            except redis.exceptions.ConnectionError as redis_error:
                consecutive_redis_errors += 1
                logger.error(
                    f"🔌 Redis connection error (attempt {consecutive_redis_errors}): {redis_error}"
                )

                if consecutive_redis_errors >= max_consecutive_errors:
                    logger.error(
                        f"💀 Too many consecutive Redis errors ({consecutive_redis_errors}), stopping worker"
                    )
                    break

                # Exponential backoff for Redis connection issues
                backoff_time = min(30, 2**consecutive_redis_errors)
                logger.info(f"⏳ Waiting {backoff_time}s before retry...")
                await asyncio.sleep(backoff_time)

                # Try to reinitialize queue manager
                try:
                    logger.info("🔄 Attempting to reinitialize Redis connection...")
                    self.queue_manager = await get_document_queue()
                    logger.info("✅ Redis connection reestablished")
                except Exception as reconnect_error:
                    logger.error(f"❌ Failed to reconnect to Redis: {reconnect_error}")

            except Exception as e:
                consecutive_redis_errors += 1
                logger.error(f"❌ Worker loop error: {e}")
                logger.error(f"🔍 Traceback: {traceback.format_exc()}")

                if consecutive_redis_errors >= max_consecutive_errors:
                    logger.error(
                        f"💀 Too many consecutive errors ({consecutive_redis_errors}), stopping worker"
                    )
                    break

                await asyncio.sleep(self.poll_interval)

        logger.info(f"🛑 Document processing worker {self.worker_id} stopped")

    async def process_task(self, task: DocumentProcessingTask) -> bool:
        """
        Process a single DocumentProcessingTask.
        Now supports both RAW content extraction and STRUCTURED extraction modes.

        Args:
            task: DocumentProcessingTask instance

        Returns:
            bool: True if successful, False otherwise
        """
        start_time = time.time()

        try:
            logger.info(f"🔄 Processing task {task.task_id}")
            logger.info(f"   🏢 Company: {task.company_id}")
            logger.info(f"   🔗 R2 URL: {task.r2_url}")
            logger.info(f"   📁 Data type: {task.data_type}")
            logger.info(f"   🏭 Industry: {task.industry}")

            # Check if this is an extraction mode task
            extraction_mode = task.metadata.get("extraction_mode", False)
            processing_type = task.metadata.get("processing_type", "document_workflow")

            if extraction_mode and processing_type == "extraction_workflow":
                logger.info(
                    "🎯 EXTRACTION MODE: Processing with AI templates and structured data"
                )
                return await self._process_extraction_task(task, start_time)
            else:
                logger.info("📄 DOCUMENT MODE: Processing with raw content only")
                return await self._process_document_task(task, start_time)

        except Exception as e:
            processing_time = time.time() - start_time
            logger.error(f"❌ Task {task.task_id} failed: {str(e)}")
            logger.error(f"❌ Processing time: {processing_time:.2f}s")

            # Send error callback
            await self._send_error_callback(task, str(e))
            return False

    async def _process_extraction_task(
        self, task: DocumentProcessingTask, start_time: float
    ) -> bool:
        """
        Process extraction task with AI templates and structured data.
        Uses ai_extraction_service for template-based extraction.
        """
        logger.info(f"🎯 Starting extraction workflow for task {task.task_id}")

        try:
            # Extract metadata for AI service
            company_id = task.metadata.get("company_id")
            industry_str = task.metadata.get("industry", "other")
            language_str = task.metadata.get("language", "vi")
            data_type = task.metadata.get("data_type", "auto")
            target_categories = task.metadata.get(
                "target_categories", ["products", "services"]
            )
            company_info = task.metadata.get("company_info")

            # Convert string enums to proper types
            try:
                industry_enum = Industry(industry_str)
            except ValueError:
                logger.warning(f"⚠️ Unknown industry '{industry_str}', using OTHER")
                industry_enum = Industry.OTHER

            try:
                language_enum = Language(language_str)
            except ValueError:
                logger.warning(
                    f"⚠️ Unknown language '{language_str}', using AUTO_DETECT"
                )
                language_enum = Language.AUTO_DETECT

            # Prepare metadata for AI service
            metadata = {
                "original_name": task.metadata.get("filename", "unknown"),
                "file_size": task.metadata.get("file_size", 0),
                "file_type": task.metadata.get("file_type", "unknown"),
                "industry": industry_enum,
                "language": language_enum,
                "extraction_timestamp": datetime.now().isoformat(),
                "task_id": task.task_id,
                "company_id": company_id,
                "data_type": data_type,
            }

            logger.info(f"🤖 Calling AI extraction service...")
            logger.info(f"   🏭 Industry: {industry_enum}")
            logger.info(f"   🌐 Language: {language_enum}")
            logger.info(f"   📋 Data Type: {data_type}")
            logger.info(f"   📊 Target Categories: {target_categories}")

            # Call AI extraction service with templates
            ai_service = get_ai_service()
            extraction_result = await ai_service.extract_from_r2_url(
                r2_url=task.r2_url,
                metadata=metadata,
                company_info=company_info,
                target_categories=target_categories,
            )

            logger.info(f"✅ AI extraction completed")
            logger.info(
                f"   📦 Products found: {len(extraction_result.get('products', []))}"
            )
            logger.info(
                f"   🔧 Services found: {len(extraction_result.get('services', []))}"
            )

            # Debug extraction_result structure
            logger.info(
                f"🔍 [DEBUG] extraction_result keys: {list(extraction_result.keys())}"
            )
            if "structured_data" in extraction_result:
                structured_data = extraction_result["structured_data"]
                logger.info(
                    f"🔍 [DEBUG] structured_data keys: {list(structured_data.keys())}"
                )
                logger.info(
                    f"🔍 [DEBUG] structured_data products count: {len(structured_data.get('products', []))}"
                )
                logger.info(
                    f"🔍 [DEBUG] structured_data services count: {len(structured_data.get('services', []))}"
                )
            else:
                logger.warning(
                    "⚠️ [DEBUG] No 'structured_data' key in extraction_result!"
                )

            # Prepare for Qdrant ingestion with optimized chunking
            logger.info(f"📤 Preparing for Qdrant ingestion with optimized chunking...")

            ingestion_data = await ai_service.prepare_for_qdrant_ingestion(
                extraction_result=extraction_result,
                user_id="system",
                document_id=task.task_id,
                company_id=company_id,
                industry=industry_enum,
                language=language_enum,
                callback_url=task.callback_url,
            )

            # Upload to Qdrant using existing logic
            await self._upload_extraction_to_qdrant(ingestion_data, task)

            # Send success callback
            processing_time = time.time() - start_time
            await self.send_callback(
                task,
                "completed",
                success=True,
                processing_time=processing_time,
                extraction_result=extraction_result,
            )

            # ✅ SAVE EXTRACTION RESULT TO FILE FOR DEBUG
            import json
            import os
            from datetime import datetime

            debug_dir = "debug_extraction_results"
            os.makedirs(debug_dir, exist_ok=True)

            # Save full extraction result
            debug_data = {
                "task_id": task.task_id,
                "timestamp": datetime.now().isoformat(),
                "company_id": task.company_id,
                "industry": task.industry,
                "processing_time": processing_time,
                "full_extraction_result": extraction_result,
                "extraction_result_keys": list(extraction_result.keys()),
                "products_count": len(extraction_result.get("products", [])),
                "services_count": len(extraction_result.get("services", [])),
                "has_structured_data_key": "structured_data" in extraction_result,
                "raw_content_length": len(extraction_result.get("raw_content", "")),
                "metadata": metadata,
            }

            debug_filename = f"extraction_result_{task.task_id}.json"
            debug_filepath = os.path.join(debug_dir, debug_filename)

            try:
                with open(debug_filepath, "w", encoding="utf-8") as f:
                    json.dump(debug_data, f, indent=2, ensure_ascii=False)
                logger.info(
                    f"💾 [DEBUG] Saved full extraction result: {debug_filepath}"
                )
            except Exception as e:
                logger.error(f"❌ [DEBUG] Failed to save extraction result: {e}")

            logger.info(f"✅ Extraction task {task.task_id} completed successfully")
            logger.info(f"   ⏱️ Processing time: {processing_time:.2f}s")

            return True

        except Exception as e:
            logger.error(f"❌ Extraction task {task.task_id} failed: {str(e)}")
            raise

    async def _process_document_task(
        self, task: DocumentProcessingTask, start_time: float
    ) -> bool:
        """
        Process regular document task with raw content only (original logic).
        """
        logger.info(f"📄 Starting document workflow for task {task.task_id}")

        try:
            # Extract RAW CONTENT from R2 URL using appropriate AI Provider
            logger.info(f"🤖 Extracting RAW CONTENT ONLY (no template extraction)...")

            # Select AI Provider based on file type and extract raw content
            raw_content, ai_provider_used = await self.extract_raw_content_from_r2(
                r2_url=task.r2_url,
                file_metadata=task.metadata,
                data_type=task.data_type,
            )

            extraction_result = {
                "raw_content": raw_content,
                "extraction_metadata": {
                    "ai_provider": ai_provider_used,
                    "extraction_type": "raw_text_only",
                    "file_name": task.metadata.get("original_name", "unknown"),
                    "file_size": len(raw_content) if raw_content else 0,
                    "extraction_timestamp": datetime.now().isoformat(),
                },
            }

            logger.info(f"✅ Raw content extraction completed")
            logger.info(
                f"   📄 Content length: {len(extraction_result.get('raw_content', ''))}"
            )
            logger.info(f"   🤖 Method: Simple text extraction (no template)")
            logger.info(f"   📄 File: {task.metadata.get('original_name', 'unknown')}")

            # If upload_to_qdrant is enabled, upload to Qdrant
            if task.upload_to_qdrant:
                await self.upload_to_qdrant(
                    task=task, extraction_result=extraction_result
                )

            processing_time = time.time() - start_time
            logger.info(f"✅ Task {task.task_id} completed in {processing_time:.2f}s")

            # Update task status to completed
            try:
                await self.queue_manager.update_task_status(
                    task.task_id,
                    "completed",
                    worker_id=self.worker_id,
                    error_message=None,
                )
                logger.info(f"✅ Task status updated to completed: {task.task_id}")
            except Exception as status_error:
                logger.warning(f"⚠️ Failed to update task status: {status_error}")

            # Send callback if provided
            if task.callback_url:
                await self.send_callback(
                    task,
                    "completed",
                    success=True,
                    error_message=None,
                    processing_time=processing_time,
                    extraction_result=extraction_result,
                )

            return True

        except Exception as e:
            logger.error(f"❌ Document task {task.task_id} failed: {str(e)}")
            raise

    async def extract_raw_content_from_r2(
        self,
        r2_url: str,
        file_metadata: dict,
        data_type: str,
    ) -> tuple[str, str]:
        """
        Extract raw text content from R2 URL without AI processing.
        Used for simple file upload workflow.

        Args:
            r2_url: R2 URL to download file from
            file_metadata: File metadata containing original_name, file_type, etc.
            data_type: Data type of the file

        Returns:
            Tuple of (raw_content, ai_provider_used)
        """
        try:
            logger.info(f"📥 Downloading file from R2: {r2_url}")

            # Download file content
            file_content = await self._download_from_r2(r2_url)

            # Get file extension from metadata
            original_name = file_metadata.get("original_name", "unknown")
            file_type = file_metadata.get("file_type", "unknown")

            # Determine file extension
            if "." in original_name:
                file_extension = "." + original_name.split(".")[-1].lower()
            elif "docx" in file_type:
                file_extension = ".docx"
            elif "pdf" in file_type:
                file_extension = ".pdf"
            elif "txt" in file_type:
                file_extension = ".txt"
            else:
                file_extension = ".txt"  # Default fallback

            logger.info(f"📄 File extension determined: {file_extension}")
            logger.info(f"📁 File type: {file_type}")
            logger.info(f"📋 Original name: {original_name}")

            # Extract text content based on file type
            raw_content = await self._extract_text_from_file(
                file_content, file_extension, file_metadata
            )

            logger.info(f"✅ Raw content extracted: {len(raw_content)} characters")

            # Return content and provider used (simple text extraction)
            return raw_content, "simple_text_extraction"

        except Exception as e:
            logger.error(f"❌ Raw content extraction failed: {str(e)}")
            raise Exception(f"Raw content extraction failed: {str(e)}")

    async def _download_from_r2(self, r2_url: str) -> bytes:
        """Download file content from R2 URL"""
        try:
            logger.info(f"🌐 [DOWNLOAD] Starting download from: {r2_url}")
            async with aiohttp.ClientSession() as session:
                async with session.get(r2_url) as response:
                    logger.info(f"📡 [DOWNLOAD] Response status: {response.status}")
                    if response.status == 200:
                        content = await response.read()
                        logger.info(f"✅ [DOWNLOAD] Downloaded {len(content)} bytes")
                        return content
                    else:
                        raise Exception(
                            f"Failed to download from R2: HTTP {response.status}"
                        )
        except Exception as e:
            logger.error(f"❌ R2 download failed: {str(e)}")
            raise Exception(f"R2 download failed: {str(e)}")

    async def _extract_text_from_file(
        self, file_content: bytes, file_extension: str, metadata: dict
    ) -> str:
        """
        Extract text content from different file types.
        Simplified version for document upload workflow.
        """
        try:
            file_name = metadata.get("original_name", "unknown")
            logger.info(f"📄 Extracting text from {file_extension} file: {file_name}")

            if file_extension == ".txt":
                # Plain text file
                text_content = file_content.decode("utf-8")
                logger.info(f"✅ TXT file decoded successfully")
                return text_content

            elif file_extension == ".docx":
                # DOCX file - extract text using python-docx
                try:
                    if Document is None:
                        logger.warning(
                            "⚠️ python-docx not installed, cannot extract DOCX content"
                        )
                        return (
                            "DOCX file detected but python-docx library not available"
                        )

                    doc_stream = io.BytesIO(file_content)
                    doc = Document(doc_stream)

                    # Extract all text from paragraphs
                    text_lines = []
                    for paragraph in doc.paragraphs:
                        if paragraph.text.strip():
                            text_lines.append(paragraph.text.strip())

                    # Extract text from tables
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                if cell.text.strip():
                                    text_lines.append(cell.text.strip())

                    text_content = "\n".join(text_lines)
                    logger.info(
                        f"✅ DOCX file processed: {len(text_lines)} text blocks"
                    )
                    return text_content

                except Exception as docx_error:
                    logger.error(f"❌ DOCX extraction failed: {docx_error}")
                    # Fallback to binary content as text (not ideal but prevents crash)
                    return f"DOCX extraction failed: {str(docx_error)}"

            elif file_extension == ".pdf":
                # PDF file - simple text extraction
                try:
                    if PyPDF2 is None:
                        logger.warning(
                            "⚠️ PyPDF2 not installed, cannot extract PDF content"
                        )
                        return "PDF file detected but PyPDF2 library not available"

                    pdf_stream = io.BytesIO(file_content)
                    pdf_reader = PyPDF2.PdfReader(pdf_stream)

                    text_lines = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_lines.append(f"=== Page {page_num + 1} ===")
                            text_lines.append(page_text.strip())

                    text_content = "\n".join(text_lines)
                    logger.info(f"✅ PDF file processed: {len(pdf_reader.pages)} pages")
                    return text_content

                except Exception as pdf_error:
                    logger.error(f"❌ PDF extraction failed: {pdf_error}")
                    return f"PDF extraction failed: {str(pdf_error)}"

            else:
                # Unknown file type - try to decode as text
                try:
                    text_content = file_content.decode("utf-8")
                    logger.info(f"✅ File decoded as UTF-8 text")
                    return text_content
                except UnicodeDecodeError:
                    logger.warning(f"⚠️ Cannot decode as UTF-8, trying latin-1")
                    try:
                        text_content = file_content.decode("latin-1")
                        return text_content
                    except Exception:
                        logger.error(f"❌ Cannot decode file as text")
                        return f"Binary file - cannot extract text content (file type: {file_extension})"

        except Exception as e:
            logger.error(f"❌ Text extraction failed: {str(e)}")
            return f"Text extraction failed: {str(e)}"

    # ===== EXTRACTION PROCESSING REMOVED =====
    # ExtractionProcessingTask is now handled by ExtractionProcessingWorker
    # DocumentProcessingWorker only handles DocumentProcessingTask

    async def send_callback(
        self,
        task: DocumentProcessingTask,
        status: str,
        success: bool = True,
        error_message: Optional[str] = None,
        processing_time: Optional[float] = None,
        extraction_result: Optional[dict] = None,
    ):
        """
        Send callback for /upload workflow - RAW CONTENT ONLY (no structured data)
        Gửi callback cho luồng /upload - CHỈ RAW CONTENT (không có structured data)
        """
        if not task.callback_url:
            return

        try:
            # Prepare callback data for /upload workflow (raw content only)
            callback_data = {
                "task_id": task.task_id,
                "company_id": task.company_id,
                "status": status,
                "success": success,
                "processing_time": processing_time,
                "timestamp": datetime.now().isoformat(),
            }

            # Add raw content if available (for /upload workflow)
            if extraction_result and "raw_content" in extraction_result:
                raw_content = extraction_result.get("raw_content", "")
                extraction_metadata = extraction_result.get("extraction_metadata", {})

                callback_data.update(
                    {
                        # ✅ RAW CONTENT ONLY - No structured data for /upload workflow
                        "raw_content": raw_content,
                        "file_processing": {
                            "ai_provider": extraction_metadata.get("ai_provider"),
                            "extraction_type": extraction_metadata.get(
                                "extraction_type", "raw_text_only"
                            ),
                            "file_name": extraction_metadata.get("file_name"),
                            "content_length": len(raw_content) if raw_content else 0,
                            "processing_method": "document_upload_workflow",
                        },
                        "file_metadata": {
                            "original_name": task.metadata.get("original_name"),
                            "file_size": task.metadata.get("file_size"),
                            "file_type": task.metadata.get("file_type"),
                            "data_type": task.data_type,
                            "industry": str(task.industry) if task.industry else None,
                            "language": str(task.language) if task.language else None,
                        },
                    }
                )

                logger.info(
                    f"📞 [UPLOAD] Sending callback with RAW CONTENT for {task.task_id}"
                )
                logger.info(f"   📄 Raw content length: {len(raw_content)} characters")
                logger.info(
                    f"   🤖 AI Provider: {extraction_metadata.get('ai_provider')}"
                )
                logger.info(
                    f"   📁 File: {task.metadata.get('original_name', 'unknown')}"
                )
            else:
                # No extraction result (e.g., error case)
                if error_message:
                    callback_data.update(
                        {
                            "error": error_message,
                            "error_details": {
                                "workflow": "document_upload",
                                "step": "content_extraction",
                                "timestamp": datetime.now().isoformat(),
                            },
                        }
                    )
                    logger.info(
                        f"📞 [UPLOAD] Sending ERROR callback for {task.task_id}"
                    )
                    logger.info(f"   ❌ Error: {error_message}")

            # Build final callback payload
            callback_payload = {
                "task_id": task.task_id,
                "status": status,
                "timestamp": datetime.now().isoformat(),
                "data": callback_data,
            }

            logger.info(
                f"📞 Sending callback for {task.task_id} to {task.callback_url}"
            )

            # Generate webhook signature
            webhook_secret = os.getenv("WEBHOOK_SECRET", "webhook-secret-for-signature")

            headers = {
                "Content-Type": "application/json",
                "X-Webhook-Source": "ai-service",
                "X-Webhook-Secret": webhook_secret,  # ✅ Simplified: Use plain text secret
                "User-Agent": "Agent8x-AI-Service/1.0",
            }

            # Send async HTTP request
            timeout = aiohttp.ClientTimeout(total=10)  # 10 second timeout

            async with aiohttp.ClientSession(timeout=timeout) as session:
                async with session.post(
                    task.callback_url,
                    json=callback_payload,
                    headers=headers,
                ) as response:
                    if response.status == 200:
                        logger.info(f"✅ Callback sent successfully for {task.task_id}")
                    else:
                        logger.warning(
                            f"⚠️ Callback returned status {response.status} for {task.task_id}"
                        )

        except Exception as e:
            logger.error(f"❌ Failed to send callback for {task.task_id}: {e}")

    async def _send_error_callback(
        self, task: DocumentProcessingTask, error_message: str
    ):
        """
        Send error callback for /upload workflow
        Gửi error callback cho luồng /upload
        """
        processing_time = time.time() - getattr(task, "_start_time", time.time())

        await self.send_callback(
            task,
            "failed",
            success=False,
            error_message=error_message,
            processing_time=processing_time,
        )

    async def _upload_extraction_to_qdrant(
        self, ingestion_data: dict, task: DocumentProcessingTask
    ):
        """Upload extraction results to Qdrant using ingestion data."""
        try:
            logger.info(f"💾 Uploading extraction chunks to Qdrant...")

            # Initialize Qdrant service
            qdrant_service = QdrantCompanyDataService(
                qdrant_url=os.getenv("QDRANT_URL"),
                qdrant_api_key=os.getenv("QDRANT_API_KEY"),
            )

            # Convert ingestion data back to QdrantDocumentChunk objects
            from src.models.unified_models import (
                QdrantDocumentChunk,
                IndustryDataType,
                Industry,
                Language,
            )

            qdrant_chunks = []
            for chunk_data in ingestion_data.get("qdrant_chunks", []):
                chunk = QdrantDocumentChunk(
                    chunk_id=chunk_data["chunk_id"],
                    company_id=chunk_data["company_id"],
                    file_id=chunk_data["file_id"],
                    content=chunk_data["content"],
                    content_for_embedding=chunk_data.get(
                        "content_for_embedding", chunk_data["content"]
                    ),
                    content_type=IndustryDataType(chunk_data["content_type"]),
                    structured_data=chunk_data["structured_data"],
                    language=Language(chunk_data["language"]),
                    industry=Industry(chunk_data["industry"]),
                    location=chunk_data["location"],
                    valid_from=(
                        datetime.fromisoformat(chunk_data["valid_from"])
                        if chunk_data["valid_from"]
                        else None
                    ),
                    valid_until=(
                        datetime.fromisoformat(chunk_data["valid_until"])
                        if chunk_data["valid_until"]
                        else None
                    ),
                    created_at=datetime.fromisoformat(chunk_data["created_at"]),
                    updated_at=datetime.fromisoformat(chunk_data["updated_at"]),
                )
                qdrant_chunks.append(chunk)

            # Add chunks to Qdrant
            if qdrant_chunks:
                result = await qdrant_service.add_document_chunks(
                    chunks=qdrant_chunks, company_id=task.metadata.get("company_id")
                )

                logger.info("✅ Successfully uploaded extraction chunks to Qdrant")
                logger.info(f"   📊 Upload result: {result.get('status')}")
                logger.info(f"   📈 Points added: {result.get('points_added', 0)}")
                logger.info(f"   🎯 Total chunks: {len(qdrant_chunks)}")

        except Exception as e:
            logger.error(f"❌ Failed to upload extraction to Qdrant: {str(e)}")
            raise

    def _create_intelligent_chunks(self, content: str) -> List[str]:
        """
        Create intelligent chunks from raw content for better vector search.

        Args:
            content: Raw text content to chunk

        Returns:
            List of content chunks
        """
        if not content or not content.strip():
            return []

        # ✅ SMART CONTENT-AWARE Configuration for chunking
        # Tối ưu cho content thực tế (loại bỏ nhiều whitespace)
        chunk_size = 2500  # ~1.5-2 trang nội dung thực
        min_chunk_size = 800  # ~0.5 trang minimum
        overlap = 200  # Overlap nhẹ cho context
        no_chunk_threshold = 2000  # Dưới 2000 chars không cần tách chunk

        chunks = []
        content = content.strip()

        # Clean excessive whitespace first
        import re

        content = re.sub(
            r"\n\s*\n\s*\n+", "\n\n", content
        )  # Remove excessive line breaks
        content = re.sub(r" {3,}", "  ", content)  # Reduce multiple spaces

        logger.info(
            f"📄 Content cleaned: {len(content)} characters after whitespace optimization"
        )

        # ✅ IMPORTANT: For small content (< 2000 chars), return as single chunk
        # Avoid unnecessary chunking overhead for short documents
        if len(content) <= no_chunk_threshold:
            logger.info(
                f"📄 Small content ({len(content)} chars) - keeping as single chunk"
            )
            return [content]
            return [content]

        # ✅ SMART CHUNKING: Try to split by content categories first
        category_chunks = self._split_by_categories(content, chunk_size)
        if category_chunks and len(category_chunks) > 1:
            logger.info(
                f"✅ Smart category-based chunking: {len(category_chunks)} sections"
            )
            return category_chunks

        # Fallback to standard chunking if no clear categories
        start = 0
        while start < len(content):
            # Calculate end position
            end = start + chunk_size

            # If this is not the last chunk, try to find good breaking points
            if end < len(content):
                # Look for sentence boundaries within the last 200 characters
                search_start = max(end - 200, start)

                # ✅ ENHANCED Priority order for breaking points
                break_points = []

                # 1. Look for section headers (uppercase lines, numbered sections)
                for i in range(search_start, end):
                    line_start = content.rfind("\n", 0, i) + 1
                    line_end = content.find("\n", i)
                    if line_end == -1:
                        line_end = len(content)

                    line = content[line_start:line_end].strip()
                    if self._is_section_header(line):
                        break_points.append((line_start, 5))  # Highest priority

                # 2. Look for paragraph breaks (double newlines)
                para_break = content.rfind("\n\n", search_start, end)
                if para_break != -1:
                    break_points.append((para_break + 2, 4))  # High priority

                # 3. Look for list items or numbered points
                for pattern in ["\n• ", "\n- ", "\n\d+\. ", "\n\d+\) "]:
                    import re

                    matches = list(re.finditer(pattern, content[search_start:end]))
                    for match in matches:
                        pos = search_start + match.start() + 1
                        break_points.append((pos, 3))  # Medium-high priority

                # 4. Look for sentence endings
                for punct in [". ", "! ", "? "]:
                    sent_break = content.rfind(punct, search_start, end)
                    if sent_break != -1:
                        break_points.append(
                            (sent_break + len(punct), 2)
                        )  # Medium priority

                # 5. Look for single newlines
                line_break = content.rfind("\n", search_start, end)
                if line_break != -1:
                    break_points.append((line_break + 1, 1))  # Low priority

                # Choose the best breaking point
                if break_points:
                    # Sort by priority (descending) then by position (descending)
                    break_points.sort(key=lambda x: (x[1], x[0]), reverse=True)
                    end = break_points[0][0]

            # Extract chunk
            chunk_text = content[start:end].strip()

            # Only add non-empty chunks that meet minimum size
            if chunk_text and len(chunk_text) >= min_chunk_size:
                chunks.append(chunk_text)

            # Move start position with overlap for context continuity
            if end < len(content):
                start = max(end - overlap, start + min_chunk_size)
            else:
                break

        # Ensure we don't have empty chunks
        chunks = [
            chunk
            for chunk in chunks
            if chunk.strip() and len(chunk.strip()) >= min_chunk_size
        ]

        logger.info(f"📊 Intelligent chunking completed:")
        logger.info(f"   📄 Original content: {len(content)} characters")
        logger.info(f"   📦 Generated chunks: {len(chunks)}")
        if chunks:
            avg_size = sum(len(chunk) for chunk in chunks) / len(chunks)
            logger.info(f"   📏 Average chunk size: {avg_size:.0f} characters")
            logger.info(
                f"   📐 Size range: {min(len(chunk) for chunk in chunks)} - {max(len(chunk) for chunk in chunks)} chars"
            )

        return chunks

    def _split_by_categories(self, content: str, max_chunk_size: int) -> List[str]:
        """
        Split content by logical categories/sections
        """
        import re

        # Look for common section patterns
        section_patterns = [
            r"\n[A-ZÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÊỀỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴĐ][A-ZÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÊỀỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴĐ\s]{5,}:?\n",  # UPPERCASE sections
            r"\n\d+[\.\)]\s+[A-ZÀÁẢÃẠ][^\n]{10,}\n",  # Numbered sections
            r"\n[IVX]+[\.\)]\s+[A-ZÀÁẢÃẠ][^\n]{5,}\n",  # Roman numerals
            r"\n#+\s+[^\n]+\n",  # Markdown headers
        ]

        sections = []
        current_pos = 0

        for pattern in section_patterns:
            matches = list(re.finditer(pattern, content, re.MULTILINE | re.IGNORECASE))
            if matches and len(matches) > 1:  # Need at least 2 sections
                logger.info(
                    f"🎯 Found {len(matches)} sections with pattern: {pattern[:20]}..."
                )

                for i, match in enumerate(matches):
                    start = current_pos if i == 0 else matches[i - 1].end()
                    end = match.start() if i < len(matches) - 1 else len(content)

                    section_content = content[start:end].strip()
                    if len(section_content) >= 500:  # Minimum meaningful section
                        # If section is too long, split it further
                        if len(section_content) <= max_chunk_size:
                            sections.append(section_content)
                        else:
                            # Recursively split large sections
                            sub_sections = self._split_large_section(
                                section_content, max_chunk_size
                            )
                            sections.extend(sub_sections)

                return sections if len(sections) > 1 else []

        return []  # No clear sections found

    def _split_large_section(self, section: str, max_size: int) -> List[str]:
        """Split a large section into smaller chunks"""
        if len(section) <= max_size:
            return [section]

        chunks = []
        start = 0
        while start < len(section):
            end = start + max_size
            if end < len(section):
                # Find best break point
                para_break = section.rfind("\n\n", start, end)
                if para_break > start:
                    end = para_break
                else:
                    sent_break = section.rfind(". ", start, end)
                    if sent_break > start:
                        end = sent_break + 2

            chunk = section[start:end].strip()
            if chunk and len(chunk) >= 500:
                chunks.append(chunk)

            start = end

        return chunks

    def _is_section_header(self, line: str) -> bool:
        """Check if a line looks like a section header"""
        if not line or len(line.strip()) < 3:
            return False

        line = line.strip()

        # Check common header patterns
        import re

        header_patterns = [
            r"^[A-ZÀÁẢÃẠ][A-ZÀÁẢÃẠĂẮẰẲẴẶÂẤẦẨẪẬÊỀỀỂỄỆÍÌỈĨỊÓÒỎÕỌÔỐỒỔỖỘƠỚỜỞỠỢÚÙỦŨỤƯỨỪỬỮỰÝỲỶỸỴĐ\s]{3,}:?$",  # ALL CAPS
            r"^\d+[\.\)]\s+[A-ZÀÁẢÃẠ]",  # Numbered: "1. Title"
            r"^[IVX]+[\.\)]\s+[A-ZÀÁẢÃẠ]",  # Roman: "I. Title"
            r"^#+\s+",  # Markdown headers
            r"^[A-ZÀÁẢÃẠ][^\n]{5,}:$",  # Title followed by colon
        ]

        for pattern in header_patterns:
            if re.match(pattern, line, re.IGNORECASE):
                return True

        return False

    async def upload_to_qdrant(
        self, task: DocumentProcessingTask, extraction_result: dict
    ):
        """Upload extracted content to Qdrant"""
        try:
            logger.info(f"📤 Uploading to Qdrant for task {task.task_id}")

            # Convert string enums to proper enum types
            try:
                industry_enum = (
                    Industry(task.industry)
                    if isinstance(task.industry, str)
                    else task.industry
                )
                language_enum = (
                    Language(task.language)
                    if isinstance(task.language, str)
                    else task.language
                )
            except ValueError as e:
                logger.warning(f"⚠️ Enum conversion error: {e}, using defaults")
                industry_enum = Industry.OTHER
                language_enum = Language.AUTO_DETECT

            # Use unified collection name consistently
            UNIFIED_COLLECTION_NAME = "multi_company_data"

            # Ensure collection exists with correct vector dimensions
            try:
                collections = self.qdrant_service.client.get_collections()
                existing_names = [col.name for col in collections.collections]

                if UNIFIED_COLLECTION_NAME not in existing_names:
                    # Create new collection
                    from qdrant_client.models import Distance, VectorParams

                    self.qdrant_service.client.create_collection(
                        collection_name=UNIFIED_COLLECTION_NAME,
                        vectors_config=VectorParams(
                            size=config.VECTOR_SIZE, distance=Distance.COSINE
                        ),
                    )
                    logger.info(
                        f"✅ Created unified collection: {UNIFIED_COLLECTION_NAME} with {config.VECTOR_SIZE} dimensions"
                    )
                else:
                    # Collection exists - check dimensions with error handling
                    try:
                        collection_info = self.qdrant_service.client.get_collection(
                            UNIFIED_COLLECTION_NAME
                        )
                        existing_vector_size = (
                            collection_info.config.params.vectors.size
                        )

                        if existing_vector_size != config.VECTOR_SIZE:
                            logger.warning(
                                f"⚠️ Dimension mismatch! Collection has {existing_vector_size}, config requires {config.VECTOR_SIZE}"
                            )
                            logger.info(
                                f"🔄 Recreating collection with correct dimensions..."
                            )

                            # Delete and recreate with correct dimensions
                            self.qdrant_service.client.delete_collection(
                                UNIFIED_COLLECTION_NAME
                            )
                            logger.info(
                                f"🗑️ Deleted old collection: {UNIFIED_COLLECTION_NAME}"
                            )

                            from qdrant_client.models import Distance, VectorParams

                            self.qdrant_service.client.create_collection(
                                collection_name=UNIFIED_COLLECTION_NAME,
                                vectors_config=VectorParams(
                                    size=config.VECTOR_SIZE, distance=Distance.COSINE
                                ),
                            )

                            # Create index for file_id to enable efficient filtering
                            from qdrant_client.models import (
                                CreateFieldIndex,
                                FieldCondition,
                                PayloadSchemaType,
                            )

                            try:
                                # Create index for file_id
                                self.qdrant_service.client.create_payload_index(
                                    collection_name=UNIFIED_COLLECTION_NAME,
                                    field_name="file_id",
                                    field_schema=PayloadSchemaType.KEYWORD,
                                )
                                logger.info(
                                    f"✅ Created file_id index for collection: {UNIFIED_COLLECTION_NAME}"
                                )

                                # Create index for company_id
                                self.qdrant_service.client.create_payload_index(
                                    collection_name=UNIFIED_COLLECTION_NAME,
                                    field_name="company_id",
                                    field_schema=PayloadSchemaType.KEYWORD,
                                )
                                logger.info(
                                    f"✅ Created company_id index for collection: {UNIFIED_COLLECTION_NAME}"
                                )
                            except Exception as index_error:
                                logger.warning(
                                    f"⚠️ Failed to create file_id index: {index_error}"
                                )

                            logger.info(
                                f"✅ Recreated collection: {UNIFIED_COLLECTION_NAME} with {config.VECTOR_SIZE} dimensions"
                            )
                        else:
                            logger.info(
                                f"✅ Using existing collection: {UNIFIED_COLLECTION_NAME} with {config.VECTOR_SIZE} dimensions"
                            )
                    except Exception as collection_info_error:
                        # If we can't get collection info (version mismatch), assume it needs recreation
                        logger.warning(
                            f"⚠️ Failed to get collection info (likely client version mismatch): {collection_info_error}"
                        )
                        logger.info(
                            f"🔄 Recreating collection to ensure compatibility..."
                        )

                        try:
                            # Delete existing collection
                            self.qdrant_service.client.delete_collection(
                                UNIFIED_COLLECTION_NAME
                            )
                            logger.info(
                                f"🗑️ Deleted existing collection: {UNIFIED_COLLECTION_NAME}"
                            )
                        except:
                            pass  # Collection might not exist or deletion failed

                        # Create new collection
                        from qdrant_client.models import Distance, VectorParams

                        self.qdrant_service.client.create_collection(
                            collection_name=UNIFIED_COLLECTION_NAME,
                            vectors_config=VectorParams(
                                size=config.VECTOR_SIZE, distance=Distance.COSINE
                            ),
                        )
                        logger.info(
                            f"✅ Recreated collection: {UNIFIED_COLLECTION_NAME} with {config.VECTOR_SIZE} dimensions"
                        )

            except Exception as collection_error:
                logger.error(f"❌ Failed to ensure collection: {collection_error}")
                # If there's a collection error, try to recreate it
                logger.info(f"🔄 Attempting to recreate collection due to error...")
                try:
                    # Delete existing collection if it exists
                    try:
                        self.qdrant_service.client.delete_collection(
                            UNIFIED_COLLECTION_NAME
                        )
                        logger.info(
                            f"🗑️ Deleted problematic collection: {UNIFIED_COLLECTION_NAME}"
                        )
                    except:
                        pass  # Collection might not exist

                    # Create new collection with correct dimensions
                    from qdrant_client.models import Distance, VectorParams

                    self.qdrant_service.client.create_collection(
                        collection_name=UNIFIED_COLLECTION_NAME,
                        vectors_config=VectorParams(
                            size=config.VECTOR_SIZE, distance=Distance.COSINE
                        ),
                    )

                    # Create index for file_id to enable efficient filtering
                    from qdrant_client.models import (
                        CreateFieldIndex,
                        FieldCondition,
                        PayloadSchemaType,
                    )

                    try:
                        # Create index for file_id
                        self.qdrant_service.client.create_payload_index(
                            collection_name=UNIFIED_COLLECTION_NAME,
                            field_name="file_id",
                            field_schema=PayloadSchemaType.KEYWORD,
                        )
                        logger.info(
                            f"✅ Created file_id index for collection: {UNIFIED_COLLECTION_NAME}"
                        )

                        # Create index for company_id
                        self.qdrant_service.client.create_payload_index(
                            collection_name=UNIFIED_COLLECTION_NAME,
                            field_name="company_id",
                            field_schema=PayloadSchemaType.KEYWORD,
                        )
                        logger.info(
                            f"✅ Created company_id index for collection: {UNIFIED_COLLECTION_NAME}"
                        )
                    except Exception as index_error:
                        logger.warning(f"⚠️ Failed to create indices: {index_error}")

                    logger.info(
                        f"✅ Recreated collection: {UNIFIED_COLLECTION_NAME} with {config.VECTOR_SIZE} dimensions"
                    )
                except Exception as recreate_error:
                    logger.error(f"❌ Failed to recreate collection: {recreate_error}")
                    raise

            # Create document chunks from extracted content
            raw_content = extraction_result.get("raw_content", "")
            if not raw_content or not raw_content.strip():
                logger.warning(f"⚠️ No content to upload for task {task.task_id}")
                return

            # Intelligent chunking: 2-3 pages per chunk or by content categories
            content_chunks = self._create_intelligent_chunks(raw_content)
            logger.info(f"📊 Created {len(content_chunks)} intelligent chunks")

            # ✅ OPTIMIZATION: Generate embeddings in batches to avoid memory issues
            try:
                logger.info(
                    f"🧠 Generating embeddings for {len(content_chunks)} chunks in batches..."
                )

                # Generate all embeddings at once using batch processing
                embeddings = await self.ai_manager.generate_embeddings_batch(
                    texts=content_chunks,
                    max_batch_size=10,  # Process 10 chunks at a time to avoid memory issues
                    timeout_seconds=600,  # 10 minute timeout for entire operation
                )

                if len(embeddings) != len(content_chunks):
                    logger.warning(
                        f"⚠️ Embedding count mismatch: {len(embeddings)} vs {len(content_chunks)}"
                    )

                logger.info(f"✅ Successfully generated {len(embeddings)} embeddings")

            except Exception as batch_embedding_error:
                logger.error(
                    f"❌ Batch embedding generation failed: {batch_embedding_error}"
                )
                # Fallback to zero vectors for all chunks
                zero_dimension = 768  # Default dimension
                embeddings = [[0.0] * zero_dimension] * len(content_chunks)
                logger.info(f"🔄 Using {len(embeddings)} zero vector fallbacks")

            # Create points for Qdrant using pre-generated embeddings
            file_chunks = []
            for i, (chunk_content, embedding) in enumerate(
                zip(content_chunks, embeddings)
            ):
                try:
                    # Generate UUID for point ID (Qdrant requirement)
                    import uuid

                    point_id = str(uuid.uuid4())

                    point_data = {
                        "id": point_id,
                        "payload": {
                            "file_id": task.metadata.get("file_id"),
                            "company_id": task.company_id,
                            "content": chunk_content,
                            "content_type": "file_document",
                            "data_type": task.data_type,
                            "industry": industry_enum.value,
                            "language": language_enum.value,
                            "tags": task.metadata.get("tags", []),
                            "original_name": task.metadata.get("original_name"),
                            "file_name": task.metadata.get(
                                "file_name", task.metadata.get("original_name")
                            ),
                            "file_size": task.metadata.get("file_size"),
                            "file_type": task.metadata.get("file_type"),
                            "uploaded_by": task.metadata.get("uploaded_by"),
                            "description": task.metadata.get("description"),
                            "ai_provider": extraction_result.get(
                                "extraction_metadata", {}
                            ).get("ai_provider"),
                            "chunk_index": i,
                            "total_chunks": len(content_chunks),
                            "chunk_size": len(chunk_content),
                            "chunking_method": "intelligent_structural",
                            "estimated_pages": round(len(chunk_content) / 1500, 1),
                            "created_at": task.created_at,
                            "r2_url": task.r2_url,
                        },
                        "vector": embedding,
                    }
                    file_chunks.append(point_data)

                except Exception as embedding_error:
                    logger.warning(
                        f"⚠️ Failed to create embedding for chunk {i}: {embedding_error}"
                    )
                    # Create point without embedding as fallback
                    # Generate UUID for point ID (Qdrant requirement)
                    fallback_point_id = str(uuid.uuid4())

                    point_data = {
                        "id": fallback_point_id,
                        "payload": {
                            "file_id": task.metadata.get("file_id"),
                            "company_id": task.company_id,
                            "content": chunk_content,
                            "content_type": "file_document",
                            "data_type": task.data_type,
                            "industry": industry_enum.value,
                            "language": language_enum.value,
                            "tags": task.metadata.get("tags", []),
                            "original_name": task.metadata.get("original_name"),
                            "file_name": task.metadata.get(
                                "file_name", task.metadata.get("original_name")
                            ),
                            "file_size": task.metadata.get("file_size"),
                            "file_type": task.metadata.get("file_type"),
                            "uploaded_by": task.metadata.get("uploaded_by"),
                            "description": task.metadata.get("description"),
                            "ai_provider": extraction_result.get(
                                "extraction_metadata", {}
                            ).get("ai_provider"),
                            "chunk_index": i,
                            "total_chunks": len(content_chunks),
                            "chunk_size": len(chunk_content),
                            "chunking_method": "intelligent_structural",
                            "estimated_pages": round(len(chunk_content) / 1500, 1),
                            "created_at": task.created_at,
                            "r2_url": task.r2_url,
                            "embedding_failed": True,
                        },
                    }
                    file_chunks.append(point_data)

            # Upload to Qdrant
            if file_chunks:
                logger.info(f"📤 Uploading {len(file_chunks)} chunks to Qdrant")
                try:
                    await self.qdrant_service.upsert_points(
                        UNIFIED_COLLECTION_NAME, file_chunks
                    )
                    logger.info(f"✅ Successfully uploaded {len(file_chunks)} chunks")

                    # Update extraction result with chunk count
                    extraction_result["chunks_created"] = len(file_chunks)
                except Exception as upload_error:
                    error_msg = str(upload_error).lower()
                    if (
                        "vector dimension error" in error_msg
                        or "expected dim" in error_msg
                    ):
                        logger.warning(
                            f"⚠️ Vector dimension mismatch detected. Recreating collection..."
                        )

                        # Delete and recreate collection with correct dimensions
                        try:
                            self.qdrant_service.client.delete_collection(
                                UNIFIED_COLLECTION_NAME
                            )
                            logger.info(
                                f"🗑️ Deleted collection with wrong dimensions: {UNIFIED_COLLECTION_NAME}"
                            )

                            from qdrant_client.models import Distance, VectorParams

                            self.qdrant_service.client.create_collection(
                                collection_name=UNIFIED_COLLECTION_NAME,
                                vectors_config=VectorParams(
                                    size=config.VECTOR_SIZE, distance=Distance.COSINE
                                ),
                            )
                            logger.info(
                                f"✅ Recreated collection: {UNIFIED_COLLECTION_NAME} with {config.VECTOR_SIZE} dimensions"
                            )

                            # Retry upload
                            await self.qdrant_service.upsert_points(
                                UNIFIED_COLLECTION_NAME, file_chunks
                            )
                            logger.info(
                                f"✅ Successfully uploaded {len(file_chunks)} chunks after recreation"
                            )
                            extraction_result["chunks_created"] = len(file_chunks)

                        except Exception as retry_error:
                            logger.error(
                                f"❌ Failed to recreate and upload: {retry_error}"
                            )
                            raise
                    else:
                        logger.error(f"❌ Upload failed: {upload_error}")
                        raise
            else:
                logger.warning("⚠️ No valid chunks to upload")

        except Exception as e:
            logger.error(f"❌ Failed to upload to Qdrant: {e}")
            raise

    def _get_webhook_secret(self) -> str:
        """
        Get webhook secret for plain text authentication
        Lấy webhook secret cho xác thực văn bản thuần túy
        """
        return os.getenv("WEBHOOK_SECRET", "webhook-secret-for-signature")

    # ==============================================================================
    # TASK STATUS AND CALLBACK SERVICES (Copy from extraction_routes.py)
    # ==============================================================================

    class TaskStatusService:
        """Service to manage task status and results"""

        def __init__(self, redis_client):
            self.redis = redis_client
            self.task_prefix = "task_status:"
            self.result_prefix = "task_result:"
            self.ttl = 3600  # 1 hour TTL for task data

        async def update_task_status(
            self, task_id: str, status: str, progress: Optional[dict] = None
        ) -> bool:
            """Update task status"""
            try:
                key = f"{self.task_prefix}{task_id}"

                # Get existing data
                existing_data = await asyncio.get_event_loop().run_in_executor(
                    None, self.redis.get, key
                )

                if not existing_data:
                    return False

                task_data = json.loads(existing_data)
                task_data["status"] = status

                if progress:
                    task_data["progress"] = progress

                if status in ["completed", "failed"]:
                    task_data["completed_at"] = datetime.now().isoformat()

                # Update with TTL
                await asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: self.redis.setex(key, self.ttl, json.dumps(task_data)),
                )
                return True
            except Exception as e:
                logger.error(f"Failed to update task status for {task_id}: {e}")
                return False

        async def store_task_result(self, task_id: str, result_data: dict) -> bool:
            """Store task result"""
            try:
                key = f"{self.result_prefix}{task_id}"

                # Add timestamp
                result_data["stored_at"] = datetime.now().isoformat()

                # Store with TTL
                await asyncio.get_event_loop().run_in_executor(
                    None,
                    lambda: self.redis.setex(key, self.ttl, json.dumps(result_data)),
                )

                # Update task status to completed
                await self.update_task_status(task_id, "completed")
                return True
            except Exception as e:
                logger.error(f"Failed to store task result for {task_id}: {e}")
                return False


class CallbackService:
    """Service to handle callback notifications"""

    @staticmethod
    def _get_webhook_secret() -> str:
        """Get webhook secret for plain text authentication"""
        return os.getenv("WEBHOOK_SECRET", "webhook-secret-for-signature")

    @staticmethod
    async def send_callback(
        callback_url: str, task_id: str, status: str, data: Optional[dict] = None
    ):
        """Send callback notification to backend"""
        if not callback_url:
            return

        try:
            callback_payload = {
                "task_id": task_id,
                "status": status,
                "timestamp": datetime.now().isoformat(),
                "data": data or {},
            }

            logger.info(f"📞 Sending callback for {task_id} to {callback_url}")

            # Generate webhook signature
            webhook_secret = os.getenv("WEBHOOK_SECRET", "webhook-secret-for-signature")

            headers = {
                "Content-Type": "application/json",
                "X-Webhook-Source": "ai-service",
                "X-Webhook-Secret": webhook_secret,  # ✅ Simplified: Use plain text secret
                "User-Agent": "Agent8x-AI-Service/1.0",
            }

            # Send async HTTP request
            timeout = aiohttp.ClientTimeout(total=10)  # 10 second timeout

            async with aiohttp.ClientSession(timeout=timeout) as session:
                async with session.post(
                    callback_url,
                    json=callback_payload,
                    headers=headers,
                ) as response:
                    if response.status == 200:
                        logger.info(f"✅ Callback sent successfully for {task_id}")
                    else:
                        logger.warning(
                            f"⚠️ Callback returned status {response.status} for {task_id}"
                        )

        except Exception as e:
            logger.error(f"❌ Failed to send callback for {task_id}: {e}")

    @staticmethod
    async def send_error_callback(
        callback_url: str,
        task_id: str,
        error_message: str,
        error_details: Optional[dict] = None,
    ):
        """Send error callback"""
        error_data = {
            "error": error_message,
            "error_details": error_details or {},
            "failed_at": datetime.now().isoformat(),
        }

        callback_service = CallbackService()
        await callback_service.send_callback(
            callback_url, task_id, "failed", error_data
        )


async def start_document_worker():
    """Start the document processing worker"""
    worker = DocumentProcessingWorker()

    # Setup graceful shutdown
    loop = asyncio.get_event_loop()
    for sig in (signal.SIGTERM, signal.SIGINT):
        loop.add_signal_handler(sig, lambda: asyncio.create_task(worker.shutdown()))

    try:
        await worker.initialize()
        await worker.run()
    except Exception as e:
        logger.error(f"❌ Worker failed: {e}")
        await worker.shutdown()
        raise


if __name__ == "__main__":
    """Run worker standalone"""
    try:
        asyncio.run(start_document_worker())
    except KeyboardInterrupt:
        logger.info("🛑 Worker stopped by user")
    except Exception as e:
        logger.error(f"❌ Worker startup failed: {e}")
