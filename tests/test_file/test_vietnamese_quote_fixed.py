#!/usr/bin/env python3
"""
Test Vietnamese quote gen        "company_info": {
            "name": "CÔNG TY TNHH ÁI VŨNG TÀU",
            "address": "123 Đường Lê Lợi, Phường 7, TP. Vũng Tàu, Bà Rịa - Vũng Tàu",
            "phone": "0254.123.4567",
            "email": "info@aivungtau.com",
            "website": "https://aivungtau.com",
            "tax_code": "0123456789",  # Changed from tax_id to tax_code
            "representative": "Nguyễn Văn B",  # Added required field
            "position": "Giám đốc",
            "logo": "https://example.com/logo.png",  # New field
            "city": "Vũng Tàu",  # New field
            "fax": "0254.123.4568",  # New field
            "social_link": "https://facebook.com/aivungtau"  # New field
        }, new enhanced format
"""

import asyncio
import requests
import json
from pathlib import Path

# Firebase token for authentication
FIREBASE_TOKEN = "eyJhbGciOiJSUzI1NiIsImtpZCI6IjJkNzU5MjUzMjFkZTczNGY1ZGNlOTc2ZmY4MDEyMzc0YjZmZjE3ODciLCJ0eXAiOiJKV1QifQ.eyJuYW1lIjoiTWljaGFlbCBMZSIsInBpY3R1cmUiOiJodHRwczovL2xoMy5nb29nbGV1c2VyY29udGVudC5jb20vYS9BQ2c4b2NKWGp5RTVoeWQ4Wk9ib0hmTVBRLTJNVjlSTnhHNmczYU9uNzA5UWdGRXpmdWRDPXM5Ni1jIiwiaXNzIjoiaHR0cHM6Ly9zZWN1cmV0b2tlbi5nb29nbGUuY29tL2FpdnVuZ3RhdS1wcm9qZWN0IiwiYXVkIjoiYWl2dW5ndGF1LXByb2plY3QiLCJhdXRoX3RpbWUiOjE3MzQ4OTc1NDMsInVzZXJfaWQiOiJ4UXNPRGMxVGNxZGJNYVV1TVBDbWc0cUlSdjIyIiwic3ViIjoieFFzT0RjMVRjcWRiTWFVdU1QQ21nNHFJUnYyMiIsImlhdCI6MTczNDg5NzU0MywiZXhwIjoxNzM0OTAxMTQzLCJlbWFpbCI6InRpZW5ob2kubGhAZ21haWwuY29tIiwiZW1haWxfdmVyaWZpZWQiOnRydWUsImZpcmViYXNlIjp7ImlkZW50aXRpZXMiOnsiZ29vZ2xlLmNvbSI6WyIxMDg5MTkxNzU4NDU4MjgwNTk1MDIiXSwiZW1haWwiOlsidGllbmhvaS5saEBnbWFpbC5jb20iXX0sInNpZ25faW5fcHJvdmlkZXIiOiJnb29nbGUuY29tIn19.NL4wnWe9HnbLHcBSuNu8CjwwcFJMPBKWP2OKCjqIiuvvF5c4E7wdRzGOPLNHBa2vFRfQq9AjUnGC5jWCwMFBxbMBXnWQFDgFDSN9LVr6GvpDdMnr7v_fCWUhp6J9xPEKReCzE-V7XbmNO-sSFSqkz4PVHe1gHUcHFe0dQ5tAhxmSSGp8kQh_vA0LfAwKV3c2A3WKwSUpLZ0_MsqhKPqmQQHtNFgfWQzFmeTGP9xvxPuiZABUKcJm1Vz7PTdp9h9B1n4AqZKXcGI4iIbZJ0jO_ADFX8nHPJe6d-bNLZNfbK9xmFe8K-VJXkrECjSb5F5SJaWZMXGUJbWXVZq5Sg"

BASE_URL = "http://localhost:8000"


async def test_vietnamese_quote():
    """Test creating a Vietnamese quote with enhanced format"""

    # First, get available templates
    print("📋 Getting available templates...")
    response = requests.get(
        f"{BASE_URL}/api/documents/templates",
        headers={"Authorization": f"Bearer {FIREBASE_TOKEN}"},
    )

    if response.status_code != 200:
        print(f"❌ Failed to get templates: {response.status_code}")
        print(response.text)
        return

    response_data = response.json()
    templates = response_data.get("templates", [])
    print(f"✅ Found {len(templates)} templates")

    # Find a quote template
    quote_template = None
    for template in templates:
        print(f"🔍 Checking template: {template.get('name', 'Unknown')}")
        if template.get("type") == "quote":
            quote_template = template
            break

    if not quote_template:
        print("❌ No quote template found")
        return

    print(f"📄 Using template: {quote_template['name']}")

    # Prepare quote request with Vietnamese data and enhanced fields
    quote_data = {
        "template_id": quote_template["_id"],
        "company_info": {
            "name": "CÔNG TY TNHH ÁI VŨNG TÀU",
            "address": "123 Đường Lê Lợi, Phường 7, TP. Vũng Tàu, Bà Rịa - Vũng Tàu",
            "phone": "0254.123.4567",
            "email": "info@aivungtau.com",
            "website": "https://aivungtau.com",
            "tax_id": "0123456789",
            "logo": "https://example.com/logo.png",  # New field
            "city": "Vũng Tàu",  # New field
            "fax": "0254.123.4568",  # New field
            "social_link": "https://facebook.com/aivungtau",  # New field
        },
        "customer_info": {
            "name": "CÔNG TY TNHH ABC TRADING",
            "contact_person": "Nguyễn Văn A",
            "address": "456 Đường Nguyễn Huệ, Quận 1, TP. Hồ Chí Minh",
            "phone": "028.123.4567",
            "email": "contact@abctrading.com",
            "tax_id": "9876543210",
        },
        "products": [
            {
                "name": "Máy tính Dell Inspiron 15",
                "description": "Laptop Dell Inspiron 15 3000 series với cấu hình mạnh mẽ",
                "quantity": 5,
                "unit": "chiếc",
                "unit_price": 15000000,
                "total_price": 75000000,
                "specifications": {
                    "CPU": "Intel Core i5-1135G7",
                    "RAM": "8GB DDR4",
                    "Storage": "256GB SSD",
                    "Display": "15.6 inch FHD",
                },
                "warranty_period": "24 tháng",
                "delivery_time": "7-10 ngày làm việc",
            },
            {
                "name": "Màn hình Samsung 24 inch",
                "description": "Màn hình LED Samsung 24 inch độ phân giải Full HD",
                "quantity": 3,
                "unit": "cái",
                "unit_price": 4500000,
                "total_price": 13500000,
                "specifications": {
                    "Size": "24 inch",
                    "Resolution": "1920x1080",
                    "Panel": "IPS",
                    "Response Time": "5ms",
                },
                "warranty_period": "36 tháng",
                "delivery_time": "3-5 ngày làm việc",
            },
        ],
        "total_amount": 88500000,
        "total_quantity": 8,
        "payment_terms": {
            "method": "Chuyển khoản ngân hàng",
            "term": "Thanh toán 50% trước khi giao hàng, 50% sau khi nghiệm thu",
            "due_days": 30,
            "bank_info": {
                "bank_name": "Ngân hàng Vietcombank",
                "account_number": "0123456789",
                "account_name": "CÔNG TY TNHH ÁI VŨNG TÀU",
            },
        },
        "additional_terms": "Giá đã bao gồm vận chuyển trong nội thành TP.HCM. Bảo hành chính hãng.",
        "vat_rate": 10.0,  # New field for VAT
        "notes": "Báo giá có hiệu lực trong 30 ngày. Liên hệ để được tư vấn chi tiết.",  # New field
    }

    print("💾 Creating Vietnamese quote with enhanced format...")
    response = requests.post(
        f"{BASE_URL}/api/documents/quotes",
        headers={
            "Authorization": f"Bearer {FIREBASE_TOKEN}",
            "Content-Type": "application/json",
        },
        json=quote_data,
    )

    if response.status_code != 200:
        print(f"❌ Failed to create quote: {response.status_code}")
        print(response.text)
        return

    result = response.json()
    print(f"✅ Quote created successfully!")
    print(f"📄 Document ID: {result['document_id']}")
    print(f"📝 File path: {result['file_path']}")
    print(f"⏱️  Processing time: {result['processing_time']:.3f}s")

    # Download and save the file
    if result.get("file_content"):
        output_path = Path(f"vietnamese_quote_{result['document_id']}.docx")

        # Decode base64 content
        import base64

        file_content = base64.b64decode(result["file_content"])

        with open(output_path, "wb") as f:
            f.write(file_content)

        print(f"💾 File saved to: {output_path}")
        print(f"📊 File size: {len(file_content):,} bytes")

        # Try to open the file for inspection
        try:
            from docx import Document

            doc = Document(output_path)
            print(f"📖 Document contains {len(doc.paragraphs)} paragraphs")

            # Show first few paragraphs for verification
            print("\n📋 Document preview (first 5 paragraphs):")
            for i, para in enumerate(doc.paragraphs[:5]):
                if para.text.strip():
                    print(f"  {i+1}: {para.text[:100]}...")

        except Exception as e:
            print(f"⚠️  Could not inspect document: {e}")

    print("\n🎉 Vietnamese quote test completed!")


if __name__ == "__main__":
    asyncio.run(test_vietnamese_quote())
